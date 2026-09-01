import json
import zipfile
from io import BytesIO, StringIO
from decimal import Decimal
from pathlib import Path
from unittest.mock import patch

from django.contrib.auth import get_user_model
from django.core.cache import cache
from django.test import TestCase
from django.urls import reverse
from docx import Document
from openpyxl import Workbook

from calculator.models import CalculatorSettings, PriceItem
from .models import CatalogCategory, CatalogMatchDecision, CatalogProduct, CatalogSupplier, CatalogSyncRun, ProductionTrainingExample, ProductionTrainingSession, ProductionTrainingTurn, ProductionType, TenderEstimate, TenderKnowledgeSource, TenderSettings
from .catalog import CatalogSyncError, GiftsXmlClient, OasisClient, _category_candidates, _category_retrieval, _expand_category_graph, catalog_candidates_for_line, parse_gifts_catalog, sync_gifts_catalog, sync_gifts_categories, sync_oasis_catalog
from .services import _VisibleTextParser, _apply_catalog_operations, _apply_psodin_calculation, _evaluate_cost_recipe, _format_html_tables, _json_from_model, _knowledge_sources_for_line, _normalize_training_hypothesis, _paper_candidates, _parse_document_decimal, _resolve_line_match, _select_catalog_category_tasks, _select_html_price_quote, _shorten_structured_item_names, _source_text_quality, _strip_shared_item_boilerplate, _technical_source_chunks, _validate_public_url, _verify_catalog_category_tasks, analyze_production_route, analyze_tender_requirements, apply_catalog_candidate, apply_verified_source_quote, build_training_hypothesis, calculate_sheet_imposition, calculate_tender, classify_production_type, detect_tender_document_type, extract_tender_source, inspect_tender_document, recognize_tender_items


class TenderTests(TestCase):
    def setUp(self):
        cache.clear()
        self.user = get_user_model().objects.create_user(username="manager", password="password")
        self.other = get_user_model().objects.create_user(username="other", password="password")
        self.payload = [{"name": "Ручка", "quantity": "10", "nmck_unit": "100", "material_unit": "40", "application_unit": "10", "logistics_unit": "5", "product_url": "https://example.com/item", "comment": "Синяя"}]

    def test_smart_upload_has_separate_nmck_and_technical_fields(self):
        self.client.force_login(self.user)

        response = self.client.get(reverse("tender_home"))

        self.assertContains(response, 'id="tender-ai-nmck-file"')
        self.assertContains(response, 'id="tender-ai-technical-file"')
        self.assertNotContains(response, 'id="tender-tech-modal"')

    def test_pending_nmck_rows_show_technical_status_and_use_duplicate_guard(self):
        self.client.force_login(self.user)

        response = self.client.get(reverse("tender_home"))

        self.assertContains(response, "data-ai-technical-status")
        self.assertContains(response, "function applyTechnicalResultToNmckRows")
        self.assertContains(response, "function findExistingTenderLine")

    def test_line_ai_button_opens_drawer_and_starts_hypothesis_immediately(self):
        self.client.force_login(self.user)

        response = self.client.get(reverse("tender_home"))

        self.assertContains(response, "openRequirements=(index,autoCalculate=false)=>")
        self.assertContains(response, "questionControl.onclick=()=>openRequirements(index,true)")
        self.assertContains(response, "shouldAutoCalculate=autoCalculate&&!info.production")
        self.assertContains(response, "if(shouldAutoCalculate)build.click()")

    def test_catalog_ui_distinguishes_empty_results_from_supplier_failure(self):
        self.client.force_login(self.user)

        response = self.client.get(reverse("tender_home"))

        self.assertContains(response, "catalogSourceStatusHtml")
        self.assertContains(response, "Каталог временно недоступен")
        self.assertContains(response, "Поиск выполнен повторно")

    def test_assistant_calculation_has_sticky_totals_and_scoped_loading_state(self):
        self.client.force_login(self.user)

        content = self.client.get(reverse("tender_home")).content.decode()
        styles = (Path(__file__).resolve().parents[1] / "static" / "core" / "index.css").read_text(encoding="utf-8")

        self.assertIn("data-route-unit-total", content)
        self.assertIn("data-route-order-total", content)
        self.assertIn("function updateRouteToolbar", content)
        self.assertIn("function setProductionBusy", content)
        self.assertIn("aria-busy", content)
        self.assertIn(".tender-production-header { position:sticky", styles)
        self.assertIn(".tender-production-result.is-loading::after", styles)
        self.assertIn(".is-loading-button::before", styles)
        self.assertIn("bottom:0", styles)

    def test_smart_import_finishes_in_the_product_list_for_nmck_or_technical_document(self):
        self.client.force_login(self.user)

        content = self.client.get(reverse("tender_home")).content.decode()

        nmck_apply = content[content.index("document.getElementById('tender-ai-add').onclick"):]
        technical_apply = content[content.index("function applyTechnicalResult(result,sourceName)"):]
        self.assertIn("aiModal.classList.remove('is-open')", nmck_apply)
        self.assertIn("aiModal.classList.remove('is-open')", technical_apply)
        self.assertIn("function isPristineTenderLine", content)

    def test_cost_row_places_comment_and_link_before_assistant_action(self):
        self.client.force_login(self.user)

        content = self.client.get(reverse("tender_home")).content.decode()

        expected_order = "${field('Комментарий','comment','text','Необязательно')}${linkControl()}${questionButton(line)}"
        self.assertIn(expected_order, content)

    def test_assistant_deduplicates_questions_and_highlights_route(self):
        self.client.force_login(self.user)

        response = self.client.get(reverse("tender_home"))
        styles = (Path(__file__).resolve().parents[1] / "static" / "core" / "index.css").read_text(encoding="utf-8")

        self.assertContains(response, "function uniqueAssistantQuestions")
        self.assertIn(".training-dialogue .training-step:nth-child(2) > header", styles)
        self.assertIn(".training-dialogue .training-step:nth-child(4)", styles)
        self.assertNotIn(".training-dialogue .training-step:nth-child(2) { margin:0 -10px", styles)

    def test_line_assistant_button_matches_compact_metric_height(self):
        styles = (Path(__file__).resolve().parents[1] / "static" / "core" / "index.css").read_text(encoding="utf-8")

        self.assertIn(".tender-line-questions.ai-route-button { min-height:27px", styles)

    def test_assistant_dialogue_uses_requested_section_order(self):
        self.client.force_login(self.user)

        content = self.client.get(reverse("tender_home")).content.decode()
        content = content[content.index("function trainingDialogueHtml"):]

        sections = [
            "Полученное ТЗ",
            "Технологический маршрут",
            "Ваши корректировки",
            "Предложенные товары",
            "Нужно уточнить",
            "Обратная связь",
            "Добавить поставщика или источник",
        ]
        positions = [content.index(section) for section in sections]
        self.assertEqual(positions, sorted(positions))

    def test_formula_includes_every_expense_in_roi(self):
        _, result = calculate_tender([{**self.payload[0], **{key: Decimal(self.payload[0][key]) for key in ("quantity", "nmck_unit", "material_unit", "application_unit", "logistics_unit")}}], Decimal("30"), Decimal("100"), Decimal("5"))
        self.assertEqual(result["rrp_total"], Decimal("700.00"))
        self.assertEqual(result["vat"], Decimal("35.00"))
        self.assertEqual(result["all_expenses"], Decimal("685.00"))
        self.assertEqual(result["net_profit"], Decimal("15.00"))
        self.assertEqual(result["roi"], Decimal("2.19"))

    def test_user_can_save_and_open_own_estimate(self):
        self.client.force_login(self.user)
        payload = [{**self.payload[0], "requirements": {"requirements": [{"label": "Материал", "value": "пластик"}], "questions": []}}]
        analysis = {"technical": {"name": "ТЗ.pdf", "matched": 1, "questions": 0}}
        response = self.client.post(reverse("tender_home"), {"tender_number": "123", "name": "Тест", "reduction_percent": "30", "russia_delivery": "100", "lines_json": json.dumps(payload), "document_analysis_json": json.dumps(analysis)})
        self.assertEqual(response.status_code, 302)
        estimate = TenderEstimate.objects.get()
        self.assertEqual(estimate.owner, self.user)
        self.assertEqual(estimate.lines.count(), 1)
        self.assertEqual(estimate.vat_rate_snapshot, Decimal("5.00"))
        self.assertFalse(estimate.summary_snapshot["is_incomplete"])
        self.assertEqual(estimate.document_analysis["technical"]["matched"], 1)
        self.assertEqual(estimate.lines.get().requirements["requirements"][0]["value"], "пластик")

    def test_saved_estimate_list_shows_status_selector_without_draft_exclamation(self):
        TenderEstimate.objects.create(
            owner=self.user,
            tender_number="123",
            name="Тест",
            summary_snapshot={"is_incomplete": True, "net_profit": "1000", "roi": "10"},
        )
        self.client.force_login(self.user)

        response = self.client.get(reverse("tender_home"))

        self.assertContains(response, 'class="saved-estimate__status-form is-draft"')
        self.assertContains(response, 'data-estimate-status-form')
        self.assertNotContains(response, 'onchange="this.form.submit()"')
        self.assertContains(response, 'const data=new FormData(form);')
        self.assertContains(response, 'body:data')
        self.assertContains(response, '<option value="draft" selected>Черновик</option>', html=True)
        self.assertContains(response, '<option value="pending">В ожидании</option>', html=True)
        self.assertContains(response, '<option value="not_participated">Не участвовали</option>', html=True)
        self.assertContains(response, '<option value="lost">Проигран</option>', html=True)
        self.assertContains(response, '<option value="won">Выигран</option>', html=True)
        self.assertNotContains(response, 'class="saved-estimate__draft"')

    def test_user_can_change_own_estimate_status(self):
        estimate = TenderEstimate.objects.create(owner=self.user, tender_number="123", name="Тест")
        updated_at = estimate.updated_at
        self.client.force_login(self.user)

        response = self.client.post(
            reverse("tender_estimate_status", args=[estimate.pk]),
            {"status": "won"},
            HTTP_X_REQUESTED_WITH="XMLHttpRequest",
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["status"], "won")
        self.assertTrue(response.json()["requires_result"])
        estimate.refresh_from_db()
        self.assertEqual(estimate.status, "won")
        self.assertEqual(estimate.updated_at, updated_at)

    def test_result_note_is_shown_for_finished_tender_and_saved(self):
        estimate = TenderEstimate.objects.create(
            owner=self.user,
            tender_number="123",
            name="Тест",
            status=TenderEstimate.LOST,
            result_notes="Победитель снизился на 30%.",
        )
        self.client.force_login(self.user)

        opened = self.client.get(reverse("tender_estimate", args=[estimate.pk]))

        self.assertContains(opened, 'data-tender-result')
        self.assertContains(opened, "Победитель снизился на 30%.")
        self.assertNotContains(opened, 'data-tender-result hidden')

        response = self.client.post(
            reverse("tender_estimate", args=[estimate.pk]),
            {
                "tender_number": "123",
                "name": "Тест",
                "reduction_percent": "20",
                "russia_delivery": "0",
                "result_notes": "Проиграли: победитель снизился на 30%.",
                "lines_json": json.dumps(self.payload),
            },
        )

        self.assertEqual(response.status_code, 302)
        estimate.refresh_from_db()
        self.assertEqual(estimate.result_notes, "Проиграли: победитель снизился на 30%.")

    def test_user_cannot_change_another_users_estimate_status(self):
        estimate = TenderEstimate.objects.create(owner=self.other, tender_number="777", name="Чужой")
        self.client.force_login(self.user)

        response = self.client.post(
            reverse("tender_estimate_status", args=[estimate.pk]),
            {"status": "lost"},
        )

        self.assertEqual(response.status_code, 404)

    def test_user_cannot_see_another_users_estimate(self):
        estimate = TenderEstimate.objects.create(owner=self.other, tender_number="777", name="Чужой")
        self.client.force_login(self.user)
        self.assertEqual(self.client.get(reverse("tender_estimate", args=[estimate.pk])).status_code, 404)
        self.assertNotContains(self.client.get(reverse("tender_home")), "Чужой")

    def test_admin_sees_all_and_can_change_owner(self):
        admin = get_user_model().objects.create_superuser(username="admin", password="password")
        estimate = TenderEstimate.objects.create(owner=self.other, tender_number="777", name="Просчёт")
        self.client.force_login(admin)
        self.assertContains(self.client.get(reverse("tender_home")), "Просчёт")
        response = self.client.post(reverse("tender_estimate", args=[estimate.pk]), {"tender_number": "777", "name": "Просчёт", "owner_id": self.user.pk, "reduction_percent": "30", "russia_delivery": "0", "lines_json": json.dumps(self.payload)})
        self.assertEqual(response.status_code, 302)
        estimate.refresh_from_db()
        self.assertEqual(estimate.owner, self.user)

    def test_vat_rate_comes_from_admin_setting(self):
        TenderSettings.objects.create(pk=1, vat_rate="7")
        self.client.force_login(self.user)
        self.client.post(reverse("tender_home"), {"tender_number": "123", "name": "Тест", "reduction_percent": "30", "russia_delivery": "0", "lines_json": json.dumps(self.payload)})
        self.assertEqual(TenderEstimate.objects.get().vat_rate_snapshot, Decimal("7.00"))

    def test_empty_optional_costs_link_and_comment_are_allowed(self):
        line = {**self.payload[0], "application_unit": "", "logistics_unit": "", "product_url": "", "comment": ""}
        self.client.force_login(self.user)
        response = self.client.post(reverse("tender_home"), {"tender_number": "123", "name": "Только материал", "reduction_percent": "30", "russia_delivery": "", "lines_json": json.dumps([line])})
        self.assertEqual(response.status_code, 302)
        saved = TenderEstimate.objects.get().lines.get()
        self.assertEqual(saved.application_unit, Decimal("0.00"))
        self.assertEqual(saved.logistics_unit, Decimal("0.00"))

    def test_incomplete_lines_are_saved_as_draft_and_reopened(self):
        invalid = {**self.payload[0], "material_unit": "", "application_unit": "", "logistics_unit": ""}
        self.client.force_login(self.user)
        response = self.client.post(reverse("tender_home"), {"tender_number": "ABC-999", "name": "Не терять", "reduction_percent": "30", "russia_delivery": "", "lines_json": json.dumps([invalid])})
        self.assertEqual(response.status_code, 302)
        estimate = TenderEstimate.objects.get()
        self.assertTrue(estimate.summary_snapshot["is_incomplete"])
        self.assertEqual(estimate.lines.get().name, "Ручка")
        reopened = self.client.get(reverse("tender_estimate", args=[estimate.pk]))
        self.assertContains(reopened, "Ручка")
        self.assertContains(reopened, 'class="saved-estimate__status-form is-draft"')
        self.assertNotContains(reopened, 'class="saved-estimate__draft"')

    def test_partially_filled_line_values_are_preserved_in_draft(self):
        partial = {"name": "", "quantity": "50", "nmck_unit": "", "material_unit": "12.50", "application_unit": "", "logistics_unit": "", "product_url": "", "comment": "Уточнить товар", "requirements": {}}
        self.client.force_login(self.user)

        response = self.client.post(reverse("tender_home"), {"tender_number": "DRAFT-1", "name": "Черновик", "reduction_percent": "30", "russia_delivery": "", "lines_json": json.dumps([partial])})

        self.assertEqual(response.status_code, 302)
        line = TenderEstimate.objects.get().lines.get()
        self.assertEqual(line.name, "")
        self.assertEqual(line.quantity, Decimal("50"))
        self.assertEqual(line.material_unit, Decimal("12.50"))
        self.assertEqual(line.comment, "Уточнить товар")

    def test_excel_preview_returns_sheets_and_rows(self):
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "НМЦК"
        sheet.append(["Наименование", "Количество", "Цена"])
        sheet.append(["Ручка", 100, 25.5])
        content = BytesIO()
        workbook.save(content)
        content.seek(0)
        content.name = "nmck.xlsx"
        self.client.force_login(self.user)

        response = self.client.post(reverse("tender_import_preview"), {"file": content}, format="multipart")

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["sheet"], "НМЦК")
        self.assertEqual(payload["rows"][1], ["Ручка", "100", "25.5"])

    def test_excel_preview_requires_login(self):
        response = self.client.post(reverse("tender_import_preview"))
        self.assertEqual(response.status_code, 302)

    def test_docx_tables_are_extracted_without_saving_file(self):
        document = Document()
        table = document.add_table(rows=2, cols=3)
        table.rows[0].cells[0].text = "Наименование"
        table.rows[0].cells[1].text = "Количество"
        table.rows[0].cells[2].text = "Цена"
        table.rows[1].cells[0].text = "Блокнот"
        table.rows[1].cells[1].text = "20"
        table.rows[1].cells[2].text = "150"
        content = BytesIO()
        document.save(content)
        content.seek(0)
        content.name = "nmck.docx"

        text, truncated = extract_tender_source(content)

        self.assertIn("Блокнот | 20 | 150", text)
        self.assertFalse(truncated)

    def test_nested_docx_tables_include_technical_values(self):
        document = Document()
        outer = document.add_table(rows=1, cols=3)
        outer.cell(0, 0).text = "Папка «Благодарность»"
        nested = outer.cell(0, 1).add_table(rows=2, cols=2)
        nested.cell(0, 0).text = "Материал"
        nested.cell(0, 1).text = "Дизайнерский картон"
        nested.cell(1, 0).text = "Плотность"
        nested.cell(1, 1).text = "не менее 290 г/м²"
        outer.cell(0, 2).text = "1000"
        content = BytesIO()
        document.save(content)
        content.seek(0)
        content.name = "tz.docx"

        text, truncated = extract_tender_source(content)

        self.assertIn("Дизайнерский картон", text)
        self.assertIn("не менее 290 г/м²", text)
        self.assertFalse(truncated)

    @patch.dict("os.environ", {"TIMEWEB_AI_API_KEY": ""})
    def test_docx_embedded_excel_is_detected_and_parsed_locally(self):
        workbook = Workbook()
        sheet = workbook.active
        sheet.append(["№", "Наименование", "Ед.", "Кол-во", "Средняя цена", "НМЦК"])
        sheet.append([1, "Календарь настенный", "шт", 100, 440, 44000])
        sheet.append([2, "Календарь настольный", "шт", 300, 401.67, 120501])
        embedded = BytesIO()
        workbook.save(embedded)

        document = Document()
        document.add_heading("Обоснование НМЦК")
        base = BytesIO()
        document.save(base)
        content = BytesIO()
        with zipfile.ZipFile(BytesIO(base.getvalue())) as source, zipfile.ZipFile(content, "w", zipfile.ZIP_DEFLATED) as target:
            for info in source.infolist():
                target.writestr(info, source.read(info.filename))
            target.writestr("word/embeddings/Microsoft_Excel_Worksheet.xlsx", embedded.getvalue())
        content.seek(0)
        content.name = "Обоснование НМЦК.docx"

        inspection = inspect_tender_document(content)
        content.seek(0)
        result = recognize_tender_items(content)

        self.assertEqual(inspection["processing_mode"], "embedded")
        self.assertEqual(inspection["components"]["embedded_spreadsheets"], 1)
        self.assertTrue(result["local_parse"])
        self.assertEqual(len(result["items"]), 2)
        self.assertEqual(sum(Decimal(value["nmck_total"]) for value in result["items"]), Decimal("164501.00"))

    def test_model_json_accepts_explanatory_wrapper_and_control_characters(self):
        result = _json_from_model('Ответ модели:\n```json\n{"items":[{"name":"строка\tс табуляцией"}]}\n```')

        self.assertEqual(result["items"][0]["name"], "строка\tс табуляцией")

    @patch("tenders.views.detect_tender_document_type", return_value="unknown")
    @patch("tenders.views.analyze_tender_requirements")
    def test_second_smart_upload_can_be_forced_to_technical_document(self, analyze, detect):
        analyze.return_value = {"document_summary": "ТЗ", "global_requirements": [], "items": [], "warnings": [], "scan_ocr": False, "usage": {}}
        content = BytesIO(b"placeholder")
        content.name = "requirements.docx"
        self.client.force_login(self.user)

        response = self.client.post(reverse("tender_document_preview"), {
            "file": content,
            "lines_json": json.dumps(self.payload),
            "document_role": "technical",
        }, format="multipart")

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["document_type"], "technical")
        self.assertIn("technical", response.json())

    @patch("tenders.views.recognize_tender_items")
    @patch("tenders.views.detect_tender_document_type", return_value="technical")
    def test_explicit_nmck_action_rejects_obvious_technical_document(self, detect, recognize):
        content = BytesIO(b"placeholder")
        content.name = "requirements.docx"
        self.client.force_login(self.user)

        response = self.client.post(reverse("tender_document_preview"), {
            "file": content,
            "lines_json": "[]",
            "document_role": "nmck",
        }, format="multipart")

        self.assertEqual(response.status_code, 422)
        self.assertIn("Загрузить ООЗ / ТЗ", response.json()["error"])
        recognize.assert_not_called()

    @patch("tenders.views.analyze_tender_requirements")
    @patch("tenders.views.detect_tender_document_type", return_value="nmck")
    def test_explicit_technical_action_rejects_obvious_nmck_document(self, detect, analyze):
        content = BytesIO(b"placeholder")
        content.name = "nmck.xlsx"
        self.client.force_login(self.user)

        response = self.client.post(reverse("tender_document_preview"), {
            "file": content,
            "lines_json": "[]",
            "document_role": "technical",
        }, format="multipart")

        self.assertEqual(response.status_code, 422)
        self.assertIn("Загрузить НМЦК", response.json()["error"])
        analyze.assert_not_called()

    @patch("tenders.services._ai_gateway_json")
    def test_technical_analysis_excludes_delivery_terms_but_keeps_shared_specs(self, gateway):
        gateway.return_value = ({
            "document_summary": "Футболки",
            "global_requirements": [
                {"label": "Срок поставки", "value": "10 дней", "source": "стр. 1"},
                {"label": "Общий цвет", "value": "чёрный", "source": "стр. 1"},
            ],
            "items": [{
                "line_index": 0,
                "source_name": "Футболка",
                "quantity": 20,
                "requirements": [{"label": "Материал", "value": "хлопок"}],
                "missing": [],
                "questions": [],
                "confidence": .9,
            }],
            "warnings": [],
        }, {})
        document = Document()
        document.add_paragraph("Футболка, хлопок, чёрный цвет. Поставка за 10 дней.")
        content = BytesIO()
        document.save(content)
        content.seek(0)
        content.name = "ТЗ.docx"

        result = analyze_tender_requirements(content, [{"name": "Футболка", "quantity": "20"}])

        self.assertEqual(result["global_requirements"], [{"label": "Общий цвет", "value": "чёрный", "source": "стр. 1"}])
        self.assertIn("Не извлекай условия поставки", gateway.call_args.args[0])
        self.assertIn("Все числовые технические характеристики значимы", gateway.call_args.args[0])

    @patch("tenders.services._ai_gateway_json")
    def test_requirements_match_is_recovered_when_model_omits_confidence(self, gateway):
        gateway.return_value = ({"document_summary": "Папки", "global_requirements": [], "items": [{"line_index": None, "source_name": "Папка Благодарность 18.12.19.190", "quantity": 1000, "requirements": [{"label": "Материал", "value": "картон"}], "missing": [], "questions": []}], "warnings": []}, {})
        document = Document()
        document.add_paragraph("Папка Благодарность, 1000 штук, материал картон")
        content = BytesIO()
        document.save(content)
        content.seek(0)
        content.name = "tz.docx"

        result = analyze_tender_requirements(content, [{"name": "Папка «Благодарность»", "quantity": "1000"}])

        self.assertEqual(result["items"][0]["line_index"], 0)
        self.assertEqual(result["items"][0]["match_status"], "matched")
        self.assertGreater(result["items"][0]["confidence"], .8)

    @patch("tenders.services._ai_gateway_json")
    def test_matched_technical_quantity_comes_from_nmck_line(self, gateway):
        gateway.return_value = ({"document_summary": "Футболка", "global_requirements": [], "items": [{"line_index": 0, "source_name": "Футболка", "quantity": 1, "requirements": [{"label": "Материал", "value": "хлопок"}], "missing": [], "questions": [], "confidence": .9}], "warnings": []}, {})
        document = Document()
        document.add_paragraph("Футболка, количество 20, материал хлопок")
        content = BytesIO()
        document.save(content)
        content.seek(0)
        content.name = "ТЗ.docx"

        result = analyze_tender_requirements(content, [{"name": "Футболка", "quantity": "20"}])

        self.assertEqual(result["items"][0]["quantity"], "20")

    @patch("tenders.views.recognize_tender_items")
    def test_ai_preview_returns_editable_items(self, recognize):
        recognize.return_value = {"items": [{"name": "Блокнот", "quantity": "20", "nmck_unit": "150.00", "nmck_total": "3000.00", "total_from_source": True, "total_matches": True, "confidence": 0.9}], "warnings": [], "usage": {}}
        content = BytesIO(b"placeholder")
        content.name = "nmck.pdf"
        self.client.force_login(self.user)

        response = self.client.post(reverse("tender_ai_import_preview"), {"file": content}, format="multipart")

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["items"][0]["name"], "Блокнот")
        recognize.assert_called_once()

    def test_ai_preview_rejects_unsupported_file(self):
        content = BytesIO(b"text")
        content.name = "nmck.txt"
        self.client.force_login(self.user)

        response = self.client.post(reverse("tender_ai_import_preview"), {"file": content}, format="multipart")

        self.assertEqual(response.status_code, 400)

    @patch("tenders.views.analyze_tender_requirements")
    def test_legacy_doc_is_accepted_for_requirements(self, analyze):
        analyze.return_value = {"document_summary": "ТЗ", "global_requirements": [], "items": [], "warnings": [], "scan_ocr": False, "usage": {}}
        content = BytesIO(b"legacy-doc-placeholder")
        content.name = "tz.doc"
        self.client.force_login(self.user)

        response = self.client.post(reverse("tender_requirements_preview"), {"file": content, "lines_json": "[]"}, format="multipart")

        self.assertEqual(response.status_code, 200)
        analyze.assert_called_once()

    def test_scanned_pdf_is_detected_for_ocr_fallback(self):
        content = BytesIO(b"%PDF-1.4")
        content.name = "scan.pdf"
        with patch("tenders.services.PdfReader") as reader:
            reader.return_value.pages = [type("Page", (), {"extract_text": lambda self: ""})()]
            text, truncated = extract_tender_source(content)
        self.assertEqual(text, "")
        self.assertFalse(truncated)

    @patch("tenders.views.inspect_tender_document")
    def test_document_preflight_reports_visual_processing_mode(self, inspect):
        inspect.return_value = {"document_type": "unknown", "processing_mode": "visual", "truncated": False, "quality": {"usable": False}}
        content = BytesIO(b"%PDF-1.4")
        content.name = "scan.pdf"
        self.client.force_login(self.user)

        response = self.client.post(reverse("tender_document_inspect"), {"file": content}, format="multipart")

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["processing_mode"], "visual")

    def test_project_contract_is_not_classified_as_technical_document(self):
        document = Document()
        document.add_heading("Проект контракта")
        document.add_paragraph("Приложение содержит описание объекта закупки и требования к товару.")
        content = BytesIO()
        document.save(content)
        content.seek(0)
        content.name = "Проект контракта.docx"

        self.assertEqual(detect_tender_document_type(content), "unknown")

    def test_technical_filename_and_content_are_classified_together(self):
        document = Document()
        document.add_heading("Описание объекта закупки")
        document.add_paragraph("Технические характеристики футболки: хлопок, плотность 180 г/м².")
        content = BytesIO()
        document.save(content)
        content.seek(0)
        content.name = "Описание объекта закупки.docx"

        self.assertEqual(detect_tender_document_type(content), "technical")

    def test_short_tz_filename_is_a_valid_technical_signal(self):
        document = Document()
        document.add_paragraph("Футболка: материал хлопок, плотность 180 г/м², нанесение DTF.")
        content = BytesIO()
        document.save(content)
        content.seek(0)
        content.name = "ТЗ_мерч_2_позиции.docx"

        self.assertEqual(detect_tender_document_type(content), "technical")

    def test_broken_nonempty_pdf_text_layer_requires_visual_fallback(self):
        quality = _source_text_quality("/i1041 /i1086 /i1083 /i1086 /i0003 /i1090 /i1077 /i1082 /i1089 /i1090")

        self.assertFalse(quality["usable"])

    def test_normal_russian_pdf_text_layer_stays_on_fast_path(self):
        quality = _source_text_quality("Описание объекта закупки. Футболка хлопок, размер XL, тираж 100 штук.")

        self.assertTrue(quality["usable"])

    def test_document_numbers_accept_comma_dot_and_group_separators(self):
        self.assertEqual(_parse_document_decimal("3 870,67"), Decimal("3870.67"))
        self.assertEqual(_parse_document_decimal("3,870.67 руб."), Decimal("3870.67"))
        self.assertEqual(_parse_document_decimal("2.129,21"), Decimal("2129.21"))

    @patch("tenders.services._pdf_page_count", return_value=1)
    @patch("tenders.services.extract_tender_source", return_value=("/i1041 /i1086 /i1083 /i1086 /i0003", False))
    @patch("tenders.services._ai_gateway_json")
    def test_broken_pdf_text_layer_uses_visual_recognition(self, gateway, _extract, _page_count):
        gateway.return_value = ({"items": [{"name": "Футболка", "quantity": "10", "nmck_unit": "100", "nmck_total": None, "confidence": .9}], "warnings": []}, {})
        content = BytesIO(b"%PDF-1.4")
        content.name = "nmck.pdf"

        result = recognize_tender_items(content)

        self.assertTrue(result["scan_ocr"])
        self.assertEqual(result["processing_mode"], "visual")
        self.assertEqual(result["items"][0]["nmck_total"], "1000.00")
        self.assertTrue(gateway.call_args.kwargs["scan_ocr"])

    @patch("tenders.services._pdf_page_count", return_value=25)
    @patch("tenders.services._scan_pdf_images", return_value=["encoded-page"])
    @patch("tenders.services._ai_gateway_json")
    def test_long_scanned_pdf_is_processed_in_bounded_page_batches(self, gateway, scan_images, _page_count):
        from .services import _visual_gateway_responses

        gateway.return_value = ({"items": [], "warnings": []}, {"prompt_tokens": 1, "completion_tokens": 1})
        content = BytesIO(b"%PDF-1.4")
        content.name = "long-scan.pdf"

        responses = _visual_gateway_responses("Распознай документ", content, max_tokens=1000)

        self.assertEqual(len(responses), 3)
        self.assertEqual(scan_images.call_count, 3)
        self.assertEqual(scan_images.call_args_list[0].kwargs, {"start_page": 0, "page_limit": 12})
        self.assertEqual(scan_images.call_args_list[-1].kwargs, {"start_page": 24, "page_limit": 12})
        self.assertEqual(gateway.call_count, 3)
        self.assertEqual(gateway.call_args.kwargs["image_data_urls"], ["data:image/jpeg;base64,encoded-page"])

    @patch("tenders.views.analyze_tender_requirements")
    def test_technical_document_preview_uses_current_lines(self, analyze):
        analyze.return_value = {"document_summary": "Печать буклетов", "global_requirements": [], "items": [{"line_index": 0, "source_name": "Буклет", "requirements": [], "missing": [], "questions": [], "confidence": .9}], "warnings": [], "scan_ocr": False, "usage": {}}
        content = BytesIO(b"placeholder")
        content.name = "tz.pdf"
        self.client.force_login(self.user)

        response = self.client.post(reverse("tender_requirements_preview"), {"file": content, "lines_json": json.dumps(self.payload)}, format="multipart")

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["items"][0]["line_index"], 0)
        analyze.assert_called_once()
        self.assertEqual(analyze.call_args.args[1][0]["name"], "Ручка")

    def test_a5_imposition_uses_parent_sheet_and_rotation(self):
        self.assertEqual(calculate_sheet_imposition(148, 210, 210, 297, 0), {"ups": 2, "rotation": True})
        self.assertEqual(calculate_sheet_imposition(148, 210, 320, 450, 0)["ups"], 4)
        self.assertEqual(calculate_sheet_imposition(148, 210, 320, 450, 3)["ups"], 4)

    def test_unknown_exact_sra3_paper_stays_a_priced_question(self):
        a4 = PriceItem.objects.create(category="paper", name="Обычная A4 80", unit_price=Decimal("0.8"))
        a3 = PriceItem.objects.create(category="paper", name="Maestro Special A3 80", unit_price=Decimal("1.7"))
        sra3 = PriceItem.objects.create(category="paper", name="Немел SRA3 120г", unit_price=Decimal("3"))
        candidates = _paper_candidates({"finished_width_mm": 148, "finished_height_mm": 210, "units_per_product": 60, "material_query": "офсетная бумага", "grammage_gsm": 80, "bleed_mm": 0}, 300, [a4, a3, sra3])

        exact_sra3 = next(value for value in candidates if value["format"] == "SRA3" and value["grammage_gsm"] == 80)
        self.assertTrue(exact_sra3["price_missing"])
        self.assertEqual(exact_sra3["ups"], 4)
        self.assertEqual(exact_sra3["sheets"], 4635)

    @patch("tenders.views.build_training_hypothesis")
    def test_production_route_preview_uses_current_position(self, analyze):
        analyze.return_value = {"stage": "training_dialogue", "product_type": "digital_sheet", "confidence": .4, "route": {"name": "Под ключ", "steps": ["Изготовление"]}, "costs": [], "totals": {}}
        self.client.force_login(self.user)

        self.user.is_superuser = True
        self.user.is_staff = True
        self.user.save(update_fields=["is_superuser", "is_staff"])

        response = self.client.post(reverse("tender_production_route_preview"), {"line_json": json.dumps({"name": "Блокнот А5", "quantity": 300, "requirements": {}})})

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["product_type"], "digital_sheet")
        self.assertTrue(response.json()["session_id"])
        analyze.assert_called_once()
        self.assertEqual(analyze.call_args.args[0]["name"], "Блокнот А5")

    def test_manager_cannot_start_ai_calculation(self):
        self.client.force_login(self.user)

        response = self.client.post(reverse("tender_production_route_preview"), {"line_json": json.dumps({"name": "Блокнот А5", "quantity": 300, "requirements": {}})})

        self.assertEqual(response.status_code, 403)

    def test_page_exposes_ai_only_for_admin(self):
        self.client.force_login(self.user)
        manager_page = self.client.get(reverse("tender_home"))
        self.assertContains(manager_page, "aiEnabled=false")

        self.user.is_superuser = True
        self.user.is_staff = True
        self.user.save(update_fields=["is_superuser", "is_staff"])
        admin_page = self.client.get(reverse("tender_home"))
        self.assertContains(admin_page, "aiEnabled=true")

    @patch("tenders.views.build_training_hypothesis")
    def test_admin_feedback_creates_structured_turn_and_updates_session(self, build):
        self.user.is_superuser = True
        self.user.is_staff = True
        self.user.save(update_fields=["is_superuser", "is_staff"])
        session = ProductionTrainingSession.objects.create(
            created_by=self.user,
            position_name="Папка",
            requirements={"requirements": []},
            current_hypothesis={"route": {"name": "Старый маршрут"}},
        )
        build.return_value = {
            "stage": "training_dialogue",
            "product_type": "binding_special",
            "route": {"name": "Подрядчик под ключ", "steps": ["Изготовление под ключ"]},
            "costs": [{"category": "application", "name": "Изготовление", "amount_total": "10000.00"}],
            "totals": {"application_unit": "100.00", "cost_total": "10000.00"},
            "understood_changes": ["Выбран подрядчик под ключ"],
        }
        self.client.force_login(self.user)
        payload = {"session_id": session.pk, "line": {"name": "Папка", "quantity": 100, "requirements": {}}, "feedback": "Считать у подрядчика под ключ"}

        response = self.client.post(reverse("tender_revise_production_hypothesis"), {"payload": json.dumps(payload)})

        self.assertEqual(response.status_code, 200)
        session.refresh_from_db()
        self.assertEqual(session.current_hypothesis["route"]["name"], "Подрядчик под ключ")
        turn = ProductionTrainingTurn.objects.get(session=session)
        self.assertEqual(turn.feedback, "Считать у подрядчика под ключ")
        self.assertEqual(turn.understood_changes, ["Выбран подрядчик под ключ"])

    def test_confirmed_dialogue_becomes_training_example(self):
        self.user.is_superuser = True
        self.user.is_staff = True
        self.user.save(update_fields=["is_superuser", "is_staff"])
        session = ProductionTrainingSession.objects.create(
            created_by=self.user,
            position_name="Папка",
            requirements={"requirements": [{"label": "Нанесение", "value": "Тиснение"}]},
            current_hypothesis={
                "product_type": "binding_special",
                "facts": ["Тиснение"],
                "route": {"name": "Под ключ", "reason": "Специализированное изделие", "steps": ["Изготовление под ключ"]},
                "costs": [{"category": "application", "name": "Изготовление", "amount_total": "10000.00"}],
                "totals": {"cost_total": "10000.00"},
                "psodin_calculation": {"authorized": True, "productivity_per_hour": "10", "tariff": "partner"},
            },
        )
        self.client.force_login(self.user)

        response = self.client.post(reverse("tender_confirm_production_type"), {"payload": json.dumps({"session_id": session.pk})})

        self.assertEqual(response.status_code, 200)
        session.refresh_from_db()
        self.assertTrue(session.is_confirmed)
        self.assertEqual(session.confirmed_example.routes[0]["name"], "Под ключ")
        self.assertEqual(session.confirmed_example.routes[0]["psodin_calculation"]["productivity_per_hour"], "10")

    def test_new_confirmation_supersedes_duplicate_without_deleting_history(self):
        self.user.is_superuser = True
        self.user.save(update_fields=["is_superuser"])
        production_type = ProductionType.objects.get(code="binding_special")
        previous = ProductionTrainingExample.objects.create(
            production_type=production_type, position_name="Папка", requirements={}, features=[],
            routes=[{"name": "Старый маршрут"}], created_by=self.user,
        )
        session = ProductionTrainingSession.objects.create(
            created_by=self.user, position_name="Папка", requirements={},
            current_hypothesis={
                "product_type": production_type.code,
                "route": {"name": "Новый маршрут", "reason": "Исправлено", "steps": ["Новый маршрут"]},
                "costs": [], "totals": {}, "learning_warnings": [],
            },
        )
        self.client.force_login(self.user)

        response = self.client.post(reverse("tender_confirm_production_type"), {"payload": json.dumps({"session_id": session.pk})})

        self.assertEqual(response.status_code, 200)
        previous.refresh_from_db()
        self.assertFalse(previous.is_active)
        self.assertEqual(previous.superseded_by_id, response.json()["example_id"])

    def test_unverified_tz_price_is_blocked_from_learning(self):
        production_type = ProductionType.objects.get(code="binding_special")
        raw = {
            "product_type": production_type.code,
            "route": {"reason": "Под ключ", "processes": [{"name": "Универсальная типография"}]},
            "costs": [{
                "category": "logistics", "name": "Логистика", "amount_total": "3000",
                "source": "дано в ТЗ", "source_type": "manager",
                "recipe": {"method": "fixed", "inputs": {"fixed_amount": "3000"}},
            }],
        }

        result = _normalize_training_hypothesis(raw, {"quantity": 1000, "requirements": {"requirements": []}}, [production_type], [])

        self.assertEqual(result["costs"][0]["source"], "Источник цены не подтверждён")
        self.assertTrue(result["learning_warnings"])

    @patch("tenders.services._ai_gateway_json")
    def test_untrained_classification_cannot_claim_high_confidence(self, gateway):
        gateway.return_value = ({"suggested_type": "digital_sheet", "confidence": .95, "reason": "Малый тираж", "features": ["100 штук"], "alternatives": [], "matched_example_ids": [], "questions": []}, {})

        result = classify_production_type({"name": "Открытка", "quantity": 100, "requirements": {}})

        self.assertEqual(result["stage"], "production_classification")
        self.assertLessEqual(result["confidence"], .45)

    @patch("tenders.services._ai_gateway_json")
    def test_classification_drops_question_already_answered_by_tz(self, gateway):
        gateway.return_value = ({
            "suggested_type": "binding_specialized",
            "confidence": .45,
            "reason": "Папка с декоративной отделкой",
            "features": ["3D конгрев герба", "горячее тиснение надписи"],
            "alternatives": [],
            "matched_example_ids": [],
            "routes": [],
            "questions": [{
                "question": "Требуется ли офсетная печать или возможно цифровая печать для нанесения текста и герба?",
                "missing_fact": "способ нанесения",
                "why_it_changes_route": "выбор технологии",
            }],
        }, {})
        line = {
            "name": "Папка «Благодарность»",
            "quantity": 1000,
            "requirements": {"requirements": [
                {"label": "Нанесение герба", "value": "Полнообъемный 3D конгрев с золотой металлизацией"},
                {"label": "Надпись", "value": "Горячее тиснение фольгой цвет золото"},
            ]},
        }

        result = classify_production_type(line)

        self.assertEqual(result["questions"], [])

    def test_only_admin_can_confirm_training_example(self):
        production_type = ProductionType.objects.get(code="digital_sheet")
        payload = {"line": {"name": "Открытка", "requirements": {"requirements": []}}, "production_type": production_type.code, "features": ["тираж 100"], "routes": [{"name": "Под ключ", "processes": [{"role": "production", "name": "Цифровая листовая печать"}]}]}
        self.client.force_login(self.user)
        denied = self.client.post(reverse("tender_confirm_production_type"), {"payload": json.dumps(payload)})
        self.assertEqual(denied.status_code, 403)

        self.user.is_superuser = True
        self.user.is_staff = True
        self.user.save(update_fields=["is_superuser", "is_staff"])
        saved = self.client.post(reverse("tender_confirm_production_type"), {"payload": json.dumps(payload)})
        self.assertEqual(saved.status_code, 200)
        example = ProductionTrainingExample.objects.get(position_name="Открытка", production_type=production_type)
        self.assertEqual(example.routes[0]["processes"][0]["name"], "Цифровая листовая печать")

    @patch.dict("os.environ", {"TIMEWEB_AI_API_KEY": ""})
    def test_multiline_nmck_xlsx_is_parsed_locally_with_final_totals(self):
        workbook = Workbook()
        sheet = workbook.active
        sheet.append(["Обоснование начальной (максимальной) цены контракта"])
        sheet.append(["№", "Наименование услуги", "", "Количество закупаемых позиций", "", "", "", "Цена Исполнителя", "", "", "", "", "", "Средняя арифметическая цена", "", "", "Начальная (максимальная) цена"])
        sheet.append([None, None, None, None, None, None, None, "Исполнитель 1"])
        sheet.append([None] * 13 + ["Средняя цена за единицу"] + [None, None, "НМЦК позиции"])
        sheet.append([1, "Услуги по изготовлению и поставке подарочной продукции (футболка подарочная № 1)", None, 20, None, None, None, 4900, 3950, 4800, None, None, None, 3870.67, None, None, 77413.40])
        stream = BytesIO()
        workbook.save(stream)
        stream.seek(0)
        stream.name = "Обоснование НМЦК.xlsx"

        self.assertEqual(detect_tender_document_type(stream), "nmck")
        result = recognize_tender_items(stream)

        self.assertTrue(result["local_parse"])
        self.assertEqual(result["items"][0]["name"], "футболка подарочная № 1")
        self.assertEqual(result["items"][0]["nmck_unit"], "3870.67")
        self.assertEqual(result["items"][0]["nmck_total"], "77413.40")

    @patch.dict("os.environ", {"TIMEWEB_AI_API_KEY": ""})
    def test_nmck_xlsx_accepts_volume_and_abbreviated_quantity_headers(self):
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "МинЦена"
        sheet.cell(3, 4, "Обоснование начальной (максимальной) цены Контракта.")
        sheet.cell(8, 1, "Начальная (максимальная) цена контракта")
        sheet.cell(8, 13, "Минимальная цена выбранная Заказчиком за единицу товара *")
        sheet.cell(8, 14, "Сумма начальной (максимальной) цены контракта")
        sheet.cell(9, 2, "Наименование товара, работ, услуг")
        sheet.cell(9, 3, "Объем")
        sheet.cell(10, 3, "Ед.изм.")
        sheet.cell(10, 4, "Кол-во")
        sheet.cell(11, 2, "Карта клиента")
        sheet.cell(11, 4, 2000)
        sheet.cell(11, 8, 24.45)
        sheet.cell(11, 13, 18.90)
        sheet.cell(11, 14, 37800)
        stream = BytesIO()
        workbook.save(stream)
        stream.seek(0)
        stream.name = "Обоснование НМЦК.xlsx"

        self.assertEqual(detect_tender_document_type(stream), "nmck")
        result = recognize_tender_items(stream)

        self.assertTrue(result["local_parse"])
        self.assertEqual(result["items"][0]["quantity"], "2000")
        self.assertEqual(result["items"][0]["nmck_unit"], "18.90")
        self.assertEqual(result["items"][0]["nmck_total"], "37800.00")

    @patch.dict("os.environ", {"TIMEWEB_AI_API_KEY": ""})
    def test_nmck_xlsx_accepts_cost_wording_without_price_word(self):
        workbook = Workbook()
        sheet = workbook.active
        sheet.append(["Наименование", "Количество", "Среднеарифметическая стоимость за единицу", "Среднеарифметическая стоимость за все кол-во товара"])
        sheet.append(["Папка картонная", 889, 1458.5, 1296606.5])
        stream = BytesIO()
        workbook.save(stream)
        stream.seek(0)
        stream.name = "НМЦК.xlsx"

        result = recognize_tender_items(stream)

        self.assertTrue(result["local_parse"])
        self.assertEqual(result["items"][0]["quantity"], "889")
        self.assertEqual(result["items"][0]["nmck_unit"], "1458.50")
        self.assertEqual(result["items"][0]["nmck_total"], "1296606.50")

    @patch.dict("os.environ", {"TIMEWEB_AI_API_KEY": ""})
    def test_nmck_table_accepts_characteristics_as_name_and_yo_in_volume(self):
        workbook = Workbook()
        sheet = workbook.active
        sheet.append(["Основные характеристики объекта закупки", "Объём", "Среднее ценовое значение", "Начальная (максимальная) цена контракта"])
        sheet.append(["Чехол на чемодан с логотипом", 200, 1713, 342600])
        stream = BytesIO()
        workbook.save(stream)
        stream.seek(0)
        stream.name = "Обоснование НМЦК.xlsx"

        result = recognize_tender_items(stream)

        self.assertTrue(result["local_parse"])
        self.assertEqual(result["items"][0]["name"], "Чехол на чемодан с логотипом")
        self.assertEqual(result["items"][0]["quantity"], "200")
        self.assertEqual(result["items"][0]["nmck_total"], "342600.00")

    @patch("tenders.services.extract_tender_source", return_value=("", False))
    def test_structured_nmck_xlsx_is_detected_when_text_classification_fails(self, _extract):
        workbook = Workbook()
        sheet = workbook.active
        sheet.cell(8, 13, "Минимальная цена выбранная Заказчиком за единицу товара")
        sheet.cell(8, 14, "Сумма начальной (максимальной) цены контракта")
        sheet.cell(9, 2, "Наименование товара, работ, услуг")
        sheet.cell(9, 4, "Кол-во")
        sheet.cell(11, 2, "Карта клиента")
        sheet.cell(11, 4, 2000)
        sheet.cell(11, 13, 18.90)
        sheet.cell(11, 14, 37800)
        stream = BytesIO()
        workbook.save(stream)
        stream.seek(0)
        stream.name = "Обоснование НМЦК.xlsx"

        self.assertEqual(detect_tender_document_type(stream), "nmck")

    @patch.dict("os.environ", {"TIMEWEB_AI_API_KEY": ""})
    def test_smart_document_endpoint_imports_structured_nmck_xlsx(self):
        workbook = Workbook()
        sheet = workbook.active
        sheet.cell(3, 4, "Обоснование начальной (максимальной) цены Контракта")
        sheet.cell(8, 13, "Минимальная цена выбранная Заказчиком за единицу товара")
        sheet.cell(8, 14, "Сумма начальной (максимальной) цены контракта")
        sheet.cell(9, 2, "Наименование товара, работ, услуг")
        sheet.cell(9, 4, "Кол-во")
        sheet.cell(11, 2, "Карта клиента")
        sheet.cell(11, 4, 2000)
        sheet.cell(11, 13, 18.90)
        sheet.cell(11, 14, 37800)
        stream = BytesIO()
        workbook.save(stream)
        stream.seek(0)
        stream.name = "Обоснование НМЦК.xlsx"
        self.client.force_login(self.user)

        response = self.client.post(reverse("tender_document_preview"), {"file": stream, "lines_json": "[]"})

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["document_type"], "nmck")
        self.assertEqual(response.json()["nmck"]["items"][0]["nmck_total"], "37800.00")

    def test_repeated_document_prefix_is_removed_from_item_names(self):
        items = [
            {"name": "полиграфической продукции: Карта «Саранск-Мордовия»"},
            {"name": "полиграфической продукции: Лифлет «Мордовия заповедная»"},
            {"name": "полиграфической продукции: Блокнот № 1"},
        ]

        names = _strip_shared_item_boilerplate(items)

        self.assertEqual(names, ["Карта «Саранск-Мордовия»", "Лифлет «Мордовия заповедная»", "Блокнот № 1"])

    @patch.dict("os.environ", {"TIMEWEB_AI_API_KEY": "test-key"})
    @patch("tenders.services._ai_gateway_json")
    def test_smart_excel_uses_one_ai_call_to_normalize_all_names(self, gateway):
        gateway.return_value = ({"items": [
            {"index": 0, "name": "Карта «Саранск-Мордовия»"},
            {"index": 1, "name": "Лифлет «Мордовия заповедная»"},
        ]}, {"prompt_tokens": 80, "completion_tokens": 30})
        items = [
            {"name": "полиграфической продукции: Карта «Саранск-Мордовия»"},
            {"name": "полиграфической продукции: Лифлет «Мордовия заповедная»"},
        ]

        normalized, usage, warning = _shorten_structured_item_names(items)

        self.assertEqual([value["name"] for value in normalized], ["Карта «Саранск-Мордовия»", "Лифлет «Мордовия заповедная»"])
        self.assertEqual(usage["prompt_tokens"], 80)
        self.assertIsNone(warning)
        gateway.assert_called_once()

    def test_large_technical_table_is_split_only_between_product_rows(self):
        header = "ОПИСАНИЕ ОБЪЕКТА ЗАКУПКИ\n№ | Наименование | Характеристики"
        rows = [f"{index} | Товар {index} | " + ("характеристика " * 70) for index in range(1, 13)]

        chunks = _technical_source_chunks("\n".join([header, *rows]), max_chars=2400)

        self.assertGreater(len(chunks), 1)
        self.assertTrue(all("ОПИСАНИЕ ОБЪЕКТА ЗАКУПКИ" in value for value in chunks))
        combined = "\n".join(chunks)
        self.assertTrue(all(f"{index} | Товар {index} |" in combined for index in range(1, 13)))

    def test_large_technical_table_keeps_repeated_characteristics_of_product_together(self):
        header = "ОПИСАНИЕ ОБЪЕКТА ЗАКУПКИ\n№ | Наименование | Параметр | Значение"
        rows = [
            *[f"1 | Шнурок | Параметр {index} | " + ("значение " * 12) for index in range(5)],
            *[f"2 | 3D-стикер | Параметр {index} | " + ("значение " * 12) for index in range(5)],
        ]

        chunks = _technical_source_chunks("\n".join([header, *rows]), max_chars=700)

        self.assertEqual(len(chunks), 2)
        self.assertEqual(sum("1 | Шнурок |" in value for value in chunks), 1)
        self.assertEqual(sum("2 | 3D-стикер |" in value for value in chunks), 1)
        self.assertTrue(any(value.count("1 | Шнурок |") == 5 for value in chunks))
        self.assertTrue(any(value.count("2 | 3D-стикер |") == 5 for value in chunks))

    @patch("tenders.services._technical_source_chunks", return_value=["первая часть", "вторая часть"])
    @patch("tenders.services.extract_tender_source", return_value=("длинный документ", False))
    @patch("tenders.services._ai_gateway_json")
    def test_partial_technical_answers_for_one_product_are_merged(self, gateway, extract, chunks):
        usage = {"prompt_tokens": 20, "completion_tokens": 20}
        gateway.side_effect = [
            ({
                "items": [{
                    "line_index": 0, "source_name": "3D стикеры", "quantity": 5000,
                    "requirements": [{"label": "Вид продукции", "value": "3D-стикер", "source": "таблица 1"}],
                    "missing": ["Материал", "Размеры"],
                    "questions": ["Какой материал используется?"], "confidence": .8,
                }],
                "global_requirements": [], "warnings": [], "document_summary": "",
            }, usage),
            ({
                "items": [{
                    "line_index": 0, "source_name": "3D стикеры", "quantity": 5000,
                    "requirements": [
                        {"label": "Материал", "value": "Полимерная смола", "source": "таблица 1"},
                        {"label": "Ширина", "value": "50 мм", "source": "таблица 1"},
                        {"label": "Высота", "value": "50 мм", "source": "таблица 1"},
                    ],
                    "missing": [], "questions": [], "confidence": .9,
                }],
                "global_requirements": [], "warnings": [], "document_summary": "",
            }, usage),
        ]

        result = analyze_tender_requirements(
            type("Upload", (), {"name": "test.docx"})(),
            [{"name": "Изготовление 3D стикеров", "quantity": 5000}],
        )

        self.assertEqual(len(result["items"]), 1)
        self.assertEqual(result["items"][0]["line_index"], 0)
        self.assertEqual(len(result["items"][0]["requirements"]), 4)
        self.assertEqual(result["items"][0]["missing"], [])
        self.assertEqual(result["items"][0]["questions"], [])

    @patch("tenders.services.extract_tender_source")
    @patch("tenders.services._ai_gateway_json")
    def test_empty_context_echoes_do_not_occupy_technical_matches(self, gateway, extract):
        extract.return_value = ("Описание объекта закупки", False)
        gateway.return_value = ({
            "items": [
                {"line_index": 0, "source_name": "Первый товар", "quantity": 10, "requirements": [], "missing": [], "questions": [], "confidence": .9},
                {"line_index": 1, "source_name": "Второй товар", "quantity": 20, "requirements": [{"label": "Материал", "value": "Бумага", "source": "таблица 1"}], "missing": [], "questions": [], "confidence": .9},
            ],
            "global_requirements": [], "warnings": [], "document_summary": "",
        }, {"prompt_tokens": 20, "completion_tokens": 20})
        lines = [{"name": "Первый товар", "quantity": 10}, {"name": "Второй товар", "quantity": 20}]

        result = analyze_tender_requirements(type("Upload", (), {"name": "test.docx"})(), lines)

        self.assertEqual(len(result["items"]), 1)
        self.assertEqual(result["items"][0]["line_index"], 1)
        self.assertIn("Первый товар", result["warnings"][0])
        self.assertIn("1 строку", result["warnings"][0])

    @patch("tenders.services.extract_tender_source")
    @patch("tenders.services._ai_gateway_json")
    def test_requirements_warn_about_every_nmck_line_missing_from_technical_document(self, gateway, extract):
        extract.return_value = ("Описание объекта закупки", False)
        gateway.return_value = ({
            "items": [{
                "line_index": 1,
                "source_name": "Карта с картонной обложкой",
                "quantity": 5000,
                "requirements": [{"label": "Материал", "value": "Картон", "source": "таблица 1"}],
                "missing": [],
                "questions": [],
                "confidence": .95,
            }],
            "global_requirements": [], "warnings": [], "document_summary": "Карты",
        }, {"prompt_tokens": 20, "completion_tokens": 20})
        lines = [
            {"name": "Карта «Саранск-Мордовия»", "quantity": 2000},
            {"name": "Карта «Саранск-Мордовия» с картонными обложками", "quantity": 5000},
        ]

        result = analyze_tender_requirements(type("Upload", (), {"name": "test.docx"})(), lines)

        self.assertEqual(result["items"][0]["line_index"], 1)
        self.assertEqual(len(result["warnings"]), 1)
        self.assertIn("ООЗ/ТЗ не покрывает 1 строку НМЦК", result["warnings"][0])
        self.assertIn("Карта «Саранск-Мордовия»", result["warnings"][0])

    def test_local_batch_index_does_not_override_a_better_name_match(self):
        lines = [
            {"name": "Лифлет Мордовия", "quantity": 3000},
            {"name": "Стикерпак", "quantity": 500},
        ]

        index, confidence, _ = _resolve_line_match(0, "Стикерпак", 500, lines, set())

        self.assertEqual(index, 1)
        self.assertGreaterEqual(confidence, .66)

    def test_training_cost_keeps_detailed_calculation_trace(self):
        production_type = ProductionType.objects.create(code="trace-test", name="Тест")
        raw = {
            "product_type": production_type.code,
            "route": {"name": "Под ключ", "steps": ["Изготовление"]},
            "costs": [{
                "category": "material", "name": "Majestic SRA3", "amount_total": 9500,
                "source": "Калькулятор PSODIN", "source_type": "calculator", "source_date": "25.08.2026",
                "basis": "25 листов × 380 ₽", "calculation_steps": ["4 изделия с листа", "100 / 4 = 25 листов"],
                "adaptation": "Рассчитано для тиража 100 шт.", "confirmed": True,
            }],
        }

        result = _normalize_training_hypothesis(raw, {"quantity": 100}, [production_type], [1])
        cost = result["costs"][0]

        self.assertEqual(cost["calculation_steps"], ["4 изделия с листа", "100 / 4 = 25 листов"])
        self.assertEqual(cost["adaptation"], "Рассчитано для тиража 100 шт.")
        self.assertEqual(cost["source_type"], "calculator")

    def test_recipe_recalculates_current_quantity_instead_of_copying_history_total(self):
        total, steps = _evaluate_cost_recipe({"method": "sheet_yield", "inputs": {"unit_price": 380, "units_per_sheet": 4, "waste_percent": 5}}, Decimal("1000"))

        self.assertEqual(total, Decimal("99940.00"))
        self.assertIn("263 листов", " ".join(steps))

    def test_recipe_applies_discount_to_cost_not_to_quantity(self):
        total, steps = _evaluate_cost_recipe({
            "method": "unit_rate",
            "inputs": {"unit_rate": 1050},
            "modifiers": [{"type": "discount_percent", "value": 15}],
        }, Decimal("5"))

        self.assertEqual(total, Decimal("4462.50"))
        self.assertIn("5250.00 ₽", steps[-1])
        self.assertIn("Скидка 15%", steps[-1])

    def test_recipe_applies_ordered_discount_and_fixed_cost_on_backend(self):
        total, steps = _evaluate_cost_recipe({
            "method": "fixed",
            "inputs": {"fixed_amount": 10000},
            "modifiers": [
                {"type": "discount_percent", "value": 10},
                {"type": "add_fixed", "value": 500},
            ],
        }, Decimal("1"))

        self.assertEqual(total, Decimal("9500.00"))
        self.assertEqual(len(steps), 3)

    def test_recipe_rejects_invalid_discount_instead_of_showing_unreliable_total(self):
        total, steps = _evaluate_cost_recipe({
            "method": "fixed",
            "inputs": {"fixed_amount": 10000},
            "modifiers": [{"type": "discount_percent", "value": 120}],
        }, Decimal("1"))

        self.assertIsNone(total)
        self.assertEqual(steps, [])

    def test_gifts_parser_filters_category_and_maps_image_url(self):
        product_xml = StringIO("""<doct><product product_id=\"v1\"><code>V-1</code><name>Жилет утеплённый</name><product_size>М-L</product_size><matherial>Полиэстер</matherial><color>фиолетовый</color><brand>Brand</brand><content>Описание</content><price><price>1200</price></price><small_image src=\"reviewer/webp/test.webp\"/><ondemand>false</ondemand></product><product product_id=\"m1\"><code>M-1</code><name>Магнит</name></product></doct>""")
        tree_xml = StringIO("""<doct><page page_id=\"10\" name=\"Одежда / Жилеты\"><product product=\"v1\" page=\"10\"/></page><page page_id=\"20\" name=\"Сувениры\"><product product=\"m1\" page=\"20\"/></page></doct>""")
        stock_xml = StringIO("""<doct><stock product_id=\"v1\"><free>7</free><dealerprice>999</dealerprice></stock></doct>""")

        result = parse_gifts_catalog(product_xml, tree_xml, stock_xml, category="жилеты")

        self.assertEqual(len(result), 1)
        self.assertEqual(result[0]["external_id"], "v1")
        self.assertEqual(result[0]["article"], "V-1")
        self.assertEqual(result[0]["total_stock"], 7)
        self.assertEqual(result[0]["discount_price"], Decimal("999.00"))
        self.assertEqual(result[0]["colors"], ["фиолетовый"])
        self.assertEqual(result[0]["image_url"], "https://files.gifts.ru/reviewer/webp/test.webp")

    def test_gifts_parser_keeps_category_map_without_category_filter(self):
        product_xml = StringIO("""<doct><product product_id="v1"><name>Лонгслив</name><code>LS-1</code></product></doct>""")
        tree_xml = StringIO("""<doct><page page_id="10" name="Одежда / Футболки с длинным рукавом"><product product="v1" page="10"/></page></doct>""")

        rows, categories = parse_gifts_catalog(product_xml, tree_xml, include_categories=True)

        self.assertEqual(rows[0]["category_ids"], ["10"])
        self.assertEqual(rows[0]["category_names"], ["Одежда / Футболки с длинным рукавом"])
        self.assertEqual(categories, [{
            "external_id": "10", "parent_external_id": "",
            "name": "Одежда / Футболки с длинным рукавом",
            "path": "Одежда / Футболки с длинным рукавом",
        }])

    def test_gifts_parser_reads_flat_product_page_links(self):
        product_xml = StringIO("""<doct><product><product_id>111501</product_id><code>PU422001</code><name>Рубашка поло</name></product></doct>""")
        tree_xml = StringIO("""
            <doct><page>
                <page><page_id>1105688</page_id><name>Футболки поло</name></page>
                <product><page>1105688</page><product>111501</product></product>
            </page></doct>
        """)

        rows = parse_gifts_catalog(product_xml, tree_xml)

        self.assertEqual(rows[0]["category_ids"], ["1105688"])

    def test_gifts_parser_uses_primary_catalog_image(self):
        product_xml = StringIO("""<doct><product product_id="93294"><code>26728.60</code><name>Жилет Kama, белый</name><super_big_image src="reviewer/webp/26/6728.60_1_500.webp?v=2"/></product></doct>""")
        tree_xml = StringIO("<doct/>")

        result = parse_gifts_catalog(product_xml, tree_xml)

        self.assertEqual(result[0]["image_url"], "https://files.gifts.ru/reviewer/webp/26/6728.60_1_500.webp?v=2")
        self.assertEqual(result[0]["product_url"], "https://gifts.ru/id/93294")

    def test_gifts_parser_maps_xml_thumbnail_to_public_reviewer_image(self):
        product_xml = StringIO("""<doct><product product_id="16224"><code>1376.89</code><name>Футболка унисекс Regent 150, лайм</name><small_image src="thumbnails/7/1376.89_648_200x200.jpg"/></product></doct>""")

        result = parse_gifts_catalog(product_xml, StringIO("<doct/>"))

        self.assertEqual(result[0]["image_url"], "https://files.gifts.ru/reviewer/thumbnails/7/1376.89_648_200x200.jpg")

    def test_gifts_parser_normalizes_protocol_relative_image_url(self):
        product_xml = StringIO("""<doct><product product_id=\"184880\"><code>03564102</code><name>Футболка унисекс Epic, белая</name><small_image src=\"//files.gifts.ru/reviewer/webp/8/03564102_2_200x200.webp?v=2\"/></product></doct>""")
        result = parse_gifts_catalog(product_xml, StringIO("<doct/>"))

        self.assertEqual(result[0]["image_url"], "https://files.gifts.ru/reviewer/webp/8/03564102_2_200x200.webp?v=2")

    def test_gifts_parser_keeps_name_color_out_of_catalog_color_field(self):
        product_xml = StringIO("""<doct><product product_id=\"lime\"><code>03564102</code><name>Футболка унисекс Regent 150, лайм</name></product><product product_id=\"khaki\"><code>03564103</code><name>Футболка унисекс Regent 150, хаки</name></product></doct>""")
        result = parse_gifts_catalog(product_xml, StringIO("<doct/>"))

        self.assertEqual(result[0]["colors"], [])
        self.assertEqual(result[0]["raw_data"]["name_colors"], ["лайм"])
        self.assertEqual(result[1]["colors"], [])
        self.assertEqual(result[1]["raw_data"]["name_colors"], ["хаки"])

    def test_gifts_parser_reads_color_from_filters_catalog(self):
        product_xml = StringIO("""<doct><product product_id=\"v1\"><code>V-1</code><name>Жилет Kama, фиолетовый</name><filters><filter><filtertypeid>21</filtertypeid><filterid>77</filterid></filter></filters></product></doct>""")
        filters_xml = StringIO("""<root><filtertypes><filtertype><filtertypeid>21</filtertypeid><filtertypename>Цвет</filtertypename><filters><filter><filterid>77</filterid><filtername>фиолетовый</filtername></filter></filters></filtertype><filtertype><filtertypeid>99</filtertypeid><filtertypename>Цвет упаковки</filtertypename><filters><filter><filterid>77</filterid><filtername>зеленый</filtername></filter></filters></filtertype></filtertypes></root>""")

        result = parse_gifts_catalog(product_xml, StringIO("<doct/>"), filters_xml=filters_xml)

        self.assertEqual(result[0]["colors"], ["фиолетовый"])

    def test_gifts_parser_finds_image_url_in_unknown_nested_xml_node(self):
        product_xml = StringIO("""<doct><product product_id=\"x\"><code>X</code><name>Товар</name><media><preview data-url=\"//files.gifts.ru/reviewer/webp/x.webp\"/></media></product></doct>""")
        result = parse_gifts_catalog(product_xml, StringIO("<doct/>"))

        self.assertEqual(result[0]["image_url"], "https://files.gifts.ru/reviewer/webp/x.webp")

    @patch.dict("os.environ", {"GIFTS_XML_USERNAME": "user", "GIFTS_XML_PASSWORD": "pass"})
    def test_gifts_client_requires_server_side_credentials(self):
        client = GiftsXmlClient()
        self.assertEqual(client.base_url, "https://api2.gifts.ru/export/v2")


    @patch.dict("os.environ", {"OASIS_API_KEY": ""})
    def test_oasis_client_requires_server_side_api_key(self):
        with self.assertRaisesMessage(CatalogSyncError, "OASIS_API_KEY не настроен"):
            OasisClient()

    @patch("tenders.catalog.urlopen")
    def test_oasis_client_retries_same_request_after_connection_reset(self, urlopen_mock):
        class Response:
            def __enter__(self):
                return self

            def __exit__(self, *args):
                return False

            def read(self):
                return b'{"items": []}'

        urlopen_mock.side_effect = [ConnectionResetError("temporary"), Response()]

        result = OasisClient(api_key="test", min_interval=0, max_attempts=2).get("/v4/products")

        self.assertEqual(result, {"items": []})
        self.assertEqual(urlopen_mock.call_count, 2)

    def test_oasis_sync_stores_compact_searchable_catalog_and_dealer_stock(self):
        class FakeClient:
            base_url = "https://api.oasiscatalog.com"

            def get(self, path, params=None):
                self.assert_path = path
                return [{"id": 3071, "parent_id": 10, "name": "Футболки", "path": "odezhda/futbolki"}]

            def pages(self, path, params=None, limit=100):
                if path == "/v4/products":
                    yield [{
                        "id": "1-000032048", "article": "3103101S", "article_base": "3103101",
                        "group_id": "100032034", "name": "Футболка Club мужская",
                        "full_name": "Футболка Club мужская, белая, S", "categories": [3071],
                        "materials": ["хлопок"], "colors": [{"name": "белый"}],
                        "attributes": [{"name": "Плотность", "value": "150 г/м²"}],
                        "branding": "Вышивка,DTF", "price": "209.00", "dealerPrice": "190.00",
                        "images": [{"small": "https://s.a-5.ru/test-small.jpg"}], "total_stock": 20,
                        "is_deleted": "0", "is_stopped": "0",
                    }]
                else:
                    yield [{
                        "id": "1-000032048", "article": "3103101S", "stock": 50,
                        "stock-remote": 10, "stock-transit": 100, "price": "209.00",
                        "price-discount": "180.00",
                    }]

        run = sync_oasis_catalog(FakeClient())
        product = CatalogProduct.objects.get()

        self.assertEqual(run.status, "success")
        self.assertEqual(run.created_count, 1)
        self.assertEqual(product.discount_price, Decimal("180.00"))
        self.assertEqual(product.total_stock, 60)
        self.assertEqual(product.category_names, ["odezhda/futbolki"])
        self.assertIn("плотность 150 г/м²", product.search_text)
        self.assertEqual(product.image_url, "https://s.a-5.ru/test-small.jpg")
        self.assertEqual(product.product_url, "https://www.oasiscatalog.com/item/1-000032048")
        self.assertTrue(product.is_active)
        self.assertEqual(product.raw_data, {"discount_group_id": None, "included_branding": None})

    def test_failed_oasis_sync_does_not_deactivate_previous_catalog(self):
        supplier = CatalogSupplier.objects.create(code="oasis", name="Oasis", base_url="https://api.oasiscatalog.com")
        product = CatalogProduct.objects.create(supplier=supplier, external_id="old", article="OLD", name="Старый товар", is_active=True, sync_marker="previous")

        class FailingClient:
            base_url = "https://api.oasiscatalog.com"

            def get(self, path, params=None):
                return []

            def pages(self, path, params=None, limit=100):
                if path == "/v4/products":
                    return
                    yield
                raise CatalogSyncError("Остатки временно недоступны")

        with self.assertRaises(CatalogSyncError):
            sync_oasis_catalog(FailingClient())

        product.refresh_from_db()
        supplier.refresh_from_db()
        self.assertTrue(product.is_active)
        self.assertEqual(supplier.sync_status, "failed")
        self.assertEqual(CatalogSyncRun.objects.get().status, "failed")

    def test_catalog_search_enforces_material_density_branding_and_stock(self):
        class Client:
            base_url = "https://api.oasiscatalog.com"
            def get(self, path, params=None):
                if path == "/v4/categories":
                    return [{"id": 10, "name": "Футболки поло", "path": "categories/tekstil/polo"}]
                return [
                    {"id": "exact", "article": "POLO-190", "group_id": "polo-exact", "name": "Футболка поло", "full_name": "Футболка поло тёмно-синяя", "materials": ["хлопок"], "colors": ["темно-синий"], "branding": ["Вышивка", "DTF"], "attributes": [{"name": "Плотность материала", "value": "190 г/м²"}], "price": 700, "discount_price": 650, "total_stock": 500, "categories": [10]},
                    {"id": "thin", "article": "POLO-160", "group_id": "polo-thin", "name": "Футболка поло", "full_name": "Футболка поло тёмно-синяя эконом", "materials": ["хлопок"], "colors": ["темно-синий"], "branding": ["Вышивка"], "attributes": [{"name": "Плотность материала", "value": "160 г/м²"}], "price": 400, "total_stock": 1000, "categories": [10]},
                ]
        line = {
            "name": "Футболка поло", "quantity": "300",
            "requirements": {"requirements": [
                {"label": "Материал", "value": "хлопок 100%"},
                {"label": "Цвет", "value": "темно-синий"},
                {"label": "Плотность", "value": "не менее 190 г/м²"},
                {"label": "Нанесение", "value": "вышивка"},
            ]},
        }

        candidates = catalog_candidates_for_line(line, limit=3, intent={"product_class": "поло"}, client=Client())

        self.assertEqual(candidates[0]["external_id"], "exact")
        self.assertEqual(candidates[0]["fit"], "exact")
        self.assertEqual(candidates[0]["price"], "650.00")
        self.assertEqual(candidates[0]["cost_total"], "195000.00")
        thin = next(value for value in candidates if value["external_id"] == "thin")
        self.assertEqual(thin["fit"], "partial")
        self.assertTrue(any("требуется не менее 190" in value for value in thin["mismatches"]))

    def test_category_candidates_rank_specific_product_above_broad_exact_category(self):
        candidates = _category_candidates({"oasis": [
            {"id": "vip", "name": "Одежда", "path": "vip/odezhda"},
            {"id": "shirts", "name": "Футболки", "path": "odezhda/futbolki"},
            {"id": "polo", "name": "Рубашки поло", "path": "odezhda/rubashki-polo"},
        ]}, {
            "name": "Футболка поло унисекс, цвет – белый",
        }, {
            "item": "рубашка поло унисекс",
            "product_class": "футболка",
            "categories": ["одежда", "футболки"],
            "synonyms": ["футболка поло", "поло унисекс"],
        })

        self.assertEqual(candidates[0]["category_id"], "polo")
        self.assertGreater(candidates[0]["specificity"], next(
            value["specificity"] for value in candidates if value["category_id"] == "vip"
        ))

    def test_category_retrieval_uses_item_before_product_class_and_preserves_sources(self):
        categories = [
            {
                "source": "large", "category_id": f"mug-{index}", "name": f"Кружки {index}",
                "parent_id": "", "path": f"Посуда > Кружки {index}",
            }
            for index in range(80)
        ] + [
            {
                "source": "large", "category_id": "polo-large", "name": "Рубашки поло",
                "parent_id": "", "path": "Одежда > Рубашки поло",
            },
            {
                "source": "large", "category_id": "shirts", "name": "Футболки",
                "parent_id": "", "path": "Одежда > Футболки",
            },
            {
                "source": "small", "category_id": "polo-small", "name": "Polo shirts",
                "parent_id": "", "path": "Textile > Polo shirts",
            },
        ]

        seeds, diagnostics = _category_retrieval(
            categories,
            {"name": "Поло унисекс"},
            {"item": "поло", "synonyms": ["рубашка поло", "polo shirt"], "product_class": "футболка"},
        )

        self.assertEqual({value["source"] for value in seeds}, {"large", "small"})
        self.assertIn("polo-large", {value["category_id"] for value in seeds})
        self.assertIn("polo-small", {value["category_id"] for value in seeds})
        large_ids = [value["category_id"] for value in seeds if value["source"] == "large"]
        self.assertLess(large_ids.index("polo-large"), large_ids.index("shirts") if "shirts" in large_ids else len(large_ids))
        self.assertEqual(diagnostics["considered_count"], len(categories))

    def test_category_graph_expansion_is_local_and_deduplicated(self):
        categories = [
            {"source": "supplier", "category_id": "root", "name": "Одежда", "parent_id": "", "path": "Одежда"},
            {"source": "supplier", "category_id": "polo", "name": "Поло", "parent_id": "root", "path": "Одежда > Поло"},
            {"source": "supplier", "category_id": "male", "name": "Мужские поло", "parent_id": "polo", "path": "Одежда > Поло > Мужские"},
            {"source": "supplier", "category_id": "female", "name": "Женские поло", "parent_id": "polo", "path": "Одежда > Поло > Женские"},
            {"source": "supplier", "category_id": "mugs", "name": "Кружки", "parent_id": "", "path": "Посуда > Кружки"},
        ]

        fragment = _expand_category_graph(categories, [{**categories[1], "retrieval_score": 1}])

        self.assertEqual({value["category_id"] for value in fragment}, {"root", "polo", "male", "female"})
        self.assertEqual(len(fragment), len({(value["source"], value["category_id"]) for value in fragment}))
        polo = next(value for value in fragment if value["category_id"] == "polo")
        self.assertEqual(set(polo["child_ids"]), {"male", "female"})

    @patch("tenders.services._ai_gateway_json")
    def test_llm_category_selection_accepts_only_real_category_ids(self, gateway):
        gateway.side_effect = [
            ({"categories": [
                {"source": "oasis", "category_id": "invented", "priority": 1},
                {"source": "oasis", "category_id": "polo", "priority": 1},
            ]}, {}),
            ({"categories": [
                {"source": "oasis", "category_id": "polo", "keep": True},
            ]}, {}),
        ]

        tasks, usage, errors = _select_catalog_category_tasks(
            {"name": "Футболка поло"},
            {"item": "рубашка поло"},
            [
                {"source": "oasis", "category_id": "polo", "name": "Рубашки поло", "path": "Одежда / Поло", "specificity": 1},
                {"source": "oasis", "category_id": "shirts", "name": "Футболки", "path": "Одежда / Футболки", "specificity": .5},
            ],
        )

        self.assertEqual(tasks, [{
            "source": "oasis", "category_id": "polo", "name": "Рубашки поло",
            "path": "Одежда / Поло", "priority": 1,
        }])
        self.assertFalse(errors)
        self.assertEqual(usage["llm_calls"], 2)
        self.assertEqual(usage["selector_llm_calls"], 1)
        self.assertEqual(usage["verifier_llm_calls"], 1)
        self.assertFalse(usage["retry_used"])

    @patch("tenders.services._ai_gateway_json")
    def test_llm_category_selection_receives_one_compact_multi_source_fragment(self, gateway):
        captured = {}

        def answer(prompt, **kwargs):
            if "SEMANTIC VERIFIER" in prompt:
                captured["verifier_prompt"] = prompt
                return ({"categories": [
                    {"source": "supplier-x", "category_id": "protective-headwear", "keep": True},
                    {"source": "supplier-y", "category_id": "helmets", "keep": True},
                ]}, {"input_tokens": 45, "output_tokens": 10})
            captured["selector_prompt"] = prompt
            return ({"categories": [
                {"source": "supplier-x", "category_id": "protective-headwear", "priority": 1},
                {"source": "supplier-y", "category_id": "helmets", "priority": 1},
            ]}, {"input_tokens": 321, "output_tokens": 123})

        gateway.side_effect = answer
        tasks, usage, errors = _select_catalog_category_tasks(
            {"name": "Каска защитная зимняя"},
            {
                "item": "защитная каска",
                "required": [{"label": "Сезон", "value": "зима", "weight": 1}],
            },
            [
                {
                    "source": "supplier-x", "category_id": "winter-helmets", "name": "Зимние каски",
                    "parent_id": "protective-headwear", "path": "Каталог > Спецодежда > Защита головы > Зимние каски",
                    "specificity": 10,
                },
                {
                    "source": "supplier-x", "category_id": "workwear", "name": "Спецодежда",
                    "parent_id": "", "path": "Каталог > Спецодежда", "specificity": 1,
                },
                {
                    "source": "supplier-x", "category_id": "protective-headwear", "name": "Защита головы",
                    "parent_id": "workwear", "path": "Каталог > Спецодежда > Защита головы", "specificity": 1,
                },
                {
                    "source": "supplier-x", "category_id": "helmets-general", "name": "Защитные каски",
                    "parent_id": "protective-headwear", "path": "Каталог > Спецодежда > Защита головы > Защитные каски",
                    "specificity": 1,
                },
                {
                    "source": "supplier-y", "category_id": "helmets", "name": "Каски",
                    "parent_id": "", "path": "Защита > Каски", "specificity": 8,
                },
            ],
        )

        prompt = captured["selector_prompt"]
        self.assertIn('"parent_id"', prompt)
        self.assertIn('"child_ids"', prompt)
        self.assertIn('"supplier-x"', prompt)
        self.assertIn('"supplier-y"', prompt)
        self.assertIn("минимальный достаточный набор", prompt)
        self.assertIn("лучше вернуть меньше категорий", prompt)
        self.assertNotIn('primary', prompt)
        self.assertNotIn('equivalent', prompt)
        self.assertNotIn('condition', prompt)
        verifier_prompt = captured["verifier_prompt"]
        self.assertIn('"item":"защитная каска"', verifier_prompt)
        self.assertNotIn('"parent_id"', verifier_prompt)
        self.assertNotIn('"child_ids"', verifier_prompt)
        self.assertNotIn('"priority"', verifier_prompt)
        self.assertNotIn("Сезон", verifier_prompt)
        self.assertEqual(gateway.call_count, 2)
        self.assertEqual({value["source"] for value in tasks}, {"supplier-x", "supplier-y"})
        self.assertEqual(usage["input_tokens"], 366)
        self.assertEqual(usage["output_tokens"], 133)
        self.assertEqual(usage["selector_input_tokens"], 321)
        self.assertEqual(usage["verifier_input_tokens"], 45)
        self.assertEqual(usage["llm_calls"], 2)
        self.assertFalse(errors)

    @patch("tenders.services._ai_gateway_json")
    def test_category_verifier_removes_only_explicit_rejections_and_defaults_to_keep(self, gateway):
        gateway.return_value = ({"categories": [
            {"source": "gifts", "category_id": "mugs", "keep": True},
            {"source": "gifts", "category_id": "sets", "keep": False},
            {"source": "gifts", "category_id": "invented", "keep": False},
        ]}, {"prompt_tokens": 50, "completion_tokens": 12})
        tasks = [
            {"source": "gifts", "category_id": "mugs", "name": "Кружки", "path": "Посуда > Кружки", "priority": 1},
            {"source": "gifts", "category_id": "sets", "name": "Подарочные наборы", "path": "Наборы > С кружками", "priority": 1},
            {"source": "gifts", "category_id": "tableware", "name": "Посуда", "path": "Посуда", "priority": 1},
        ]

        kept, usage, decisions = _verify_catalog_category_tasks("кружка", tasks)

        self.assertEqual([value["category_id"] for value in kept], ["mugs", "tableware"])
        self.assertEqual(usage["prompt_tokens"], 50)
        self.assertEqual(decisions[("gifts", "sets")], False)
        prompt = gateway.call_args.args[0]
        self.assertIn("SEMANTIC VERIFIER", prompt)
        self.assertNotIn('"priority"', prompt)

    @patch("tenders.services._ai_gateway_json")
    def test_category_verifier_does_not_trigger_semantic_retry(self, gateway):
        gateway.side_effect = [
            ({"categories": [{"source": "supplier", "category_id": "accessories", "priority": 1}]}, {}),
            ({"categories": [{"source": "supplier", "category_id": "accessories", "keep": False}]}, {}),
        ]

        tasks, usage, errors = _select_catalog_category_tasks(
            {"name": "Кружка"}, {"item": "кружка"}, [{
                "source": "supplier", "category_id": "accessories", "name": "Для кружек",
                "parent_id": "", "path": "Упаковка > Для кружек", "specificity": 1,
            }],
        )

        self.assertEqual(tasks, [])
        self.assertEqual(gateway.call_count, 2)
        self.assertFalse(usage["retry_used"])
        self.assertFalse(errors)

    @patch("tenders.services._ai_gateway_json")
    def test_llm_category_selection_retries_terms_only_after_empty_selection(self, gateway):
        gateway.side_effect = [
            ({"categories": []}, {"prompt_tokens": 100, "completion_tokens": 10}),
            ({"search_terms": ["защитный шлем", "каска"]}, {"prompt_tokens": 40, "completion_tokens": 8}),
            ({"categories": [{"source": "supplier-a", "category_id": "helmet", "priority": 1}]}, {"prompt_tokens": 120, "completion_tokens": 10}),
            ({"categories": [{"source": "supplier-a", "category_id": "helmet", "keep": True}]}, {"prompt_tokens": 30, "completion_tokens": 5}),
        ]
        tasks, usage, errors = _select_catalog_category_tasks(
            {"name": "Защитная каска"}, {"item": "каска"}, [
                {
                    "source": "supplier-a", "category_id": "helmet", "name": "Защитные шлемы и каски",
                    "parent_id": "", "path": "Каталог A > Товар A", "specificity": 1,
                },
            ],
        )

        self.assertEqual(gateway.call_count, 4)
        self.assertEqual(tasks[0]["category_id"], "helmet")
        self.assertEqual(usage["prompt_tokens"], 290)
        self.assertEqual(usage["completion_tokens"], 33)
        self.assertEqual(usage["llm_calls"], 4)
        self.assertEqual(usage["selector_llm_calls"], 2)
        self.assertEqual(usage["verifier_llm_calls"], 1)
        self.assertEqual(usage["semantic_attempts"], 2)
        self.assertTrue(usage["retry_used"])
        self.assertFalse(errors)

    def test_catalog_search_uses_llm_selected_real_category_instead_of_broad_category(self):
        requested_categories = []
        seen_category_candidates = []

        class Client:
            base_url = "https://api.oasiscatalog.com"

            def get(self, path, params=None):
                if path == "/v4/categories":
                    return [
                        {"id": "vip", "name": "Одежда", "path": "vip/odezhda"},
                        {"id": "polo", "name": "Рубашки поло", "path": "odezhda/rubashki-polo"},
                    ]
                requested_categories.append(params.get("category"))
                if params.get("category") == "polo":
                    return [{
                        "id": "right", "article": "POLO-1", "group_id": "polo-1",
                        "name": "Рубашка поло мужская", "full_name": "Рубашка поло мужская, белая",
                        "categories": ["polo"], "colors": ["белый"], "materials": ["хлопок"],
                        "attributes": [{"name": "Плотность", "value": "180 г/м2"}],
                        "total_stock": 100, "price": 500,
                    }]
                return [{
                    "id": "wrong", "article": "MITTEN-1", "group_id": "mitten-1",
                    "name": "Варежки", "categories": ["vip"], "total_stock": 100, "price": 100,
                }]

        def selector(line, intent, candidates, attempted):
            seen_category_candidates.extend(candidates)
            return [{"source": "oasis", "category_id": "polo", "priority": 1}], {}, []

        result = catalog_candidates_for_line(
            {"name": "Футболка поло унисекс", "quantity": 10, "requirements": {"requirements": []}},
            limit=3,
            intent={
                "item": "рубашка поло", "product_class": "футболка",
                "categories": ["одежда", "футболки"], "synonyms": ["футболка поло"],
            },
            client=Client(), category_selector=selector, include_diagnostics=True,
        )

        self.assertEqual(requested_categories, ["polo"])
        self.assertEqual(result["candidates"][0]["external_id"], "right")
        self.assertEqual(result["attempts"][0]["category_tasks"][0]["category_id"], "polo")
        self.assertEqual(seen_category_candidates[0]["category_id"], "polo")

    def test_llm_category_selector_receives_complete_fixed_map(self):
        seen_category_ids = []
        seen_parent_ids = {}

        class Client:
            base_url = "https://api.oasiscatalog.com"

            def get(self, path, params=None):
                if path == "/v4/categories":
                    return [
                        {"id": "clothes", "name": "Одежда", "path": "odezhda"},
                        {"id": "raglan", "parent_id": "clothes", "name": "Регланы", "path": "odezhda/reglany"},
                        {"id": "mugs", "name": "Кружки", "path": "posuda/kruzhki"},
                    ]
                return []

        def selector(line, intent, candidates, attempted):
            seen_category_ids.extend(value["category_id"] for value in candidates)
            seen_parent_ids.update({value["category_id"]: value.get("parent_id", "") for value in candidates})
            return [], {}, []

        catalog_candidates_for_line(
            {"name": "Лонгслив", "quantity": 10},
            intent={"item": "лонгслив", "synonyms": ["футболка с длинным рукавом"]},
            client=Client(), category_selector=selector,
        )

        self.assertEqual(set(seen_category_ids), {"clothes", "raglan", "mugs"})
        self.assertEqual(seen_parent_ids["raglan"], "clothes")

    @patch.dict("os.environ", {"KNOWLEDGE_SYNC_TOKEN": "test-token"})
    @patch("tenders.views._select_catalog_category_tasks")
    def test_category_selection_test_runs_without_catalogue_product_search(self, selector):
        supplier = CatalogSupplier.objects.create(
            code="supplier-x", name="Supplier X", base_url="https://supplier.example",
        )
        CatalogCategory.objects.create(
            supplier=supplier, external_id="root", name="Одежда", path="Каталог > Одежда",
        )
        CatalogCategory.objects.create(
            supplier=supplier, external_id="polo", parent_external_id="root",
            name="Поло", path="Каталог > Одежда > Поло",
        )
        session = ProductionTrainingSession.objects.create(
            created_by=self.user,
            position_name="Футболка поло унисекс, цвет – белый",
            requirements={"requirements": [{"label": "Цвет", "value": "белый"}]},
            current_hypothesis={"catalog_intent": {"item": "рубашка поло"}},
        )
        selector.return_value = ([{
            "source": "supplier-x", "category_id": "polo", "name": "Поло",
            "path": "Каталог > Одежда > Поло", "priority": 1,
        }], {
            "input_tokens": 100, "output_tokens": 20, "llm_calls": 1,
            "semantic_attempts": 1, "retry_used": False,
            "semantic_representation": {"item": "рубашка поло", "search_terms": ["рубашка поло"]},
            "retrieval": {"considered_count": 2, "candidate_count": 1, "fragment_count": 2, "by_source": {"supplier-x": 2}},
        }, [])

        response = self.client.post(
            reverse("tender_category_selection_test", args=[session.pk]),
            HTTP_AUTHORIZATION="Bearer test-token",
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["category_tasks"][0]["category_id"], "polo")
        self.assertEqual(response.json()["usage"]["llm_calls"], 1)
        candidates = selector.call_args.args[2]
        self.assertEqual({value["category_id"] for value in candidates}, {"root", "polo"})
        self.assertEqual(next(value for value in candidates if value["category_id"] == "polo")["parent_id"], "root")
        self.assertEqual(CatalogProduct.objects.count(), 0)

    def test_catalog_search_returns_nearest_candidate_when_required_density_differs(self):
        class Client:
            base_url = "https://api.oasiscatalog.com"

            def get(self, path, params=None):
                if path == "/v4/categories":
                    return [{"id": "long", "name": "Лонгсливы", "path": "odezhda/longslivy"}]
                return [{
                    "id": "near", "article": "LS-1", "group_id": "long-1",
                    "name": "Футболка с длинным рукавом", "full_name": "Футболка с длинным рукавом белая",
                    "categories": ["long"], "colors": ["белый"],
                    "materials": ["хлопок 100%, плотность 190 г/м2"],
                    "total_stock": 100, "price": 700,
                }]

        result = catalog_candidates_for_line(
            {"name": "Лонгслив", "quantity": 50, "requirements": {"requirements": [
                {"label": "Цвет", "value": "белый"},
                {"label": "Плотность", "value": "141 г/м2"},
            ]}},
            intent={
                "item": "лонгслив", "synonyms": ["футболка с длинным рукавом"],
                "required": [{"label": "Плотность", "value": "141 г/м2", "weight": 1}],
                "constraints": [{
                    "field": "density", "operator": "gte", "values": ["141"],
                    "level": "required", "weight": 1, "missing_policy": "reject",
                }],
            },
            client=Client(),
        )

        self.assertEqual(result[0]["external_id"], "near")
        self.assertEqual(result[0]["fit"], "partial")
        self.assertTrue(any("Плотность" in value for value in result[0]["unknown"]))

    def test_catalog_search_falls_back_to_server_category_priority_when_selector_fails(self):
        requested_categories = []

        class Client:
            base_url = "https://api.oasiscatalog.com"

            def get(self, path, params=None):
                if path == "/v4/categories":
                    return [
                        {"id": "vip", "name": "Одежда", "path": "vip/odezhda"},
                        {"id": "polo", "name": "Рубашки поло", "path": "odezhda/rubashki-polo"},
                    ]
                requested_categories.append(params.get("category"))
                return [{
                    "id": "right", "article": "POLO-1", "group_id": "polo-1",
                    "name": "Рубашка поло мужская", "full_name": "Рубашка поло мужская, белая",
                    "categories": ["polo"], "colors": ["белый"], "materials": ["хлопок"],
                    "total_stock": 100, "price": 500,
                }]

        def broken_selector(*args):
            raise RuntimeError("LLM unavailable")

        result = catalog_candidates_for_line(
            {"name": "Футболка поло унисекс", "quantity": 10, "requirements": {"requirements": []}},
            intent={
                "item": "рубашка поло", "product_class": "футболка",
                "categories": ["одежда"], "synonyms": ["футболка поло"],
            },
            client=Client(), category_selector=broken_selector, include_diagnostics=True,
        )

        self.assertEqual(requested_categories, ["polo"])
        self.assertEqual(result["candidates"][0]["external_id"], "right")
        self.assertTrue(result["category_errors"])

    def test_selected_gifts_category_does_not_depend_on_cached_search_text(self):
        gifts = CatalogSupplier.objects.create(code="gifts", name="gifts.ru", base_url="https://gifts.ru")
        CatalogCategory.objects.create(
            supplier=gifts, external_id="vacuum", name="Термокружки", path="Посуда / Термокружки",
        )
        CatalogProduct.objects.create(
            supplier=gifts, external_id="travel-mug", article="G-1", name="Термокружка Voyager",
            full_name="Термокружка Voyager, 500 мл", category_ids=["vacuum"],
            total_stock=100, discount_price=500, search_text="",
        )

        class Client:
            base_url = "https://api.oasiscatalog.com"

            def get(self, path, params=None):
                return []

        def selector(line, intent, candidates, attempted):
            return [{"source": "gifts", "category_id": "vacuum", "priority": 1}], {}, []

        result = catalog_candidates_for_line(
            {"name": "Термокружка", "quantity": 10},
            intent={"item": "термокружка", "categories": ["термокружки"]},
            client=Client(), category_selector=selector,
        )

        self.assertEqual([value["external_id"] for value in result], ["travel-mug"])

    def test_selected_gifts_category_is_filtered_before_candidate_limit(self):
        gifts = CatalogSupplier.objects.create(code="gifts", name="gifts.ru", base_url="https://gifts.ru")
        CatalogCategory.objects.create(
            supplier=gifts, external_id="long-sleeve", name="Лонгсливы",
            path="Одежда / Футболки с длинным рукавом",
        )
        CatalogProduct.objects.bulk_create([
            CatalogProduct(
                supplier=gifts, external_id=f"other-{index}", name="Кружка",
                category_ids=["mugs"], total_stock=10, search_text="кружка",
            )
            for index in range(1501)
        ])
        CatalogProduct.objects.create(
            supplier=gifts, external_id="wanted-long-sleeve", article="LS-1",
            name="Лонгслив унисекс", category_ids=["long-sleeve"],
            total_stock=100, discount_price=700, search_text="лонгслив унисекс",
        )

        class Client:
            base_url = "https://api.oasiscatalog.com"

            def get(self, path, params=None):
                return []

        def selector(line, intent, candidates, attempted):
            return [{"source": "gifts", "category_id": "long-sleeve", "priority": 1}], {}, []

        result = catalog_candidates_for_line(
            {"name": "Лонгслив", "quantity": 10},
            intent={"item": "лонгслив", "synonyms": ["футболка с длинным рукавом"]},
            client=Client(), category_selector=selector,
        )

        self.assertEqual([value["external_id"] for value in result], ["wanted-long-sleeve"])

    def test_catalog_search_combines_cached_gifts_with_oasis_by_relevance(self):
        gifts = CatalogSupplier.objects.create(code="gifts", name="gifts.ru", base_url="https://api2.gifts.ru/export/v2")
        CatalogProduct.objects.create(
            supplier=gifts, external_id="gifts-vest", article="G-1", name="Жилет утеплённый",
            full_name="Жилет утеплённый чёрный", materials=["полиэстер"], colors=["черный"],
            total_stock=20, discount_price=1200, search_text="жилет утепленный черный полиэстер одежда",
            product_url="https://gifts.ru/catalog/G-1",
        )

        class Client:
            base_url = "https://api.oasiscatalog.com"

            def get(self, path, params=None):
                if path == "/v4/categories":
                    return [{"id": 10, "name": "Жилеты", "path": "categories/odezhda/zhilety"}]
                return [{"id": "oasis-vest", "article": "O-1", "name": "Жилет для работы", "full_name": "Жилет для работы спецодежда", "materials": ["полиэстер"], "colors": ["черный"], "total_stock": 100, "categories": [10]}]

        line = {
            "name": "Жилет", "quantity": "10",
            "requirements": {"requirements": [{"label": "Материал", "value": "полиэстер"}, {"label": "Цвет", "value": "черный"}]},
        }

        candidates = catalog_candidates_for_line(line, limit=3, intent={"product_class": "жилет"}, client=Client())

        self.assertEqual({value["supplier_code"] for value in candidates}, {"oasis", "gifts"})
        self.assertEqual(candidates[0]["supplier_code"], "gifts")

    def test_purple_vest_from_gifts_is_ranked_before_nonmatching_oasis_vest(self):
        gifts = CatalogSupplier.objects.create(code="gifts", name="gifts.ru", base_url="https://api2.gifts.ru/export/v2")
        CatalogProduct.objects.create(
            supplier=gifts, external_id="249789", article="26728.78", name="Жилет детский Kama Kids, фиолетовый",
            full_name="Жилет детский Kama Kids, фиолетовый", materials=["полиэстер 100%"], colors=["фиолетовый"],
            total_stock=45, discount_price=2600, search_text="жилет детский kama kids фиолетовый полиэстер одежда",
            product_url="https://gifts.ru/id/249789",
        )

        class Client:
            base_url = "https://api.oasiscatalog.com"

            def get(self, path, params=None):
                if path == "/v4/categories":
                    return [{"id": 10, "name": "Жилеты", "path": "categories/odezhda/zhilety"}]
                return [{"id": "oasis-vest", "article": "O-1", "name": "Жилет", "full_name": "Жилет чёрный", "colors": ["черный"], "total_stock": 100, "categories": [10]}]

        line = {"name": "Жилет", "quantity": "10", "requirements": {"requirements": [{"label": "Цвет", "value": "фиолетовый"}]}}
        candidates = catalog_candidates_for_line(line, limit=3, intent={"product_class": "жилет"}, client=Client())

        self.assertEqual(candidates[0]["external_id"], "249789")
        self.assertEqual(candidates[0]["supplier_code"], "gifts")
        self.assertEqual(candidates[0]["fit"], "exact")

    def test_gifts_sync_persists_filtered_xml_rows_without_images(self):
        class Client:
            base_url = "https://api2.gifts.ru/export/v2"

            def open(self, path):
                payloads = {
                    "catalogue/product.xml": "<doct><product product_id='v1'><code>V-1</code><name>Жилет фиолетовый</name><matherial>полиэстер</matherial><small_image src='reviewer/v.webp'/></product></doct>",
                    "catalogue/tree.xml": "<doct><page page_id='1' name='Жилеты'><product product='v1'/></page></doct>",
                    "catalogue/stock.xml": "<doct><stock product_id='v1'><free>12</free><dealerprice>1000</dealerprice></stock></doct>",
                    "catalogue/filters.xml": "<root><filtertypes/></root>",
                }
                return StringIO(payloads[path])

        run = sync_gifts_catalog(Client(), category="жилеты")
        product = CatalogProduct.objects.get(supplier__code="gifts", external_id="v1")

        self.assertEqual(run.status, "success")
        self.assertEqual(product.total_stock, 12)
        self.assertEqual(product.image_url, "https://files.gifts.ru/reviewer/v.webp")

    def test_gifts_category_sync_uses_category_only_xml_and_preserves_hierarchy(self):
        class Client:
            base_url = "https://api2.gifts.ru/export/v2"
            opened = []
            xml = """<doct><page page_id='10' name='Одежда'><page page_id='20' name='Поло'><page page_id='30' name='Мужские поло'/></page></page></doct>"""

            def open(self, path):
                self.opened.append(path)
                return StringIO(self.xml)

        client = Client()
        categories = sync_gifts_categories(client)

        self.assertEqual(client.opened, ["catalogue/treeWithoutProducts.xml"])
        self.assertEqual(categories["20"], "Одежда > Поло")
        polo = CatalogCategory.objects.get(supplier__code="gifts", external_id="20")
        self.assertEqual(polo.parent_external_id, "10")
        self.assertEqual(polo.path, "Одежда > Поло")

        client.xml = """<doct><page page_id='10' name='Текстиль'><page page_id='20' name='Поло'/></page></doct>"""
        categories = sync_gifts_categories(client)

        self.assertEqual(categories["20"], "Текстиль > Поло")
        self.assertEqual(CatalogCategory.objects.filter(supplier__code="gifts").count(), 3)
        self.assertFalse(CatalogCategory.objects.get(supplier__code="gifts", external_id="30").is_active)

    @patch.dict("os.environ", {"TIMEWEB_EMBEDDINGS_ENABLED": "1", "TIMEWEB_EMBEDDING_MODEL": "test-embedding"})
    @patch("tenders.services._embedding_vectors")
    def test_category_sync_embeds_only_missing_or_changed_semantic_text(self, embedding_vectors):
        embedding_vectors.side_effect = lambda texts, model=None: [
            [float(index + 1), 1.0] for index, _ in enumerate(texts)
        ]

        class Client:
            base_url = "https://api2.gifts.ru/export/v2"
            xml = """<doct><page page_id='10' name='Одежда'><page page_id='20' name='Поло'/></page></doct>"""

            def open(self, path):
                return StringIO(self.xml)

        client = Client()
        sync_gifts_categories(client)
        first_call_count = embedding_vectors.call_count
        polo = CatalogCategory.objects.get(supplier__code="gifts", external_id="20")

        self.assertTrue(polo.embedding)
        self.assertEqual(polo.embedding_model, "test-embedding")
        self.assertTrue(polo.embedding_text_hash)
        self.assertEqual(first_call_count, 1)

        sync_gifts_categories(client)
        self.assertEqual(embedding_vectors.call_count, first_call_count)

        client.xml = """<doct><page page_id='10' name='Одежда'><page page_id='20' name='Рубашки поло'/></page></doct>"""
        sync_gifts_categories(client)

        self.assertEqual(embedding_vectors.call_count, first_call_count + 1)
        self.assertEqual(len(embedding_vectors.call_args.args[0]), 1)

    @patch("tenders.catalog.catalog_candidates_for_line")
    @patch("tenders.services._ai_gateway_json")
    def test_llm_interprets_words_before_backend_catalog_search(self, gateway, catalog_search):
        gateway.return_value = ({
            "product_type": "textile_merch", "summary": "Футболка", "confidence": .5,
            "facts": [], "route": {"reason": "Готовое изделие", "processes": [{"name": "Закупка готового изделия"}]},
            "costs": [], "questions": [], "assumptions": [], "matched_example_ids": [], "understood_changes": [],
            "catalog_intent": {"product_class": "футболка", "synonyms": ["майка"], "hard_constraints": ["лаймово-зелёный цвет"], "preferences": []},
        }, {})
        catalog_search.return_value = [{
            "id": "shirt", "external_id": "shirt", "supplier_code": "oasis",
            "article": "SHIRT-1", "name": "Футболка", "price": "500.00",
            "stock": 500, "url": "https://www.oasiscatalog.com/item/shirt",
            "fit": "exact", "matches": ["Тип товара: футболка"],
            "mismatches": [], "unknown": [],
        }]
        line = {"name": "Майка брендированная", "quantity": 160, "requirements": {"requirements": []}}

        result = build_training_hypothesis(line)

        self.assertEqual(result["catalog_intent"]["product_class"], "футболка")
        self.assertEqual(result["catalog_candidates"][0]["name"], "Футболка")
        self.assertEqual(catalog_search.call_args.kwargs["intent"]["synonyms"], ["майка"])
        self.assertEqual(result["catalog_selection"]["selection_mode"], "automatic")
        self.assertEqual(result["totals"]["material_unit"], "500.00")
        self.assertEqual(result["totals"]["cost_total"], "80000.00")
        self.assertEqual(result["costs"][0]["source_type"], "catalog")

    @patch("tenders.catalog.catalog_candidates_for_line")
    @patch("tenders.services._ai_gateway_json")
    def test_llm_catalog_plan_keeps_specific_category_from_full_title(self, gateway, catalog_search):
        gateway.return_value = ({
            "product_type": "textile_merch", "summary": "Поло", "confidence": .8,
            "facts": [], "route": {"reason": "Готовое изделие", "processes": [{"name": "Закупка готового изделия"}]},
            "costs": [], "questions": [], "assumptions": [], "matched_example_ids": [], "understood_changes": [],
            "catalog_intent": {
                "product_class": "футболка", "item": "рубашка поло",
                "categories": ["поло", "рубашка поло", "polo shirt"], "synonyms": ["polo shirt"],
                "required": [{"label": "Цвет", "value": "белый", "weight": 1}],
                "preferred": [{"label": "Материал", "value": "хлопок", "weight": .8}],
                "secondary": [], "search_fields": ["category", "name", "attributes"],
                "ranking": [{"criterion": "соответствие типу товара", "weight": 1}],
                "fallback_queries": [{"terms": ["поло", "рубашка поло"], "relaxable": False}],
            },
        }, {})
        catalog_search.return_value = []

        result = build_training_hypothesis({"name": "Футболка поло унисекс", "quantity": 10, "requirements": {"requirements": []}})

        intent = result["catalog_intent"]
        self.assertEqual(intent["categories"], ["поло", "рубашка поло", "polo shirt"])
        self.assertEqual(intent["required"][0]["label"], "Цвет")
        self.assertEqual(catalog_search.call_args.kwargs["intent"]["categories"][0], "поло")

    @patch("tenders.catalog.catalog_candidates_for_line")
    @patch("tenders.services._ai_gateway_json")
    def test_empty_catalog_result_retries_with_untried_real_categories(self, gateway, catalog_search):
        initial = {
            "product_type": "textile_merch", "summary": "Термокружка", "confidence": .6,
            "facts": [], "route": {"reason": "Готовое изделие", "processes": [{"name": "Закупка готового изделия"}]},
            "costs": [], "questions": [], "assumptions": [], "matched_example_ids": [], "understood_changes": [],
            "catalog_intent": {
                "item": "термокружка", "categories": ["термокружка"],
                "required": [{"label": "Объём", "value": "500 мл", "weight": 1}],
                "preferred": [], "secondary": [], "ranking": [],
            },
        }
        gateway.return_value = (initial, {})
        catalog_search.side_effect = [
            {
                "candidates": [],
                "sources": {"oasis": {"status": "success", "message": "", "received": 20}, "gifts": {"status": "success", "message": "", "received": 10}},
                "attempts": [{
                    "mode": "selected_categories", "candidate_count": 0,
                    "category_tasks": [{"source": "oasis", "category_id": "mugs"}],
                }],
            },
            {
                "candidates": [{
                    "id": "mug", "external_id": "mug", "supplier_code": "gifts", "supplier_name": "gifts.ru",
                    "article": "M-1", "name": "Термокружка", "price": None, "stock": 20, "url": "",
                    "fit": "partial", "matches": ["Тип товара: термокружка"], "mismatches": [], "unknown": ["Объём не указан"],
                }],
                "sources": {"oasis": {"status": "success", "message": "", "received": 20}, "gifts": {"status": "success", "message": "", "received": 10}},
                "attempts": [{
                    "mode": "selected_categories", "candidate_count": 1,
                    "category_tasks": [{"source": "gifts", "category_id": "vacuum-mugs"}],
                }],
            },
        ]

        result = build_training_hypothesis({"name": "Термокружка 500 мл", "quantity": 10, "requirements": {"requirements": []}})

        self.assertEqual(catalog_search.call_count, 2)
        self.assertEqual(result["catalog_candidates"][0]["id"], "mug")
        self.assertEqual(result["catalog_intent"]["required"][0]["label"], "Объём")
        self.assertEqual(len(result["catalog_attempts"]), 2)
        self.assertEqual(catalog_search.call_args.kwargs["excluded_category_tasks"], [
            {"source": "oasis", "category_id": "mugs"},
        ])

    @patch("tenders.catalog.catalog_candidates_for_line")
    @patch("tenders.services._ai_gateway_json")
    def test_empty_category_attempts_finish_with_full_text_search(self, gateway, catalog_search):
        gateway.return_value = ({
            "product_type": "textile_merch", "summary": "Поло", "confidence": .6,
            "facts": [], "route": {"reason": "Готовое изделие", "processes": [{"name": "Закупка готового изделия"}]},
            "costs": [], "questions": [], "assumptions": [], "matched_example_ids": [], "understood_changes": [],
            "catalog_intent": {"item": "рубашка поло", "categories": ["поло"]},
        }, {})
        source_status = {"oasis": {"status": "success", "message": "", "received": 20}}
        catalog_search.side_effect = [
            {"candidates": [], "sources": source_status, "attempts": [{
                "mode": "selected_categories", "category_tasks": [{"source": "oasis", "category_id": "polo-1"}],
            }]},
            {"candidates": [], "sources": source_status, "attempts": [{
                "mode": "selected_categories", "category_tasks": [{"source": "oasis", "category_id": "polo-2"}],
            }]},
            {"candidates": [], "sources": source_status, "attempts": [{
                "mode": "full_text", "category_tasks": [],
            }]},
        ]

        result = build_training_hypothesis(
            {"name": "Футболка поло", "quantity": 10, "requirements": {"requirements": []}},
        )

        self.assertEqual(catalog_search.call_count, 3)
        self.assertTrue(catalog_search.call_args.kwargs["force_full_text"])
        self.assertEqual([attempt["mode"] for attempt in result["catalog_attempts"]], [
            "selected_categories", "selected_categories", "full_text",
        ])

    @patch("tenders.catalog.catalog_candidates_for_line", return_value=[])
    @patch("tenders.services._ai_gateway_json")
    def test_feedback_can_rebuild_catalog_ranking_weights(self, gateway, catalog_search):
        gateway.return_value = ({
            "product_type": "textile_merch", "summary": "Промо-футболка", "confidence": .7,
            "facts": [], "route": {"reason": "Готовое изделие", "processes": [{"name": "Закупка готового изделия"}]},
            "costs": [], "questions": [], "assumptions": [], "matched_example_ids": [],
            "understood_changes": ["Цена важнее состава, цвет обязателен"],
            "catalog_intent": {
                "item": "футболка", "categories": ["футболка"],
                "required": [{"label": "Цвет", "value": "синий", "weight": 1}],
                "preferred": [{"label": "Состав", "value": "хлопок", "weight": .2}],
                "secondary": [],
                "ranking": [{"criterion": "цена", "weight": 1}, {"criterion": "состав", "weight": .2}],
            },
            "catalog_operations": [
                {"op": "set_priority", "field": "price", "priority": "critical"},
                {"op": "set_priority", "field": "material", "priority": "low"},
                {"op": "require", "field": "color", "values": ["синий"]},
                {"op": "prefer", "field": "material", "values": ["хлопок"], "priority": "low"},
            ],
        }, {})

        result = build_training_hypothesis(
            {"name": "Промо-футболка", "quantity": 1000, "requirements": {"requirements": []}},
            current={"catalog_intent": {"categories": ["футболка"], "ranking": [{"criterion": "состав", "weight": 1}]}},
            feedback="Для этого промо-тиража цена важнее состава, но синий цвет обязателен.",
        )

        weights = {value["criterion"]: value["weight"] for value in result["catalog_intent"]["ranking"]}
        self.assertEqual(weights, {"price": 1.0, "material": .3})
        constraints = {(value["field"], value["level"]): value for value in result["catalog_intent"]["constraints"]}
        self.assertEqual(constraints[("color", "required")]["values"], ["синий"])
        self.assertEqual(constraints[("material", "preferred")]["weight"], .3)

    def test_catalog_feedback_operations_patch_existing_plan_without_losing_rules(self):
        current = {
            "categories": ["поло"], "synonyms": ["рубашка поло"],
            "required": [
                {"label": "Цвет", "value": "белый", "weight": 1},
                {"label": "Пол", "value": "мужской или унисекс, исключить женский", "weight": 1},
            ],
            "ranking": [{"criterion": "состав", "weight": .6}],
        }
        operations = [
            {"op": "set_missing_policy", "field": "gender", "value": "allow_with_penalty"},
            {"op": "forbid", "field": "gender", "values": ["female"]},
            {"op": "add_alias", "values": ["футболка поло"]},
            {"op": "set_priority", "field": "price", "relation": "higher_than", "target_field": "material"},
        ]

        updated, applied, errors = _apply_catalog_operations(current, operations)

        self.assertFalse(errors)
        self.assertEqual(len(applied), 4)
        self.assertEqual([value["label"] for value in updated["required"]], ["Цвет"])
        self.assertEqual(updated["synonyms"], ["рубашка поло", "футболка поло"])
        self.assertEqual(updated["constraints"][0]["operator"], "not_in")
        self.assertEqual(updated["constraints"][0]["missing_policy"], "allow_with_penalty")
        weights = {value["criterion"]: value["weight"] for value in updated["ranking"]}
        self.assertGreater(weights["price"], weights["состав"])

    def test_catalog_contract_rejects_unknown_operation_instead_of_silently_applying_it(self):
        updated, applied, errors = _apply_catalog_operations(
            {"categories": ["поло"]},
            [{"op": "magically_fix", "field": "gender", "values": ["female"]}],
        )

        self.assertEqual(updated["categories"], ["поло"])
        self.assertFalse(applied)
        self.assertIn("magically_fix", errors[0])

    def test_catalog_contract_accepts_declared_filter_ranking_alias_and_source_operations(self):
        operations = [
            {"op": "allow", "field": "color", "values": ["синий", "белый"]},
            {"op": "forbid", "field": "gender", "values": ["female"]},
            {"op": "require", "field": "branding", "values": ["вышивка"]},
            {"op": "prefer", "field": "material", "values": ["хлопок"]},
            {"op": "deprioritize", "field": "material", "values": ["полиэстер"]},
            {"op": "ignore", "field": "name"},
            {"op": "set_priority", "field": "price", "priority": "high"},
            {"op": "add_alias", "values": ["футболка поло"]},
            {"op": "remove_alias", "values": ["старая категория"]},
            {"op": "set_missing_policy", "field": "gender", "value": "allow_with_penalty"},
            {"op": "lte", "field": "price", "values": ["500"]},
            {"op": "gte", "field": "stock", "values": ["100"]},
            {"op": "between", "field": "density", "values": ["180", "220"]},
            {"op": "source_only", "values": ["oasis"]},
            {"op": "prefer_source", "values": ["gifts"]},
            {"op": "set_scope", "values": ["одежда", "текстиль"]},
            {"op": "remove_rule", "field": "branding"},
        ]

        updated, applied, errors = _apply_catalog_operations(
            {"categories": ["поло"], "synonyms": ["старая категория"]}, operations,
        )

        self.assertFalse(errors)
        self.assertEqual(len(applied), len(operations))
        self.assertEqual(updated["synonyms"], ["футболка поло"])
        self.assertEqual(updated["allowed_sources"], ["oasis"])
        self.assertEqual(updated["preferred_sources"], ["gifts"])
        self.assertEqual(updated["rule_scope"], ["одежда", "текстиль"])
        self.assertFalse(any(value["field"] == "branding" for value in updated["constraints"]))
        self.assertEqual(applied[9], {
            "op": "set_missing_policy", "field": "gender", "value": "allow_with_penalty",
        })

    def test_catalog_contract_compiles_deprioritized_missing_value_to_missing_policy(self):
        updated, applied, errors = _apply_catalog_operations({"categories": ["поло"]}, [
            {"op": "allow", "field": "gender", "values": ["male", "unisex"]},
            {"op": "forbid", "field": "gender", "values": ["female"]},
            {"op": "deprioritize", "field": "gender", "values": ["missing"]},
        ])

        self.assertFalse(errors)
        self.assertEqual(applied[-1], {"op": "set_missing_policy", "field": "gender", "value": "allow_with_penalty"})
        self.assertTrue(all(value["missing_policy"] == "allow_with_penalty" for value in updated["constraints"]))

    @patch("tenders.catalog.catalog_candidates_for_line", return_value=[])
    @patch("tenders.services._ai_gateway_json")
    def test_feedback_uses_catalog_operations_as_patch_over_current_intent(self, gateway, catalog_search):
        gateway.return_value = ({
            "product_type": "textile_merch", "summary": "Поло", "confidence": .7,
            "facts": [], "route": {"reason": "Закупка", "processes": [{"name": "Закупка готового изделия"}]},
            "costs": [], "questions": [], "assumptions": [], "matched_example_ids": [],
            "understood_changes": ["Женские модели исключены"],
            "catalog_intent": {"categories": ["ошибочно переписанная категория"]},
            "catalog_operations": [
                {"op": "forbid", "field": "gender", "values": ["female"]},
                {"op": "set_missing_policy", "field": "gender", "value": "allow_with_penalty"},
            ],
        }, {})
        current = {"catalog_intent": {
            "categories": ["поло"],
            "required": [{"label": "Цвет", "value": "белый", "weight": 1}],
        }}

        result = build_training_hypothesis(
            {"name": "Поло унисекс", "quantity": 50, "requirements": {"requirements": []}},
            current=current, feedback="Исключи женские модели, модели без пола допустимы.",
        )

        self.assertEqual(result["catalog_intent"]["categories"], ["поло"])
        self.assertEqual(result["catalog_intent"]["required"][0]["label"], "Цвет")
        self.assertEqual(result["catalog_intent"]["constraints"][0]["values"], ["female"])
        self.assertEqual(len(result["catalog_operations_applied"]), 2)

    @patch("tenders.catalog.catalog_candidates_for_line", return_value=[])
    @patch("tenders.services._ai_gateway_json")
    def test_empty_catalog_part_uses_compact_feedback_translator(self, gateway, catalog_search):
        hypothesis = {
            "product_type": "textile_merch", "summary": "Поло", "confidence": .7,
            "facts": [], "route": {"reason": "Закупка", "processes": [{"name": "Закупка готового изделия"}]},
            "costs": [], "questions": [], "assumptions": [], "matched_example_ids": [], "understood_changes": [],
        }
        gateway.side_effect = [
            (hypothesis, {}),
            ({"operations": [{"op": "forbid", "field": "gender", "values": ["female"]}]}, {"prompt_tokens": 20, "completion_tokens": 10}),
        ]
        current = {"catalog_intent": {"categories": ["поло"], "required": [{"label": "Цвет", "value": "белый", "weight": 1}]}}

        result = build_training_hypothesis(
            {"name": "Поло унисекс", "quantity": 50, "requirements": {"requirements": []}},
            current=current, feedback="Женские модели исключи.",
        )

        self.assertEqual(gateway.call_count, 2)
        self.assertEqual(result["catalog_intent"]["categories"], ["поло"])
        self.assertEqual(result["catalog_intent"]["constraints"][0]["operator"], "not_in")
        self.assertEqual(result["usage"], {"prompt_tokens": 20, "completion_tokens": 10})

    @patch("tenders.catalog.catalog_candidates_for_line", return_value=[])
    @patch("tenders.services._ai_gateway_json")
    def test_empty_feedback_translation_preserves_plan_and_blocks_confirmation(self, gateway, catalog_search):
        hypothesis = {
            "product_type": "textile_merch", "summary": "Поло", "confidence": .7,
            "facts": [], "route": {"reason": "Закупка", "processes": [{"name": "Закупка готового изделия"}]},
            "costs": [], "questions": [], "assumptions": [], "matched_example_ids": [], "understood_changes": [],
        }
        gateway.side_effect = [(hypothesis, {}), ({"operations": []}, {})]
        current = {"catalog_intent": {"categories": ["поло"], "required": [{"label": "Цвет", "value": "белый", "weight": 1}]}}

        result = build_training_hypothesis(
            {"name": "Поло унисекс", "quantity": 50, "requirements": {"requirements": []}},
            current=current, feedback="Исключи неподходящие модели.",
        )

        self.assertEqual(result["catalog_intent"]["categories"], ["поло"])
        self.assertTrue(result["catalog_contract_errors"])
        self.assertTrue(result["learning_warnings"])

    @patch("tenders.views.refresh_training_example_embedding", return_value=True)
    def test_confirmed_feedback_weights_are_saved_with_training_example(self, refresh_embedding):
        self.user.is_superuser = True
        self.user.save(update_fields=["is_superuser"])
        self.client.force_login(self.user)
        production_type = ProductionType.objects.create(code="catalog-feedback", name="Каталожный товар")
        intent = {
            "categories": ["футболка"],
            "required": [{"label": "Цвет", "value": "синий", "weight": 1}],
            "preferred": [{"label": "Состав", "value": "хлопок", "weight": .2}],
            "ranking": [{"criterion": "цена", "weight": 1}, {"criterion": "состав", "weight": .2}],
        }
        session = ProductionTrainingSession.objects.create(
            created_by=self.user,
            position_name="Промо-футболка",
            requirements={"requirements": []},
            current_hypothesis={
                "stage": "training_dialogue", "product_type": production_type.code,
                "route": {"name": "Закупка готового изделия", "reason": "Подтверждено", "steps": ["Закупка готового изделия"], "processes": [{"name": "Закупка готового изделия"}]},
                "costs": [], "totals": {}, "catalog_intent": intent, "learning_warnings": [],
            },
        )

        response = self.client.post(reverse("tender_confirm_production_type"), {
            "payload": json.dumps({"session_id": session.pk, "line": {"name": "Промо-футболка", "quantity": 1000}}),
        })

        self.assertEqual(response.status_code, 200)
        example = ProductionTrainingExample.objects.get(pk=response.json()["example_id"])
        self.assertEqual(example.routes[0]["catalog_intent"]["ranking"], intent["ranking"])

    @patch("tenders.catalog.catalog_candidates_for_line", side_effect=ValueError("broken external row"))
    @patch("tenders.services._ai_gateway_json")
    def test_catalog_failure_does_not_discard_valid_production_hypothesis(self, gateway, catalog_search):
        gateway.return_value = ({
            "product_type": "textile_merch", "summary": "Футболка", "confidence": .5,
            "facts": [], "route": {"reason": "Готовое изделие", "processes": [{"name": "Закупка готового изделия"}]},
            "costs": [], "questions": [], "assumptions": [], "matched_example_ids": [], "understood_changes": [],
            "catalog_intent": {"product_class": "футболка", "synonyms": ["майка"], "hard_constraints": [], "preferences": []},
        }, {})

        result = build_training_hypothesis({"name": "Майка", "quantity": 10, "requirements": {"requirements": []}})

        self.assertEqual(result["product_type"], "textile_merch")
        self.assertEqual(result["catalog_candidates"], [])
        self.assertIn("Oasis", result["catalog_warning"])

    @patch("tenders.catalog.catalog_candidates_for_line", return_value=[])
    @patch("tenders.services._ai_gateway_json")
    def test_llm_catalog_price_is_not_used_without_current_catalog_candidate(self, gateway, catalog_search):
        gateway.return_value = ({
            "product_type": "textile_merch", "summary": "Лонгслив", "confidence": .8,
            "facts": [], "route": {"reason": "Готовое изделие", "processes": [{"name": "Закупка готового изделия"}]},
            "costs": [{
                "process_name": "Закупка готового изделия", "category": "material",
                "name": "Старый товар из обучения", "amount_total": 38000,
                "source": "gifts.ru", "source_type": "catalog",
                "recipe": {"method": "unit_rate", "inputs": {"unit_rate": 760}},
            }],
            "questions": [], "assumptions": [], "matched_example_ids": [1], "understood_changes": [],
            "catalog_intent": {"item": "лонгслив"},
        }, {})

        result = build_training_hypothesis(
            {"name": "Лонгслив", "quantity": 50, "requirements": {"requirements": []}},
        )

        self.assertFalse(any(value.get("source_type") == "catalog" for value in result["costs"]))
        self.assertEqual(result["totals"]["material_unit"], "0.00")
        self.assertTrue(any("текущим поиском" in value for value in result["learning_warnings"]))

    def test_catalog_search_does_not_repeat_color_variants_as_alternatives(self):
        class Client:
            base_url = "https://api.oasiscatalog.com"
            def get(self, path, params=None):
                if path == "/v4/categories":
                    return [{"id": 10, "name": "Футболки", "path": "categories/tekstil/futbolki"}]
                return [{"id": external_id, "article": external_id, "group_id": "same-shirt", "name": "Футболка", "full_name": f"Футболка {color}", "colors": [color], "materials": ["хлопок"], "price": 500, "total_stock": 100, "categories": [10]} for external_id, color in (("blue", "синий"), ("red", "красный"))]

        candidates = catalog_candidates_for_line({"name": "Футболка", "quantity": 10}, limit=3, intent={"product_class": "футболка"}, client=Client())

        self.assertEqual(len(candidates), 1)

    def test_catalog_search_reads_later_pages_and_matches_lime_to_green_apple(self):
        target = {
            "id": "00000008300", "article": "3100868S", "group_id": "apple-shirt", "color_group_id": "00000008300",
            "name": "Футболка Super Heavy Super Club мужская",
            "full_name": "Футболка Super Heavy Super Club мужская, зеленое яблоко",
            "colors": [{"name": "зеленое яблоко"}], "materials": ["хлопок"],
            "attributes": [{"name": "Плотность", "value": "180 г/м2"}],
            "branding": ["DTF (Полноцвет)"], "discount_price": "510.60",
            "total_stock": 200, "categories": [10],
        }

        class Client:
            base_url = "https://api.oasiscatalog.com"
            def get(self, path, params=None):
                if path == "/v4/categories":
                    return [{"id": 10, "name": "Футболки", "path": "categories/tekstil/futbolki"}]
                if params.get("offset") == 0:
                    return [{
                        "id": f"dummy-{index}", "article": f"D-{index}", "group_id": f"dummy-{index}",
                        "name": "Футболка", "full_name": "Футболка зеленая",
                        "colors": [{"name": "зеленый"}], "materials": ["хлопок"],
                        "attributes": [{"name": "Плотность", "value": "180 г/м2"}],
                        "branding": ["DTF"], "price": "500", "total_stock": 200, "categories": [10],
                    } for index in range(500)]
                return [target]

        line = {"name": "Майка брендированная", "quantity": 160, "requirements": {"requirements": [
            {"label": "Материал", "value": "хлопок"},
            {"label": "Цвет", "value": "лаймово-зелёный"},
            {"label": "Плотность", "value": "не менее 180 г/м²"},
            {"label": "Нанесение", "value": "DTF"},
        ]}}

        candidates = catalog_candidates_for_line(line, limit=3, intent={"product_class": "футболка"}, client=Client())

        self.assertEqual(candidates[0]["external_id"], "00000008300")
        self.assertEqual(candidates[0]["fit"], "exact")
        self.assertEqual(candidates[0]["price"], "510.60")
        self.assertEqual(candidates[0]["supplier_name"], "Oasis")
        self.assertEqual(candidates[0]["supplier_site"], "oasiscatalog.com")
        self.assertTrue(any("семейство: lime" in value for value in candidates[0]["matches"]))

    def test_catalog_search_uses_name_shade_as_soft_hint_with_explicit_parent_color(self):
        gifts = CatalogSupplier.objects.create(code="gifts", name="gifts.ru", base_url="https://api2.gifts.ru/export/v2")
        CatalogProduct.objects.create(
            supplier=gifts, external_id="lime-shirt", article="1376.89", name="Футболка унисекс Regent 150, лайм",
            full_name="Футболка унисекс Regent 150, лайм", materials=["хлопок"], colors=["зеленый"],
            raw_data={"name_colors": ["лайм"]}, total_stock=100, discount_price=404,
            search_text="футболка лайм зеленый хлопок", product_url="https://gifts.ru/id/16224",
        )

        class Client:
            base_url = "https://api.oasiscatalog.com"

            def get(self, path, params=None):
                if path == "/v4/categories":
                    return [{"id": 10, "name": "Футболки", "path": "categories/tekstil/futbolki"}]
                return []

        line = {"name": "Футболка", "quantity": "10", "requirements": {"requirements": [{"label": "Цвет", "value": "лайм"}]}}
        candidates = catalog_candidates_for_line(line, limit=1, intent={"product_class": "футболка"}, client=Client())

        self.assertEqual(candidates[0]["external_id"], "lime-shirt")
        self.assertEqual(candidates[0]["fit"], "exact")
        self.assertTrue(any("Цвет: зеленый" in value for value in candidates[0]["matches"]))

    def test_catalog_search_uses_price_after_equal_relevance(self):
        class Client:
            base_url = "https://api.oasiscatalog.com"

            def get(self, path, params=None):
                if path == "/v4/categories":
                    return [{"id": 10, "name": "Футболки", "path": "categories/tekstil/futbolki"}]
                return [
                    {"id": "expensive", "article": "E", "group_id": "expensive", "name": "Футболка", "full_name": "Футболка", "materials": ["хлопок"], "colors": ["белый"], "price": "700", "categories": [10], "total_stock": 100},
                    {"id": "cheap", "article": "C", "group_id": "cheap", "name": "Футболка", "full_name": "Футболка", "materials": ["хлопок"], "colors": ["белый"], "price": "500", "categories": [10], "total_stock": 100},
                ]

        line = {"name": "Футболка", "quantity": "10", "requirements": {"requirements": [{"label": "Материал", "value": "хлопок"}, {"label": "Цвет", "value": "белый"}]}}
        candidates = catalog_candidates_for_line(line, limit=2, intent={"product_class": "футболка"}, client=Client())

        self.assertEqual([value["external_id"] for value in candidates], ["cheap", "expensive"])

    def test_catalog_search_prioritizes_requirements_over_product_name_similarity(self):
        class Client:
            base_url = "https://api.oasiscatalog.com"

            def get(self, path, params=None):
                if path == "/v4/categories":
                    return [{"id": 10, "name": "Ручки", "path": "categories/office/pens"}]
                return [
                    {
                        "id": "popular", "article": "P", "group_id": "popular",
                        "name": "Ручка шариковая Popular", "full_name": "Ручка шариковая Popular, зеленая",
                        "materials": ["металл"], "colors": ["зеленый"], "attributes": [],
                        "categories": [10], "total_stock": 100, "price": "84",
                    },
                    {
                        "id": "gold", "article": "G", "group_id": "gold",
                        "name": "Ручка шариковая Euro Gold", "full_name": "Ручка шариковая Euro Gold, зеленая",
                        "materials": ["металл"], "colors": ["зеленый"],
                        "attributes": [{"name": "Чернила", "value": "синие"}, {"name": "Механизм", "value": "поворотный"}],
                        "categories": [10], "total_stock": 100, "price": "12.80",
                    },
                    {
                        "id": "chrome", "article": "C", "group_id": "chrome",
                        "name": "Ручка шариковая Euro Chrome", "full_name": "Ручка шариковая Euro Chrome, зеленая",
                        "materials": ["металл"], "colors": ["зеленый"],
                        "attributes": [{"name": "Чернила", "value": "синие"}, {"name": "Механизм", "value": "поворотный"}],
                        "categories": [10], "total_stock": 100, "price": "10.60",
                    },
                ]

        line = {"name": "Ручка, зелёная, материал – металл, чернила синие, механизм поворотный", "quantity": "10", "requirements": {"requirements": [
            {"label": "Материал", "value": "металл"},
            {"label": "Цвет", "value": "зелёная"},
            {"label": "Чернила", "value": "синие"},
            {"label": "Механизм", "value": "поворотный"},
            {"label": "Нанесение", "value": "гравировка 1+0"},
        ]}}

        candidates = catalog_candidates_for_line(line, limit=3, intent={"product_class": "ручка"}, client=Client())

        self.assertEqual([value["external_id"] for value in candidates], ["chrome", "gold", "popular"])
        self.assertTrue(any("Чернила" in value for value in candidates[0]["matches"]))
        self.assertTrue(any("Механизм" in value for value in candidates[0]["matches"]))

    def test_catalog_search_uses_llm_specific_categories_instead_of_generic_class(self):
        class Client:
            base_url = "https://api.oasiscatalog.com"

            def get(self, path, params=None):
                if path == "/v4/categories":
                    return [
                        {"id": 10, "name": "Футболки", "path": "categories/textile/tshirts"},
                        {"id": 11, "name": "Поло", "path": "categories/textile/polo"},
                    ]
                return [
                    {"id": "shirt", "article": "S", "group_id": "shirt", "name": "Футболка", "full_name": "Футболка белая", "colors": ["белый"], "categories": [10], "total_stock": 100, "price": "500"},
                    {"id": "polo", "article": "P", "group_id": "polo", "name": "Футболка поло", "full_name": "Футболка поло белая", "colors": ["белый"], "categories": [11], "total_stock": 100, "price": "600"},
                ]

        line = {"name": "Футболка поло унисекс", "quantity": "10", "requirements": {"requirements": [{"label": "Цвет", "value": "белый"}]}}
        intent = {"product_class": "футболка", "item": "поло", "categories": ["поло", "рубашка поло"], "synonyms": [], "required": [], "preferred": [], "secondary": []}

        candidates = catalog_candidates_for_line(line, limit=3, intent=intent, client=Client())

        self.assertEqual([value["external_id"] for value in candidates], ["polo"])

    def test_catalog_search_excludes_product_that_breaks_llm_required_constraint(self):
        class Client:
            base_url = "https://api.oasiscatalog.com"

            def get(self, path, params=None):
                if path == "/v4/categories":
                    return [{"id": 10, "name": "Футболки", "path": "categories/textile/tshirts"}]
                return [
                    {
                        "id": "cotton", "article": "C", "group_id": "cotton", "name": "Футболка",
                        "full_name": "Футболка хлопковая белая", "materials": ["хлопок"], "colors": ["белый"],
                        "categories": [10], "total_stock": 100, "price": "900",
                    },
                    {
                        "id": "cheap-polyester", "article": "P", "group_id": "polyester", "name": "Футболка",
                        "full_name": "Футболка из полиэстера белая", "materials": ["полиэстер"], "colors": ["белый"],
                        "categories": [10], "total_stock": 100, "price": "100",
                    },
                ]

        line = {"name": "Футболка", "quantity": 10, "requirements": {"requirements": []}}
        intent = {
            "categories": ["футболка"],
            "required": [{"label": "Состав", "value": "хлопок", "weight": 1}],
            "ranking": [{"criterion": "цена", "weight": 1}],
        }

        candidates = catalog_candidates_for_line(line, limit=3, intent=intent, client=Client())

        self.assertEqual([value["external_id"] for value in candidates], ["cotton"])

    def test_catalog_constraints_exclude_forbidden_value_and_read_fact_from_name_or_attribute(self):
        class Client:
            base_url = "https://api.oasiscatalog.com"

            def get(self, path, params=None):
                if path == "/v4/categories":
                    return [{"id": 10, "name": "Поло", "path": "categories/textile/polo"}]
                return [
                    {"id": "female-attribute", "article": "F1", "group_id": "f1", "name": "Поло Boston", "full_name": "Поло Boston белое", "attributes": [{"name": "Пол", "value": "женский"}], "colors": ["белый"], "categories": [10], "total_stock": 100, "price": 100},
                    {"id": "female-name", "article": "F2", "group_id": "f2", "name": "Поло Boston женское", "full_name": "Поло Boston женское, белое", "colors": ["белый"], "categories": [10], "total_stock": 100, "price": 90},
                    {"id": "male", "article": "M", "group_id": "m", "name": "Поло Laguna мужское", "full_name": "Поло Laguna мужское, белое", "attributes": [{"name": "Пол", "value": "мужской"}], "colors": ["белый"], "categories": [10], "total_stock": 100, "price": 200},
                    {"id": "unspecified", "article": "U", "group_id": "u", "name": "Поло Base", "full_name": "Поло Base, белое", "colors": ["белый"], "categories": [10], "total_stock": 100, "price": 80},
                ]

        intent = {
            "categories": ["поло"],
            "constraints": [
                {
                    "field": "gender", "operator": "in", "values": ["male", "unisex"],
                    "level": "required", "weight": 1, "missing_policy": "allow_with_penalty",
                },
                {
                    "field": "gender", "operator": "not_in", "values": ["female"],
                    "level": "required", "weight": 1, "missing_policy": "allow_with_penalty",
                },
            ],
        }

        candidates = catalog_candidates_for_line(
            {"name": "Поло унисекс", "quantity": 10, "requirements": {"requirements": []}},
            limit=10, intent=intent, client=Client(),
        )

        self.assertEqual([value["external_id"] for value in candidates], ["male", "unspecified"])
        self.assertTrue(any("Пол" in value for value in candidates[0]["matches"]))
        self.assertEqual(sum("Пол" in value for value in candidates[0]["matches"]), 1)
        self.assertTrue(any("Пол не указан" in value for value in candidates[1]["unknown"]))

    def test_catalog_constraints_apply_numeric_limits_before_ranking(self):
        class Client:
            base_url = "https://api.oasiscatalog.com"

            def get(self, path, params=None):
                if path == "/v4/categories":
                    return [{"id": 10, "name": "Поло", "path": "categories/textile/polo"}]
                return [
                    {"id": "valid", "article": "V", "group_id": "v", "name": "Поло", "full_name": "Поло белое", "attributes": [{"name": "Плотность", "value": "190 г/м²"}], "categories": [10], "total_stock": 150, "price": 450},
                    {"id": "expensive", "article": "E", "group_id": "e", "name": "Поло", "full_name": "Поло белое", "attributes": [{"name": "Плотность", "value": "190 г/м²"}], "categories": [10], "total_stock": 150, "price": 700},
                    {"id": "thin", "article": "T", "group_id": "t", "name": "Поло", "full_name": "Поло белое", "attributes": [{"name": "Плотность", "value": "150 г/м²"}], "categories": [10], "total_stock": 150, "price": 300},
                ]

        intent = {"categories": ["поло"], "constraints": [
            {"field": "price", "operator": "lte", "values": ["500"], "level": "required", "weight": 1, "missing_policy": "reject"},
            {"field": "density", "operator": "between", "values": ["180", "220"], "level": "required", "weight": 1, "missing_policy": "reject"},
            {"field": "stock", "operator": "gte", "values": ["100"], "level": "required", "weight": 1, "missing_policy": "reject"},
        ]}

        candidates = catalog_candidates_for_line({"name": "Поло", "quantity": 100}, limit=10, intent=intent, client=Client())

        self.assertEqual([value["external_id"] for value in candidates], ["valid"])

    def test_catalog_search_can_prioritize_price_over_preferred_material(self):
        class Client:
            base_url = "https://api.oasiscatalog.com"

            def get(self, path, params=None):
                if path == "/v4/categories":
                    return [{"id": 10, "name": "Футболки", "path": "categories/textile/tshirts"}]
                return [
                    {
                        "id": "cotton", "article": "C", "group_id": "cotton", "name": "Футболка",
                        "full_name": "Футболка хлопковая белая", "materials": ["хлопок"], "colors": ["белый"],
                        "categories": [10], "total_stock": 100, "price": "900",
                    },
                    {
                        "id": "cheap-polyester", "article": "P", "group_id": "polyester", "name": "Футболка",
                        "full_name": "Футболка из полиэстера белая", "materials": ["полиэстер"], "colors": ["белый"],
                        "categories": [10], "total_stock": 100, "price": "100",
                    },
                ]

        line = {"name": "Футболка", "quantity": 10, "requirements": {"requirements": []}}
        intent = {
            "categories": ["футболка"],
            "preferred": [{"label": "Материал", "value": "хлопок", "weight": .2}],
            "ranking": [{"criterion": "цена", "weight": 1}],
        }

        candidates = catalog_candidates_for_line(line, limit=2, intent=intent, client=Client())

        self.assertEqual([value["external_id"] for value in candidates], ["cheap-polyester", "cotton"])

    def test_catalog_search_returns_source_diagnostics_instead_of_hiding_oasis_failure(self):
        gifts = CatalogSupplier.objects.create(code="gifts", name="gifts.ru", base_url="https://api2.gifts.ru/export/v2")
        CatalogProduct.objects.create(
            supplier=gifts, external_id="shirt", article="G", name="Футболка", full_name="Футболка белая",
            colors=["белый"], total_stock=100, discount_price=500, search_text="футболка белая",
        )

        class Client:
            base_url = "https://api.oasiscatalog.com"

            def get(self, path, params=None):
                raise CatalogSyncError("Oasis временно недоступен")

        result = catalog_candidates_for_line(
            {"name": "Футболка", "quantity": 10},
            limit=3,
            intent={"categories": ["футболка"]},
            client=Client(),
            include_diagnostics=True,
        )

        self.assertEqual(result["candidates"][0]["supplier_code"], "gifts")
        self.assertEqual(result["sources"]["oasis"]["status"], "failed")
        self.assertIn("временно недоступен", result["sources"]["oasis"]["message"])
        self.assertEqual(result["sources"]["gifts"]["status"], "success")

    def test_gifts_retrieval_uses_product_entity_before_broad_characteristics(self):
        gifts = CatalogSupplier.objects.create(code="gifts", name="gifts.ru", base_url="https://api2.gifts.ru/export/v2")
        CatalogProduct.objects.bulk_create([
            CatalogProduct(
                supplier=gifts, external_id=f"mug-{index}", article=f"M-{index}", name="Кружка синяя",
                full_name="Кружка синяя", colors=["синий"], total_stock=100, discount_price=100,
                search_text="кружка синяя",
            )
            for index in range(1500)
        ])
        CatalogProduct.objects.create(
            supplier=gifts, external_id="shirt", article="S", name="Футболка синяя",
            full_name="Футболка синяя", colors=["синий"], total_stock=100, discount_price=500,
            search_text="футболка синяя",
        )

        class Client:
            base_url = "https://api.oasiscatalog.com"

            def get(self, path, params=None):
                raise CatalogSyncError("Oasis временно недоступен")

        candidates = catalog_candidates_for_line(
            {"name": "Футболка", "quantity": 10},
            limit=3,
            intent={
                "categories": ["футболка"],
                "required": [{"label": "Цвет", "value": "синий", "weight": 1}],
            },
            client=Client(),
        )

        self.assertEqual([value["external_id"] for value in candidates], ["shirt"])

    def test_catalog_entity_match_does_not_treat_polo_as_towel_prefix(self):
        class Client:
            base_url = "https://api.oasiscatalog.com"

            def get(self, path, params=None):
                if path == "/v4/categories":
                    return [{"id": 10, "name": "Текстиль", "path": "categories/textile"}]
                return [{
                    "id": "towel", "article": "T", "group_id": "towel", "name": "Полотенце",
                    "full_name": "Полотенце синее", "colors": ["синий"], "categories": [10],
                    "total_stock": 100, "price": 300,
                }]

        candidates = catalog_candidates_for_line(
            {"name": "Поло", "quantity": 10},
            limit=3,
            intent={"categories": ["поло"]},
            client=Client(),
        )

        self.assertEqual(candidates, [])

    def test_catalog_search_does_not_prefer_supplier_when_relevance_and_price_are_equal(self):
        gifts = CatalogSupplier.objects.create(code="gifts", name="gifts.ru", base_url="https://api2.gifts.ru/export/v2")
        CatalogProduct.objects.create(
            supplier=gifts, external_id="gifts-shirt", article="G", name="Футболка", full_name="Футболка Б",
            materials=["хлопок"], colors=["белый"], discount_price=500, total_stock=100,
            search_text="футболка б хлопок белый",
        )

        class Client:
            base_url = "https://api.oasiscatalog.com"

            def get(self, path, params=None):
                if path == "/v4/categories":
                    return [{"id": 10, "name": "Футболки", "path": "categories/tekstil/futbolki"}]
                return [{"id": "oasis-shirt", "article": "O", "group_id": "oasis-shirt", "name": "Футболка", "full_name": "Футболка А", "materials": ["хлопок"], "colors": ["белый"], "price": "500", "categories": [10], "total_stock": 100}]

        line = {"name": "Футболка", "quantity": "10", "requirements": {"requirements": [{"label": "Материал", "value": "хлопок"}, {"label": "Цвет", "value": "белый"}]}}
        candidates = catalog_candidates_for_line(line, limit=2, intent={"product_class": "футболка"}, client=Client())

        self.assertEqual([value["supplier_code"] for value in candidates], ["oasis", "gifts"])

    def test_catalog_search_uses_gifts_text_when_oasis_category_is_missing(self):
        gifts = CatalogSupplier.objects.create(code="gifts", name="gifts.ru", base_url="https://api2.gifts.ru/export/v2")
        CatalogProduct.objects.create(
            supplier=gifts, external_id="longsleeve-white", article="LS-1",
            name="Лонгслив унисекс", full_name="Лонгслив унисекс, белый",
            colors=["белый"], total_stock=100, discount_price=500,
            search_text="лонгслив унисекс белый хлопок футболка с длинным рукавом",
        )

        class Client:
            base_url = "https://api.oasiscatalog.com"

            def get(self, path, params=None):
                if path == "/v4/categories":
                    return [{"id": 10, "name": "Брелоки", "path": "categories/accessories"}]
                return []

        line = {"name": "Лонгслив, унисекс", "quantity": "10", "requirements": {"requirements": [{"label": "Цвет", "value": "белый"}]}}
        candidates = catalog_candidates_for_line(
            line,
            limit=3,
            intent={"item": "лонгслив", "categories": ["лонгслив"], "synonyms": ["long sleeve"]},
            client=Client(),
        )

        self.assertEqual(candidates[0]["external_id"], "longsleeve-white")
        self.assertEqual(candidates[0]["supplier_code"], "gifts")

    def test_catalog_search_keeps_gifts_when_oasis_is_unavailable(self):
        gifts = CatalogSupplier.objects.create(code="gifts", name="gifts.ru", base_url="https://api2.gifts.ru/export/v2")
        CatalogProduct.objects.create(
            supplier=gifts, external_id="shirt", article="S-1", name="Футболка",
            full_name="Футболка белая", colors=["белый"], total_stock=100,
            discount_price=500, search_text="футболка белая хлопок",
        )

        class Client:
            base_url = "https://api.oasiscatalog.com"

            def get(self, path, params=None):
                raise CatalogSyncError("Oasis недоступен")

        candidates = catalog_candidates_for_line(
            {"name": "Футболка", "quantity": "10", "requirements": {"requirements": [{"label": "Цвет", "value": "белый"}]}},
            limit=3, intent={"product_class": "футболка"}, client=Client(),
        )

        self.assertEqual(candidates[0]["external_id"], "shirt")
        self.assertEqual(candidates[0]["supplier_code"], "gifts")

    def test_catalog_search_excludes_zero_stock_and_ranks_shortage_after_available(self):
        class Client:
            base_url = "https://api.oasiscatalog.com"

            def get(self, path, params=None):
                if path == "/v4/categories":
                    return [{"id": 10, "name": "Футболки", "path": "categories/tekstil/futbolki"}]
                return [
                    {"id": "available", "article": "A", "group_id": "available", "name": "Футболка", "full_name": "Футболка белая", "colors": ["белый"], "materials": ["хлопок"], "categories": [10], "total_stock": 100, "price": 900},
                    {"id": "shortage", "article": "S", "group_id": "shortage", "name": "Футболка", "full_name": "Футболка белая", "colors": ["белый"], "materials": ["полиэстер"], "categories": [10], "total_stock": 5, "price": 100},
                    {"id": "empty", "article": "E", "group_id": "empty", "name": "Футболка", "full_name": "Футболка белая", "colors": ["белый"], "materials": ["хлопок"], "categories": [10], "total_stock": 0, "price": 1},
                ]

        line = {"name": "Футболка", "quantity": "10", "requirements": {"requirements": [{"label": "Материал", "value": "хлопок"}, {"label": "Цвет", "value": "белый"}]}}
        candidates = catalog_candidates_for_line(line, limit=3, intent={"product_class": "футболка"}, client=Client())

        self.assertEqual([value["external_id"] for value in candidates], ["available", "shortage"])
        self.assertTrue(any("Недостаточный остаток" in value for value in candidates[1]["mismatches"]))

    def test_catalog_search_does_not_call_a_shirt_with_long_sleeves_a_longsleeve(self):
        class Client:
            base_url = "https://api.oasiscatalog.com"

            def get(self, path, params=None):
                if path == "/v4/categories":
                    return [{"id": 10, "name": "Одежда", "path": "categories/odezhda"}]
                return [{
                    "id": "shirt", "article": "SH-1", "group_id": "shirt",
                    "name": "Рубашка женская", "full_name": "Рубашка женская с длинным рукавом",
                    "description": "Футболка с длинным рукавом в описании модели", "categories": [10], "total_stock": 100,
                    "price": 900,
                }]

        candidates = catalog_candidates_for_line(
            {"name": "Лонгслив", "quantity": "10"},
            limit=3, intent={"product_class": "лонгслив"}, client=Client(),
        )

        self.assertEqual(candidates, [])

    def test_selected_catalog_product_is_recalculated_on_backend_and_replaces_material_cost(self):
        production_type = ProductionType.objects.create(code="catalog-product", name="Каталожный товар")
        line = {
            "name": "Футболка поло", "quantity": "300",
            "requirements": {"requirements": [
                {"label": "Материал", "value": "хлопок"},
                {"label": "Цвет", "value": "темно-синий"},
                {"label": "Плотность", "value": "не менее 190 г/м²"},
                {"label": "Нанесение", "value": "вышивка"},
            ]},
        }
        hypothesis = {
            "product_type": production_type.code, "confidence": .8, "matched_example_ids": [1],
            "route": {"reason": "Нужен готовый товар и нанесение", "processes": [{"name": "Закупка материала"}, {"name": "Нанесение"}]},
            "questions": ["Какова цена закупки готовой футболки?", "Какова цена нанесения?"],
            "costs": [{
                "category": "material", "name": "Старая ручная цена", "amount_total": "99999",
                "source": "Введено администратором", "source_type": "manager", "process_name": "Закупка материала",
            }],
            "catalog_candidates": [{"id": "exact", "external_id": "exact", "supplier_code": "oasis", "article": "POLO-190", "name": "Футболка поло тёмно-синяя", "price": "650.00", "stock": 500, "url": "https://www.oasiscatalog.com/item/exact", "fit": "exact", "matches": ["Тип товара: поло"], "mismatches": [], "unknown": []}],
        }

        result = apply_catalog_candidate(hypothesis, line, "exact")

        self.assertEqual(result["catalog_selection"]["id"], "exact")
        self.assertEqual(result["product_type"], production_type.code)
        self.assertIn({"code": production_type.code, "name": production_type.name}, result["production_types"])
        self.assertEqual(result["totals"]["material_unit"], "650.00")
        self.assertEqual(result["totals"]["cost_total"], "195000.00")
        self.assertEqual(result["costs"][0]["source_type"], "catalog")
        self.assertEqual(result["costs"][0]["calculation_steps"][-1], "300 шт. × 650.00 ₽/шт. = 195000.00 ₽")
        self.assertEqual(result["route"]["steps"][0], "Закупка готового изделия")
        self.assertEqual(result["questions"], ["Какова цена нанесения?"])
        self.assertIn("поставщика Oasis", result["route"]["reason"])
        self.assertEqual(result["sources"][-1]["supplier_name"], "Oasis")
        self.assertEqual(result["sources"][-1]["price"], "650.00")

    def test_admin_can_use_partial_catalog_candidate_with_visible_mismatches(self):
        production_type = ProductionType.objects.create(code="catalog-partial", name="Каталожный товар")
        line = {"name": "Футболка", "quantity": "10", "requirements": {"requirements": []}}
        hypothesis = {
            "product_type": production_type.code, "confidence": .5,
            "route": {"reason": "Каталог", "processes": [{"name": "Закупка готового изделия"}]},
            "costs": [], "catalog_candidates": [{
                "id": "partial", "external_id": "partial", "supplier_code": "other",
                "supplier_name": "Другой поставщик", "supplier_site": "catalog.example",
                "article": "PART-1", "name": "Футболка зелёная", "price": "400.00",
                "stock": 100, "url": "https://www.oasiscatalog.com/item/partial",
                "fit": "partial", "matches": ["Тип товара: футболка"],
                "mismatches": ["Плотность ниже требования"], "unknown": [],
            }],
        }

        result = apply_catalog_candidate(hypothesis, line, "partial")

        self.assertEqual(result["totals"]["material_unit"], "400.00")
        self.assertEqual(result["totals"]["cost_total"], "4000.00")
        self.assertEqual(result["catalog_selection"]["selection_mode"], "manual")
        self.assertEqual(result["catalog_selection"]["accepted_mismatches"], ["Плотность ниже требования"])
        self.assertEqual(result["sources"][-1]["supplier_name"], "Другой поставщик")
        self.assertIn("поставщика Другой поставщик", result["route"]["reason"])

    def test_catalog_choice_becomes_training_data_only_after_confirmation(self):
        admin = get_user_model().objects.create_superuser(username="catalog-admin", password="password")
        production_type = ProductionType.objects.create(code="catalog-training", name="Каталожный товар")
        line = {"name": "Футболка", "quantity": "10", "requirements": {"requirements": [
            {"label": "Материал", "value": "хлопок"}, {"label": "Цвет", "value": "белый"},
        ]}}
        session = ProductionTrainingSession.objects.create(
            created_by=admin, position_name="Футболка", requirements=line["requirements"],
            current_hypothesis={
                "stage": "training_dialogue", "product_type": production_type.code, "confidence": .5,
                "route": {"name": "Закупка", "reason": "Каталог", "steps": ["Закупка готового изделия"], "processes": [{"name": "Закупка готового изделия"}]},
                "costs": [], "matched_example_ids": [],
                "catalog_candidates": [{"id": "shirt", "external_id": "shirt", "supplier_code": "oasis", "article": "SHIRT-1", "name": "Футболка белая", "price": "500.00", "stock": 100, "url": "https://www.oasiscatalog.com/item/shirt", "fit": "exact", "matches": ["Тип товара: футболка"], "mismatches": [], "unknown": []}],
            },
        )
        self.client.force_login(admin)

        selected = self.client.post(reverse("tender_select_catalog_product"), {
            "payload": json.dumps({"session_id": session.pk, "line": line, "product_id": "shirt"}),
        })

        self.assertEqual(selected.status_code, 200)
        decision = CatalogMatchDecision.objects.get()
        self.assertFalse(decision.is_confirmed)
        self.assertIsNone(decision.product)
        self.assertEqual(decision.product_external_id, "shirt")
        self.assertEqual(selected.json()["totals"]["cost_total"], "5000.00")

        confirmed = self.client.post(reverse("tender_confirm_production_type"), {
            "payload": json.dumps({"session_id": session.pk, "line": line}),
        })

        self.assertEqual(confirmed.status_code, 200)
        decision.refresh_from_db()
        self.assertTrue(decision.is_confirmed)
        example = ProductionTrainingExample.objects.get(pk=confirmed.json()["example_id"])
        self.assertEqual(example.routes[0]["catalog_selection"]["id"], "shirt")

    def test_route_name_contains_only_universal_processes(self):
        production_type = ProductionType.objects.create(code="route-test", name="Тест маршрута")
        raw = {
            "product_type": production_type.code,
            "route": {"reason": "Тираж и отделка", "processes": [
                {"name": "Поставка дизайнерской бумаги Majestic для папки", "details": ["Majestic"]},
                {"name": "Изготовление папок с резкой, биговкой и тиснением в универсальной типографии", "details": ["резка", "биговка"]},
            ]},
            "costs": [],
        }

        result = _normalize_training_hypothesis(raw, {"quantity": 100}, [production_type], [1])

        self.assertEqual(result["route"]["name"], "Закупка материала → Универсальная типография")
        self.assertEqual(result["route"]["processes"][1]["details"], ["резка", "биговка"])

    def test_turnkey_manufacturing_is_not_labelled_as_material_purchase(self):
        production_type = ProductionType.objects.create(code="turnkey-test", name="Тест под ключ")
        raw = {
            "product_type": production_type.code,
            "route": {
                "reason": "Заказать изготовление под ключ в цифровой типографии: бумага, печать 4+4, выборочный УФ-лак и резка.",
                "processes": [{"name": "Закупка материала", "details": ["Типография предоставляет бумагу"]}],
            },
            "costs": [],
        }

        result = _normalize_training_hypothesis(raw, {"quantity": 700}, [production_type], [])

        self.assertEqual(result["route"]["name"], "Цифровая типография под ключ")
        self.assertTrue(result["route"]["is_turnkey"])

    def test_separate_material_purchase_remains_a_separate_route_process(self):
        production_type = ProductionType.objects.create(code="split-route-test", name="Раздельный маршрут")
        raw = {
            "product_type": production_type.code,
            "route": {
                "reason": "Бумагу покупаем сами и передаём типографии.",
                "processes": [
                    {"name": "Закупка материала", "details": ["Бумага Majestic"]},
                    {"name": "Цифровая типография под ключ", "details": ["Печать и отделка"]},
                ],
            },
            "costs": [],
        }

        result = _normalize_training_hypothesis(raw, {"quantity": 700}, [production_type], [])

        self.assertEqual(result["route"]["name"], "Закупка материала → Цифровая типография под ключ")
        self.assertFalse(result["route"]["is_turnkey"])

    def test_psodin_work_is_calculated_by_existing_backend_calculator(self):
        CalculatorSettings.objects.update_or_create(pk=1, defaults={
            "hourly_rate": Decimal("1000"), "time_coefficient": Decimal("1.5"), "partner_discount": Decimal("15"),
        })
        hypothesis = {
            "costs": [
                {"category": "material", "name": "Готовая DTF-плёнка", "amount_total": "2000", "source_type": "supplier"},
                {"category": "application", "name": "Выдуманная работа", "amount_total": "9999", "source": "Калькулятор PSODIN", "source_type": "calculator"},
            ],
            "questions": ["Уточнить часы PSODIN"],
        }
        raw = {"psodin_calculation": {"process_name": "Нанесение DTF в PSODIN", "productivity_per_hour": 10, "tariff": "partner"}}

        result = _apply_psodin_calculation(hypothesis, raw, {"quantity": 50}, feedback="Работу делаем в PSODIN")

        psodin_cost = next(item for item in result["costs"] if item["source_type"] == "calculator")
        self.assertEqual(psodin_cost["amount_total"], "6375.00")
        self.assertEqual(result["psodin_calculation"]["exact_hours"], "5")
        self.assertEqual(result["psodin_calculation"]["billed_hours"], "5")
        self.assertEqual(result["totals"]["cost_total"], "8375.00")
        self.assertIn("Скидка 15.00% к стоимости", psodin_cost["calculation_steps"][-1])

    def test_psodin_backend_requests_productivity_instead_of_inventing_price(self):
        hypothesis = {"costs": [], "questions": []}

        result = _apply_psodin_calculation(hypothesis, {"psodin_calculation": {"requested": True}}, {"quantity": 50}, feedback="Считаем работу в PSODIN")

        self.assertEqual(result["psodin_calculation"]["status"], "missing_productivity")
        self.assertEqual(result["costs"], [])
        self.assertIn("Сколько изделий в час", result["questions"][0])

    def test_confirmed_psodin_productivity_is_reused_without_new_feedback(self):
        hypothesis = {"costs": [], "questions": []}
        confirmed = {"authorized": True, "productivity_per_hour": "10", "tariff": "partner"}

        result = _apply_psodin_calculation(hypothesis, {}, {"quantity": 50}, confirmed=confirmed)

        self.assertEqual(result["psodin_calculation"]["status"], "calculated")
        self.assertEqual(result["psodin_calculation"]["exact_hours"], "5")

    def test_manual_logistics_is_not_labelled_as_tz_source(self):
        production_type = ProductionType.objects.create(code="source-test", name="Тест источника")
        raw = {
            "product_type": production_type.code,
            "route": {"steps": ["Универсальная типография"]},
            "costs": [{"category": "logistics", "name": "Логистика", "amount_total": 3000, "source": "Дано в ТЗ", "source_type": "supplier"}],
        }

        result = _normalize_training_hypothesis(raw, {"quantity": 1000, "logistics_unit": 3}, [production_type], [1])

        self.assertEqual(result["costs"][0]["source"], "Введено администратором в расчёте")
        self.assertEqual(result["costs"][0]["source_type"], "manager")

    def test_manual_fixed_logistics_gets_backend_recipe_without_warning(self):
        production_type = ProductionType.objects.create(code="manual-logistics", name="Ручная логистика")
        raw = {
            "product_type": production_type.code,
            "route": {"steps": ["Закупка готового изделия", "Нанесение"]},
            "costs": [{
                "category": "logistics",
                "name": "Межцеховая доставка",
                "amount_total": 1000,
                "source": "Введено администратором",
                "source_type": "manager",
                "recipe": {"method": "none", "inputs": {}},
            }],
        }

        result = _normalize_training_hypothesis(raw, {"quantity": 160}, [production_type], [])

        logistics = result["costs"][0]
        self.assertEqual(logistics["amount_total"], "1000.00")
        self.assertEqual(logistics["recipe"], {"method": "fixed", "inputs": {"fixed_amount": "1000.00"}})
        self.assertEqual(logistics["calculation_steps"], ["Фиксированная стоимость на тираж: 1000.00 ₽"])
        self.assertNotIn("нет проверяемой серверной формулы", " ".join(result["learning_warnings"]))

    def test_html_price_table_keeps_rows_and_backend_selects_exact_tier(self):
        html = """
            <div>КВАДРАТНЫЕ 70х70 см полиэфирный шелк, горячий рез</div>
            <div>Основная цена 325 ₽</div><div>от 101 шт. — 250 ₽</div>
            <table>
              <tr><th>Косынки, банданы ТРЕУГОЛЬНЫЕ</th><th>до 10 шт</th><th>11-20 шт</th><th>21-50 шт</th><th>51-100 шт</th><th>101-200 шт</th><th>201-500 шт</th><th>от 500 шт</th></tr>
              <tr><td>70х70х100 см полиэфирный шелк, горячий рез</td><td>225 ₽</td><td>215 ₽</td><td>210 ₽</td><td>200 ₽</td><td>185 ₽</td><td>165 ₽</td><td>155 ₽</td></tr>
              <tr><td>70х70х100 см армани/мокрый шелк, оверлок</td><td>380 ₽</td><td>365 ₽</td><td>360 ₽</td><td>345 ₽</td><td>325 ₽</td><td>285 ₽</td><td>275 ₽</td></tr>
              <tr><th>Цены на косынки КВАДРАТНЫЕ</th><th>до 10 шт</th><th>11-20 шт</th><th>21-50 шт</th><th>51-100 шт</th><th>101-200 шт</th><th>201-500 шт</th><th>от 500 шт</th></tr>
              <tr><td>70х70 см армани/мокрый шелк, оверлок</td><td>635 ₽</td><td>620 ₽</td><td>615 ₽</td><td>575 ₽</td><td>545 ₽</td><td>474 ₽</td><td>465 ₽</td></tr>
            </table>
        """
        parser = _VisibleTextParser()
        parser.feed(html)

        quote = _select_html_price_quote(parser.tables, {
            "line": {"name": "Платок", "quantity": 600, "requirements": {}},
            "feedback": "Найди платок армани, оверлок, 70х70х100, тираж 600 штук",
        })

        self.assertIn("СТРОКА 3: 70х70х100 см армани/мокрый шелк, оверлок", _format_html_tables(parser.tables))
        self.assertEqual(quote["row_label"], "70х70х100 см армани/мокрый шелк, оверлок")
        self.assertEqual(quote["tier"], "от 500 шт")
        self.assertEqual(quote["unit_price"], "275.00")
        self.assertEqual(quote["amount_total"], "165000.00")
        self.assertEqual(quote["confidence"], "exact")

    def test_html_price_table_refuses_ambiguous_product_row(self):
        html = """
            <table>
              <tr><th>Товар</th><th>до 100 шт</th><th>от 101 шт</th></tr>
              <tr><td>Платок армани красный</td><td>400 ₽</td><td>300 ₽</td></tr>
              <tr><td>Платок армани синий</td><td>410 ₽</td><td>310 ₽</td></tr>
            </table>
        """
        parser = _VisibleTextParser()
        parser.feed(html)

        quote = _select_html_price_quote(parser.tables, {
            "line": {"name": "Платок армани", "quantity": 600, "requirements": {}},
            "feedback": "Найди цену платка армани",
        })

        self.assertIsNone(quote)

    def test_html_price_table_respects_rowspan_and_colspan_headers(self):
        html = """
            <table>
              <tr><th rowspan="2">Товар</th><th colspan="2">Цена по тиражу</th></tr>
              <tr><th>до 100 шт</th><th>от 101 шт</th></tr>
              <tr><td>Платок 70х70 армани</td><td>400 ₽</td><td>300 ₽</td></tr>
            </table>
        """
        parser = _VisibleTextParser()
        parser.feed(html)

        quote = _select_html_price_quote(parser.tables, {
            "line": {"name": "Платок 70х70 армани", "quantity": 600, "requirements": {}},
        })

        self.assertEqual(parser.tables[0][1], ["Товар", "до 100 шт", "от 101 шт"])
        self.assertEqual(quote["unit_price"], "300.00")

    def test_backend_verified_web_quote_overrides_llm_price(self):
        production_type = ProductionType.objects.create(code="web-price", name="Цена с сайта")
        line = {"name": "Платок", "quantity": 600, "requirements": {}}
        hypothesis = {
            "product_type": production_type.code,
            "confidence": .5,
            "route": {"processes": [{"name": "Изготовление под ключ"}], "reason": "Поставщик"},
            "costs": [{
                "category": "material", "process_name": "Изготовление под ключ", "name": "Платок армани",
                "amount_total": "150000", "source": "Pro-flag", "source_type": "supplier",
                "source_url": "https://pro-flag.ru/price", "recipe": {"method": "unit_rate", "inputs": {"unit_rate": "250"}},
            }],
            "questions": ["Какова точная цена платка?"],
        }
        quote = {
            "confidence": "exact", "method": "html_table_tier",
            "row_label": "70х70х100 см армани/мокрый шелк, оверлок",
            "tier": "от 500 шт", "quantity": "600", "unit_price": "275.00", "amount_total": "165000.00",
        }

        result = apply_verified_source_quote(
            hypothesis, line, quote, "Pro-flag · платки", "https://pro-flag.ru/price"
        )

        self.assertEqual(result["costs"][0]["recipe"], {"method": "unit_rate", "inputs": {"unit_rate": "275.00"}})
        self.assertEqual(result["costs"][0]["amount_total"], "165000.00")
        self.assertEqual(result["totals"]["cost_unit"], "275.00")
        self.assertEqual(result["questions"], [])
        self.assertEqual(result["verified_source_quote"]["tier"], "от 500 шт")

    def test_private_url_cannot_be_used_as_calculation_source(self):
        with self.assertRaisesMessage(Exception, "Локальные и служебные адреса"):
            _validate_public_url("http://127.0.0.1/price")

    @patch("tenders.views.build_training_hypothesis")
    @patch("tenders.views.extract_calculation_source")
    def test_admin_can_attach_source_to_specific_cost(self, extract, rebuild):
        self.user.is_superuser = True
        self.user.is_staff = True
        self.user.save(update_fields=["is_superuser", "is_staff"])
        production_type = ProductionType.objects.get(code="digital_sheet")
        current = {"stage": "training_dialogue", "product_type": production_type.code, "route": {"name": "Закупка материала", "steps": ["Закупка материала"]}, "costs": [{"name": "Бумага", "amount_total": "1000"}], "totals": {}}
        session = ProductionTrainingSession.objects.create(created_by=self.user, position_name="Папка", requirements={}, current_hypothesis=current)
        extract.return_value = {
            "content": "Majestic SRA3 — 380 руб./лист", "source_type": "link", "url": "https://supplier.example/price",
            "structured_data": {"price_quote": {
                "confidence": "exact", "method": "html_table_tier", "row_label": "Majestic SRA3",
                "tier": "от 1 шт", "quantity": "100", "unit_price": "380.00", "amount_total": "38000.00",
            }},
        }
        rebuilt = {**current, "costs": [{"name": "Бумага", "amount_total": "9500", "calculation_steps": ["25 листов × 380 ₽"]}], "understood_changes": ["Цена бумаги пересчитана"]}
        rebuild.return_value = rebuilt
        self.client.force_login(self.user)

        response = self.client.post(reverse("tender_add_calculation_source"), {
            "payload": json.dumps({"session_id": session.pk, "line": {"name": "Папка", "quantity": 100}, "cost_index": 0, "supplier_name": "Дубль В", "source_url": "https://supplier.example/price"})
        })

        self.assertEqual(response.status_code, 200)
        source = TenderKnowledgeSource.objects.get()
        self.assertEqual(source.supplier_name, "Дубль В")
        self.assertFalse(source.is_active)
        self.assertEqual(response.json()["costs"][0]["source_id"], source.pk)
        self.assertEqual(response.json()["costs"][0]["amount_total"], "38000.00")
        self.assertEqual(response.json()["costs"][0]["source_type"], "supplier")
        self.assertTrue(response.json()["sources"][0]["is_pending"])
        self.assertIn("Majestic SRA3", rebuild.call_args.kwargs["feedback"])

        confirm_response = self.client.post(reverse("tender_confirm_production_type"), {
            "payload": json.dumps({"session_id": session.pk, "line": {"name": "Папка", "quantity": 100}}),
        })

        self.assertEqual(confirm_response.status_code, 200)
        source.refresh_from_db()
        self.assertTrue(source.is_active)

    @patch("tenders.views.build_training_hypothesis")
    @patch("tenders.views.extract_calculation_source")
    def test_admin_can_attach_source_before_any_cost_exists(self, extract, rebuild):
        self.user.is_superuser = True
        self.user.is_staff = True
        self.user.save(update_fields=["is_superuser", "is_staff"])
        production_type = ProductionType.objects.get(code="digital_sheet")
        current = {
            "stage": "training_dialogue",
            "product_type": production_type.code,
            "route": {"name": "Цифровая типография под ключ", "steps": ["Цифровая типография под ключ"]},
            "costs": [],
            "totals": {},
        }
        session = ProductionTrainingSession.objects.create(
            created_by=self.user, position_name="Визитки", requirements={}, current_hypothesis=current,
        )
        extract.return_value = {
            "content": "Sirio Pearl SRA3 — 380 руб./лист",
            "source_type": "image",
            "url": "https://bereg.example/paper",
        }
        rebuild.return_value = {
            **current,
            "costs": [{"name": "Бумага Sirio Pearl", "amount_total": "9500"}],
            "understood_changes": ["Добавлен раздельный маршрут"],
        }
        self.client.force_login(self.user)

        response = self.client.post(reverse("tender_add_calculation_source"), {
            "payload": json.dumps({
                "session_id": session.pk,
                "line": {"name": "Визитки с выборочным УФ-лаком", "quantity": 700},
                "supplier_name": "Берег",
                "source_url": "https://bereg.example/paper",
                "feedback": "Рассмотри закупку бумаги отдельным маршрутом.",
            })
        })

        self.assertEqual(response.status_code, 200)
        source = TenderKnowledgeSource.objects.get()
        self.assertEqual(source.structured_data["scope"], "position")
        self.assertFalse(source.is_active)
        self.assertEqual(response.json()["sources"][0]["supplier_name"], "Берег")
        feedback = rebuild.call_args.kwargs["feedback"]
        self.assertIn("Рассмотри закупку бумаги отдельным маршрутом", feedback)
        self.assertIn("Sirio Pearl", feedback)
        self.assertNotIn("cost_index", json.loads(response.wsgi_request.POST["payload"]))

    @patch("tenders.views.build_training_hypothesis")
    @patch("tenders.views.extract_calculation_source")
    def test_admin_can_attach_multiple_sources_in_one_recalculation(self, extract, rebuild):
        self.user.is_superuser = True
        self.user.is_staff = True
        self.user.save(update_fields=["is_superuser", "is_staff"])
        production_type = ProductionType.objects.get(code="digital_sheet")
        current = {
            "stage": "training_dialogue",
            "product_type": production_type.code,
            "route": {"name": "Комбинированный маршрут", "steps": ["Изготовление шнурков", "Закупка вкладышей"]},
            "costs": [],
            "totals": {},
        }
        session = ProductionTrainingSession.objects.create(
            created_by=self.user, position_name="Шнурок с вкладышем", requirements={}, current_hypothesis=current,
        )
        extract.side_effect = [
            {"content": "Изготовление шнурков 500 шт. по 80 руб.", "source_type": "link", "url": "https://lanyard.example/"},
            {"content": "Готовые вкладыши 500 шт. по 12 руб.", "source_type": "link", "url": "https://insert.example/"},
        ]
        rebuild.return_value = {**current, "sources": [], "understood_changes": ["Маршрут разделён на два процесса"]}
        self.client.force_login(self.user)

        response = self.client.post(reverse("tender_add_calculation_source"), {
            "payload": json.dumps({
                "session_id": session.pk,
                "line": {"name": "Шнурок для телефона с вкладышем", "quantity": 500},
                "feedback": "Шнурки изготавливаем на заказ, вкладыши закупаем готовыми.",
                "sources": [
                    {"supplier_name": "Шнурки", "title": "Изготовление шнурков", "source_url": "https://lanyard.example/"},
                    {"supplier_name": "Вкладыши", "title": "Готовые вкладыши", "source_url": "https://insert.example/"},
                ],
            }),
        })

        self.assertEqual(response.status_code, 200)
        self.assertEqual(TenderKnowledgeSource.objects.count(), 2)
        self.assertEqual(len(response.json()["sources"]), 2)
        self.assertTrue(all(value["is_pending"] for value in response.json()["sources"]))
        rebuild.assert_called_once()
        feedback = rebuild.call_args.kwargs["feedback"]
        self.assertIn("ИСТОЧНИК № 1", feedback)
        self.assertIn("ИСТОЧНИК № 2", feedback)
        self.assertIn("Шнурки изготавливаем на заказ", feedback)
        self.assertIn("Готовые вкладыши", feedback)

    def test_only_relevant_knowledge_sources_are_selected_for_future_calculations(self):
        self.user.is_superuser = True
        self.user.save(update_fields=["is_superuser"])
        TenderKnowledgeSource.objects.create(
            title="Sirio Pearl и Majestic", supplier_name="Берег", source_type="link",
            url="https://bereg.example/designer-paper", content_summary="Дизайнерская перламутровая бумага Sirio Pearl SRA3 300 г/м²",
            created_by=self.user,
        )
        TenderKnowledgeSource.objects.create(
            title="Хлопковая ткань", supplier_name="Текстиль", source_type="text",
            content_summary="Ткань для пошива футболок", created_by=self.user,
        )

        sources = _knowledge_sources_for_line({
            "name": "Визитки на Sirio Pearl",
            "requirements": {"requirements": [{"label": "Бумага", "value": "перламутровая дизайнерская Sirio Pearl 300 г/м²"}]},
        })

        self.assertEqual([value["supplier"] for value in sources], ["Берег"])

    @patch("tenders.services._ai_gateway_json")
    def test_known_embossing_and_spring_are_recovered_from_catalog(self, gateway):
        setup = PriceItem.objects.create(category="embossing", name="Приладка (количество клише)", aliases="горячее тиснение", unit_name="клише", unit_price="1000")
        hits = PriceItem.objects.create(category="embossing", name="Ударов", aliases="тиснение фольгой", unit_name="удар", unit_price="4")
        spring = PriceItem.objects.create(category="postpress", name="Пружина в бобине, 30см", aliases="металлическая пружина", unit_price="3")
        gateway.return_value = ({"selected_source": "psodin_sheet", "calculator": "sheet", "route": "internal", "confidence": .9, "reason": "Нужна проверка", "components": [{"name": "Блок", "source": "internal", "kind": "sheet", "finished_width_mm": 148, "finished_height_mm": 210, "units_per_product": 60, "material_query": "офсетная бумага", "grammage_gsm": 80, "operation_item_ids": []}], "questions": [], "warnings": []}, {})

        result = analyze_production_route({"name": "Блокнот на металлической пружине с горячим тиснением", "quantity": 300, "requirements": {}})

        embossing = [value for value in result["cost_options"] if value["group"] == "embossing"]
        spring_options = [value for value in result["cost_options"] if value["group"] == "spring"]
        self.assertTrue({setup.pk, hits.pk}.issubset({value["catalog_item_id"] for value in embossing}))
        self.assertEqual(spring_options[0]["catalog_item_id"], spring.pk)
        self.assertIn("2 шт. с 30 см", spring_options[0]["calculation"])

    @patch("tenders.services._ai_gateway_json")
    def test_outsourced_component_is_not_sent_to_internal_calculator(self, gateway):
        PriceItem.objects.create(category="postpress", name="Пружина в бобине, 30см", aliases="пружина", unit_price="3")
        gateway.return_value = ({"selected_source": "supplier_price", "route": "outsourcing", "confidence": .9, "reason": "Закупаем готовое", "components": [{"name": "Готовый товар", "source": "outsourcing", "source_reason": "Нет технологии", "kind": "material", "operation_item_ids": []}], "questions": [], "warnings": []}, {})

        result = analyze_production_route({"name": "Готовый сувенир на пружине", "quantity": 100, "requirements": {}})

        self.assertEqual(result["route"], "outsourcing")
        self.assertEqual(result["cost_options"], [])

    @patch("tenders.services._ai_gateway_json")
    def test_manager_answers_are_used_and_canon_is_a_supported_calculator(self, gateway):
        gateway.return_value = ({"selected_source": "psodin_canon", "calculator": "canon", "route": "internal", "confidence": .9, "reason": "Широкоформатная печать", "components": [], "questions": [], "warnings": []}, {})

        result = analyze_production_route({"name": "Постер A1", "quantity": 2, "manager_answers": {"Материал?": "Plain paper 80g"}, "requirements": {}})

        self.assertEqual(result["calculator"], "canon")
        self.assertIn("Plain paper 80g", gateway.call_args.args[0])

    @patch("tenders.services._ai_gateway_json")
    def test_supplier_source_cannot_accidentally_use_print_shop_catalog(self, gateway):
        PriceItem.objects.create(category="embossing", name="Приладка", aliases="тиснение", unit_price="1000")
        gateway.return_value = ({
            "selected_source": "supplier_price",
            "calculator": "sheet",
            "route": "internal",
            "confidence": .9,
            "reason": "Нужен специализированный поставщик",
            "components": [{"name": "Пакет", "source": "internal", "kind": "sheet", "operation_item_ids": []}],
            "questions": [],
            "warnings": [],
        }, {})

        result = analyze_production_route({"name": "Подарочный пакет А3 с тиснением", "quantity": 100, "requirements": {}})

        self.assertEqual(result["selected_source"], "supplier_price")
        self.assertEqual(result["calculator"], "none")
        self.assertEqual(result["cost_options"], [])
