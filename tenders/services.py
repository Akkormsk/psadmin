import json
import logging
import os
import re
import base64
import hashlib
import ipaddress
import math
import socket
import struct
import zipfile
import time
from xml.etree import ElementTree
from concurrent.futures import ThreadPoolExecutor
from html.parser import HTMLParser
from difflib import SequenceMatcher
from io import BytesIO
from decimal import Decimal, InvalidOperation, ROUND_CEILING, ROUND_HALF_UP
from pathlib import Path
from urllib.error import HTTPError, URLError
from urllib.parse import urlparse
from urllib.request import Request, urlopen

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader
import pypdfium2 as pdfium
import xlrd
import olefile
from PIL import Image
from django.core.cache import cache
from django.utils import timezone


logger = logging.getLogger(__name__)


MONEY = Decimal("0.01")
AI_MAX_SOURCE_CHARS = 120_000
AI_MAX_SCAN_PAGES = 12
AI_SCAN_MAX_SIDE = 2600
OFFICE_EMBEDDED_FILE_LIMIT = 20
OFFICE_EMBEDDED_BYTES_LIMIT = 40 * 1024 * 1024
OFFICE_IMAGE_BYTES_LIMIT = 12 * 1024 * 1024
WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


class TenderAIError(Exception):
    pass


class _VisibleTextParser(HTMLParser):
    def __init__(self):
        super().__init__()
        self.parts = []
        self.hidden = 0
        self.tables = []
        self._table = None
        self._row = None
        self._cell = None
        self._cell_span = (1, 1)
        self._rowspans = {}
        self._next_cell_index = 0

    def handle_starttag(self, tag, attrs):
        if tag in {"script", "style", "noscript", "svg"}:
            self.hidden += 1
            return
        if self.hidden:
            return
        if tag == "table" and self._table is None:
            self._table = []
            self._rowspans = {}
        elif tag == "tr" and self._table is not None:
            width = max(self._rowspans.keys(), default=-1) + 1
            self._row = [""] * width
            for column, (remaining, value) in list(self._rowspans.items()):
                self._row[column] = value
                if remaining <= 1:
                    del self._rowspans[column]
                else:
                    self._rowspans[column] = (remaining - 1, value)
            self._next_cell_index = 0
        elif tag in {"th", "td"} and self._row is not None:
            self._cell = []
            attributes = dict(attrs)
            try:
                colspan = max(1, min(40, int(attributes.get("colspan", 1))))
                rowspan = max(1, min(300, int(attributes.get("rowspan", 1))))
            except (TypeError, ValueError):
                colspan, rowspan = 1, 1
            self._cell_span = (colspan, rowspan)

    def handle_endtag(self, tag):
        if tag in {"script", "style", "noscript", "svg"} and self.hidden:
            self.hidden -= 1
            return
        if self.hidden:
            return
        if tag in {"th", "td"} and self._cell is not None:
            value = _cell_text(" ".join(self._cell))
            while self._next_cell_index < len(self._row) and self._row[self._next_cell_index]:
                self._next_cell_index += 1
            colspan, rowspan = self._cell_span
            required = self._next_cell_index + colspan
            if len(self._row) < required:
                self._row.extend([""] * (required - len(self._row)))
            for column in range(self._next_cell_index, required):
                self._row[column] = value
                if rowspan > 1:
                    self._rowspans[column] = (rowspan - 1, value)
            self._next_cell_index = required
            self._cell = None
        elif tag == "tr" and self._row is not None:
            if any(self._row):
                self._table.append(self._row[:40])
            self._row = None
        elif tag == "table" and self._table is not None:
            if self._table:
                self.tables.append(self._table[:300])
            self._table = None

    def handle_data(self, data):
        if not self.hidden and data.strip():
            value = data.strip()
            self.parts.append(value)
            if self._cell is not None:
                self._cell.append(value)


def _cell_text(value):
    if value is None:
        return ""
    return re.sub(r"\s+", " ", str(value)).strip()


def _normalized_text(value):
    return re.sub(r"[^a-zа-я0-9]+", " ", _cell_text(value).lower().replace("ё", "е")).strip()


def _compact_item_name(value):
    """Remove procurement boilerplate while keeping the recognizable product name."""
    source = _cell_text(value).strip(" \t\r\n.;,")
    if not source:
        return ""
    generic = re.search(
        r"(?:оказание\s+)?услуг(?:и)?\s+по\s+(?:изготовлению|производству|поставке)|"
        r"изготовлени[юя]\s+и\s+поставк[ае]|поставк[ае]\s+(?:товара|продукции)",
        source,
        flags=re.IGNORECASE,
    )
    if generic and "(" in source:
        # NMCK documents often put the actual product after a generic service name.
        candidate = source.rsplit("(", 1)[-1].rstrip(" )].;,")
        if len(candidate) >= 3:
            source = candidate
    source = re.sub(
        r"^(?:оказание\s+)?услуг(?:и)?\s+по\s+(?:изготовлению|производству|поставке)\s+(?:и\s+поставке\s+)?",
        "",
        source,
        flags=re.IGNORECASE,
    )
    source = re.sub(r"^(?:изготовление\s+и\s+поставка|поставка)\s+", "", source, flags=re.IGNORECASE)
    source = _cell_text(source).strip(" \t\r\n()[]{}.;,")
    if len(source) > 140:
        source = source[:140].rsplit(" ", 1)[0].rstrip(" ,.;:-")
    return source[:500]


def _strip_shared_item_boilerplate(items):
    """Remove a repeated document-level prefix without guessing product semantics."""
    names = [_compact_item_name(item.get("name")) for item in items]
    if len(names) < 2:
        return names
    prefixes = {}
    for name in names:
        match = re.match(r"^(.{3,120}?)\s*[:;—–-]\s+(.{2,})$", name)
        if not match:
            continue
        key = re.sub(r"\s+", " ", match.group(1).lower().replace("ё", "е")).strip()
        prefixes.setdefault(key, []).append(match.group(1))
    threshold = max(2, math.ceil(len(names) * .7))
    shared = next((values[0] for values in prefixes.values() if len(values) >= threshold), "")
    if not shared:
        return names
    pattern = re.compile(rf"^{re.escape(shared)}\s*[:;—–-]\s+", flags=re.IGNORECASE)
    return [pattern.sub("", name, count=1).strip() or name for name in names]


def _shorten_structured_item_names(items):
    """Use one small AI call for semantic names; keep a deterministic fallback."""
    fallback = _strip_shared_item_boilerplate(items)
    for item, name in zip(items, fallback):
        item["name"] = name[:500]
    if not os.getenv("TIMEWEB_AI_API_KEY", "").strip() or not items:
        return items, {"prompt_tokens": 0, "completion_tokens": 0}, None
    payload = [{"index": index, "name": item["name"]} for index, item in enumerate(items)]
    schema = '{"items":[{"index":0,"name":"Короткое рабочее название"}]}'
    prompt = f"""Сократи рабочие наименования товарных позиций из одного документа НМЦК.
Рассматривай список целиком: удаляй одинаковые канцелярские вводные, описание закупочной услуги и повторяющиеся слова, которые не помогают отличить позиции.
Обязательно сохраняй вид продукции (карта, лифлет, блокнот, футболка и т. п.), собственное название, номер варианта и характеристики, отличающие одну позицию от другой.
Не обобщай разные товары, не добавляй сведения и не меняй порядок. Желательная длина — до 80 символов.
Верни каждый исходный index ровно один раз. Только JSON: {schema}

ПОЗИЦИИ:
{json.dumps(payload, ensure_ascii=False)}"""
    try:
        result, usage = _ai_gateway_json(prompt, max_tokens=min(2200, 300 + len(items) * 80))
        received = {}
        for value in result.get("items", []) if isinstance(result.get("items"), list) else []:
            try:
                index = int(value.get("index"))
            except (TypeError, ValueError, AttributeError):
                continue
            name = _compact_item_name(value.get("name")) if isinstance(value, dict) else ""
            if 0 <= index < len(items) and name:
                received[index] = name
        if len(received) != len(items):
            return items, usage, "Часть названий сокращена локально: модель вернула неполный список."
        for index, item in enumerate(items):
            item["name"] = received[index][:500]
        return items, usage, None
    except TenderAIError:
        return items, {"prompt_tokens": 0, "completion_tokens": 0}, "Названия сокращены локально: AI Gateway был недоступен."


def _decimal_text(value):
    text = format(value, "f")
    return text.rstrip("0").rstrip(".") if "." in text else text


def _parse_document_decimal(value):
    """Parse a number copied from Russian or international tender documents."""
    if value in (None, ""):
        return None
    if isinstance(value, Decimal):
        return value
    if isinstance(value, (int, float)):
        return Decimal(str(value))
    text = str(value).strip().replace("\u00a0", "").replace("\u202f", "").replace(" ", "")
    text = re.sub(r"[^0-9,\.\-+]", "", text).strip(".,")
    if not text or text in {"-", "+", ".", ","}:
        return None
    comma, dot = text.rfind(","), text.rfind(".")
    if comma >= 0 and dot >= 0:
        decimal_separator = "," if comma > dot else "."
        thousands_separator = "." if decimal_separator == "," else ","
        text = text.replace(thousands_separator, "").replace(decimal_separator, ".")
    elif comma >= 0:
        parts = text.split(",")
        text = "".join(parts) if len(parts) > 2 or (len(parts[-1]) == 3 and len(parts[0]) > 3) else ".".join(parts)
    elif dot >= 0:
        parts = text.split(".")
        text = "".join(parts) if len(parts) > 2 or (len(parts[-1]) == 3 and len(parts[0]) > 3) else ".".join(parts)
    try:
        return Decimal(text)
    except InvalidOperation:
        return None


def _source_text_quality(text):
    """Reject non-empty but unusable PDF text layers before they block OCR."""
    text = text or ""
    letters = re.findall(r"[A-Za-zА-Яа-яЁё]", text)
    words = re.findall(r"[A-Za-zА-Яа-яЁё]{3,}", text)
    cyrillic = re.findall(r"[А-Яа-яЁё]", text)
    slash_codes = re.findall(r"/(?:i)?\d+", text, flags=re.IGNORECASE)
    replacement_chars = text.count("�")
    cyrillic_share = len(cyrillic) / max(1, len(letters))
    suspicious_share = (len(slash_codes) + replacement_chars) / max(1, len(words))
    usable = len(words) >= 3 and cyrillic_share >= .15 and suspicious_share <= .8
    return {
        "usable": usable,
        "word_count": len(words),
        "cyrillic_share": round(cyrillic_share, 3),
        "suspicious_tokens": len(slash_codes) + replacement_chars,
    }


def _requires_visual_recognition(upload, source):
    return Path(upload.name).suffix.lower() == ".pdf" and not _source_text_quality(source)["usable"]


def _pdf_unreadable_pages(upload):
    if Path(upload.name).suffix.lower() != ".pdf":
        return []
    upload.seek(0)
    try:
        reader = PdfReader(upload)
        pages = [index for index, page in enumerate(reader.pages[:100]) if not _source_text_quality(page.extract_text() or "")["usable"]]
    finally:
        upload.seek(0)
    return pages


def _table_text(rows):
    result = []
    for row in rows:
        values = [_cell_text(value) for value in row]
        while values and not values[-1]:
            values.pop()
        if any(values):
            result.append(" | ".join(values))
    return "\n".join(result)


def _docx_cell_text(cell):
    """Read all text in a Word cell, including nested tables and controls."""
    return _cell_text(" ".join(cell._tc.xpath(".//w:t/text()")))


def _docx_table_rows(table):
    """Collapse repeated wrappers produced by horizontally merged cells."""
    for row in table.rows:
        values = []
        seen_cells = set()
        for cell in row.cells:
            if cell._tc in seen_cells:
                continue
            seen_cells.add(cell._tc)
            values.append(_docx_cell_text(cell))
        yield values


def _safe_office_zip(upload):
    upload.seek(0)
    data = upload.read()
    upload.seek(0)
    try:
        return zipfile.ZipFile(BytesIO(data))
    except (zipfile.BadZipFile, OSError) as exc:
        raise TenderAIError("Файл Office повреждён или имеет неподдерживаемую внутреннюю структуру.") from exc


def _xlsx_text(upload):
    upload.seek(0)
    workbook = load_workbook(upload, read_only=True, data_only=True)
    parts = []
    try:
        for sheet in workbook.worksheets:
            if not sheet.max_row or not sheet.max_column:
                sheet.calculate_dimension(force=True)
            rows = sheet.iter_rows(
                min_row=1,
                max_row=min(sheet.max_row or 1, 1000),
                max_col=min(sheet.max_column or 1, 60),
                values_only=True,
            )
            parts.append(f"ЛИСТ: {sheet.title}\n{_table_text(rows)}")
    finally:
        workbook.close()
        upload.seek(0)
    return "\n\n".join(parts)


def _xls_text(data):
    book = xlrd.open_workbook(file_contents=data)
    parts = []
    for sheet in book.sheets():
        rows = (sheet.row_values(index) for index in range(min(sheet.nrows, 1000)))
        parts.append(f"ЛИСТ: {sheet.name}\n{_table_text(rows)}")
    return "\n\n".join(parts)


def _embedded_spreadsheet(name, data):
    """Return a named in-memory XLS/XLSX, including common OLE .bin wrappers."""
    suffix = Path(name).suffix.lower()
    candidates = [(suffix, data)]
    if suffix == ".bin" and olefile.isOleFile(BytesIO(data)):
        try:
            container = olefile.OleFileIO(BytesIO(data))
            for stream_name in container.listdir(streams=True, storages=False):
                if stream_name[-1].lower() in {"package", "\x01ole10native"}:
                    payload = container.openstream(stream_name).read()
                    marker = payload.find(b"PK\x03\x04")
                    if marker >= 0:
                        candidates.insert(0, (".xlsx", payload[marker:]))
            container.close()
        except Exception:
            pass
        candidates.append((".xls", data))
    for kind, payload in candidates:
        stream = BytesIO(payload)
        if kind in {".xlsx", ".xlsm"}:
            stream.name = f"{Path(name).stem}.xlsx"
            try:
                return stream, _xlsx_text(stream)
            except Exception:
                continue
        if kind == ".xls":
            stream.name = f"{Path(name).stem}.xls"
            try:
                return stream, _xls_text(payload)
            except Exception:
                continue
    return None, ""


def _docx_package(upload):
    """Read visible Word content plus safe embedded spreadsheets and raster images."""
    upload.seek(0)
    document = Document(upload)
    parts = [_cell_text(paragraph.text) for paragraph in document.paragraphs if _cell_text(paragraph.text)]
    for index, table in enumerate(document.tables, start=1):
        parts.append(f"ТАБЛИЦА {index}\n{_table_text(_docx_table_rows(table))}")

    embedded, images, warnings = [], [], []
    archive = _safe_office_zip(upload)
    try:
        # python-docx omits top-level content controls, text boxes, headers,
        # footers and notes. Recover their text directly from OOXML.
        xml_names = [
            name for name in archive.namelist()
            if name == "word/document.xml"
            or re.fullmatch(r"word/(?:header|footer)\d+\.xml", name)
            or name in {"word/footnotes.xml", "word/endnotes.xml", "word/comments.xml"}
        ]
        visible_text = "\n".join(parts)
        extra_text = []
        for name in xml_names:
            try:
                root = ElementTree.fromstring(archive.read(name))
            except (ElementTree.ParseError, KeyError):
                continue
            values = [_cell_text(node.text) for node in root.iter(f"{{{WORD_NS}}}t") if _cell_text(node.text)]
            joined = " ".join(values)
            if joined and joined not in visible_text:
                extra_text.append(f"ДОПОЛНИТЕЛЬНЫЙ ТЕКСТ WORD ({Path(name).name})\n{joined}")
        parts.extend(extra_text)

        embedded_bytes = 0
        for name in [value for value in archive.namelist() if value.startswith("word/embeddings/")][:OFFICE_EMBEDDED_FILE_LIMIT]:
            info = archive.getinfo(name)
            if info.file_size <= 0 or embedded_bytes + info.file_size > OFFICE_EMBEDDED_BYTES_LIMIT:
                warnings.append(f"Пропущено слишком большое вложение: {Path(name).name}")
                continue
            data = archive.read(name)
            embedded_bytes += len(data)
            suffix = Path(name).suffix.lower()
            stream = BytesIO(data)
            stream.name = Path(name).name
            spreadsheet, text = _embedded_spreadsheet(stream.name, data)
            if spreadsheet is not None:
                embedded.append(spreadsheet)
                parts.append(f"ВСТРОЕННЫЙ EXCEL: {spreadsheet.name}\n{text}")
            else:
                warnings.append(f"Обнаружено неподдерживаемое вложение Word: {stream.name}")

        image_bytes = 0
        for name in [value for value in archive.namelist() if value.startswith("word/media/")]:
            suffix = Path(name).suffix.lower()
            if suffix not in {".png", ".jpg", ".jpeg", ".webp"}:
                # EMF/WMF commonly duplicate an embedded spreadsheet preview.
                if not embedded:
                    warnings.append(f"Обнаружено изображение неподдерживаемого формата: {Path(name).name}")
                continue
            info = archive.getinfo(name)
            if image_bytes + info.file_size > OFFICE_IMAGE_BYTES_LIMIT:
                warnings.append("Часть изображений Word пропущена из-за ограничения размера.")
                break
            data = archive.read(name)
            image_bytes += len(data)
            try:
                image = Image.open(BytesIO(data)).convert("RGB")
                image.thumbnail((AI_SCAN_MAX_SIDE, AI_SCAN_MAX_SIDE))
                output = BytesIO()
                image.save(output, format="JPEG", quality=92, optimize=True)
                images.append(f"data:image/jpeg;base64,{base64.b64encode(output.getvalue()).decode('ascii')}")
            except Exception:
                warnings.append(f"Не удалось прочитать изображение: {Path(name).name}")
    finally:
        archive.close()
        upload.seek(0)
    return {
        "text": "\n".join(value for value in parts if value).strip(),
        "embedded_spreadsheets": embedded,
        "images": images,
        "warnings": warnings,
        "components": {
            "paragraphs": len(document.paragraphs),
            "tables": len(document.tables),
            "embedded_spreadsheets": len(embedded),
            "images": len(images),
        },
    }


def _pdf_package(upload):
    upload.seek(0)
    reader = PdfReader(upload)
    pages = []
    for index, page in enumerate(reader.pages[:100], start=1):
        page_text = (page.extract_text() or "").strip()
        if page_text:
            pages.append(f"СТРАНИЦА {index}\n{page_text}")
    embedded, warnings = [], []
    try:
        attachments = reader.attachments or {}
    except Exception:
        attachments = {}
    total_bytes = 0
    for name, payloads in list(attachments.items())[:OFFICE_EMBEDDED_FILE_LIMIT]:
        for payload in payloads if isinstance(payloads, list) else [payloads]:
            if not isinstance(payload, bytes) or total_bytes + len(payload) > OFFICE_EMBEDDED_BYTES_LIMIT:
                warnings.append(f"Пропущено слишком большое вложение PDF: {name}")
                continue
            total_bytes += len(payload)
            spreadsheet, text = _embedded_spreadsheet(name, payload)
            if spreadsheet is not None:
                embedded.append(spreadsheet)
                pages.append(f"ВСТРОЕННЫЙ EXCEL PDF: {spreadsheet.name}\n{text}")
            else:
                warnings.append(f"Обнаружено неподдерживаемое вложение PDF: {name}")
    upload.seek(0)
    return {
        "text": "\n\n".join(pages).strip(),
        "embedded_spreadsheets": embedded,
        "images": [],
        "warnings": warnings,
        "components": {
            "pages": len(reader.pages),
            "embedded_spreadsheets": len(embedded),
        },
    }


def _document_package(upload):
    # Some integrations provide only a named upload proxy and delegate the
    # actual extraction to extract_tender_source. Do not turn that supported
    # path into a hard failure merely because the proxy is not file-like.
    if not callable(getattr(upload, "seek", None)) or not callable(getattr(upload, "read", None)):
        return None
    suffix = Path(upload.name).suffix.lower()
    if suffix in {".docx", ".docm"}:
        return _docx_package(upload)
    if suffix == ".pdf":
        try:
            return _pdf_package(upload)
        except Exception:
            upload.seek(0)
            return None
    return None


def _extract_legacy_doc(upload):
    """Read the main text stream from a binary Word 97-2003 document."""
    upload.seek(0)
    try:
        container = olefile.OleFileIO(BytesIO(upload.read()))
        word_stream = container.openstream("WordDocument").read()
        flags = struct.unpack_from("<H", word_stream, 10)[0]
        table_name = "1Table" if flags & 0x0200 else "0Table"
        table_stream = container.openstream(table_name).read()

        offset = 32
        csw = struct.unpack_from("<H", word_stream, offset)[0]
        offset += 2 + csw * 2
        cslw = struct.unpack_from("<H", word_stream, offset)[0]
        offset += 2 + cslw * 4
        pair_count = struct.unpack_from("<H", word_stream, offset)[0]
        offset += 2
        if pair_count <= 33:
            raise ValueError("CLX table is missing")
        fc_clx, lcb_clx = struct.unpack_from("<II", word_stream, offset + 33 * 8)
        clx = table_stream[fc_clx:fc_clx + lcb_clx]
        position = 0
        while position < len(clx) and clx[position] == 0x01:
            block_size = struct.unpack_from("<H", clx, position + 1)[0]
            position += 3 + block_size
        if position >= len(clx) or clx[position] != 0x02:
            raise ValueError("piece table is missing")
        plc_size = struct.unpack_from("<I", clx, position + 1)[0]
        plc = clx[position + 5:position + 5 + plc_size]
        piece_count = (len(plc) - 4) // 12
        cp_values = struct.unpack_from(f"<{piece_count + 1}I", plc, 0)
        pcd_offset = (piece_count + 1) * 4
        pieces = []
        for index in range(piece_count):
            char_count = cp_values[index + 1] - cp_values[index]
            raw_fc = struct.unpack_from("<I", plc, pcd_offset + index * 8 + 2)[0]
            compressed = bool(raw_fc & 0x40000000)
            file_offset = raw_fc & 0x3FFFFFFF
            if compressed:
                file_offset //= 2
                raw_text = word_stream[file_offset:file_offset + char_count]
                piece = raw_text.decode("cp1251", errors="replace")
            else:
                raw_text = word_stream[file_offset:file_offset + char_count * 2]
                piece = raw_text.decode("utf-16le", errors="replace")
            pieces.append(piece)
        container.close()
    except Exception as exc:
        raise TenderAIError("Не удалось прочитать старый Word .doc. Попробуйте пересохранить его как .docx.") from exc
    text = "".join(pieces)
    text = text.replace("\x07", " | ").replace("\x0b", "\n").replace("\x0c", "\n")
    text = re.sub(r"[\x00-\x06\x08-\x0a\x0e-\x1f]", "", text)
    return "\n".join(value.strip() for value in text.splitlines() if value.strip())


def extract_tender_source(upload):
    """Extract text/tables locally. The uploaded file is never persisted."""
    suffix = Path(upload.name).suffix.lower()
    if suffix in {".xlsx", ".xlsm"}:
        text = _xlsx_text(upload)
    elif suffix == ".xls":
        book = xlrd.open_workbook(file_contents=upload.read())
        parts = []
        for sheet in book.sheets():
            rows = (sheet.row_values(index) for index in range(min(sheet.nrows, 1000)))
            parts.append(f"ЛИСТ: {sheet.name}\n{_table_text(rows)}")
        text = "\n\n".join(parts)
    elif suffix in {".docx", ".docm"}:
        text = _docx_package(upload)["text"]
    elif suffix == ".doc":
        text = _extract_legacy_doc(upload)
    elif suffix == ".pdf":
        text = _pdf_package(upload)["text"]
    else:
        raise TenderAIError("Поддерживаются .xlsx, .xlsm, .xls, .doc, .docx, .docm и .pdf.")
    text = text.strip()
    if not text and suffix == ".pdf":
        return "", False
    if not text:
        raise TenderAIError("В документе не найден текст. Сканированные PDF пока не поддерживаются.")
    return text[:AI_MAX_SOURCE_CHARS], len(text) > AI_MAX_SOURCE_CHARS


def _control_nmck_total(source):
    patterns = (
        r"начальн\w*\s*\(максимальн\w*\)\s*цен\w*\s*контракт\w*[^\d]{0,80}([\d\s]+(?:[,.]\d{1,2})?)",
        r"итог\w*\s*нмцк[^\d]{0,40}([\d\s]+(?:[,.]\d{1,2})?)",
    )
    values = []
    for pattern in patterns:
        for match in re.finditer(pattern, source, flags=re.IGNORECASE):
            value = _parse_document_decimal(match.group(1))
            if value and value > 0:
                values.append(_money(value))
    return values[-1] if values else None


def _nmck_validation_warnings(items, source):
    warnings = []
    if not items:
        return ["Документ похож на НМЦК, но структурированные позиции не найдены."]
    control_total = _control_nmck_total(source)
    if control_total is not None:
        rows_total = _money(sum((_parse_document_decimal(value.get("nmck_total")) or Decimal("0")) for value in items))
        tolerance = max(Decimal("0.10"), Decimal(len(items)) * Decimal("0.02"))
        if abs(rows_total - control_total) > tolerance:
            warnings.append(f"Сумма найденных позиций {rows_total} ₽ не совпадает с итогом документа {control_total} ₽.")
    return warnings


def _pdf_page_count(upload):
    upload.seek(0)
    document = pdfium.PdfDocument(upload.read())
    page_count = len(document)
    document.close()
    upload.seek(0)
    if not page_count:
        raise TenderAIError("В PDF нет страниц.")
    return page_count


def _scan_pdf_images(upload, start_page=0, page_limit=None):
    """Render scanned PDF pages in memory for multimodal recognition."""
    upload.seek(0)
    document = pdfium.PdfDocument(upload.read())
    page_count = len(document)
    if not page_count:
        raise TenderAIError("В PDF нет страниц.")
    start_page = max(0, int(start_page or 0))
    end_page = page_count if page_limit is None else min(page_count, start_page + max(1, int(page_limit)))
    if start_page >= page_count:
        document.close()
        upload.seek(0)
        return []
    images = []
    for page_number in range(start_page, end_page):
        page = document[page_number]
        # Small tender tables contain negations and decimal values where one
        # missed glyph changes the commercial meaning. Render above screen DPI
        # before the bounded thumbnail instead of enlarging a blurry raster.
        bitmap = page.render(scale=3)
        image = bitmap.to_pil().convert("RGB")
        image.thumbnail((AI_SCAN_MAX_SIDE, AI_SCAN_MAX_SIDE))
        output = BytesIO()
        image.save(output, format="JPEG", quality=92, optimize=True)
        images.append(base64.b64encode(output.getvalue()).decode("ascii"))
        bitmap.close()
        page.close()
    document.close()
    upload.seek(0)
    return images


def _scan_pdf_selected_images(upload, page_numbers):
    page_numbers = sorted({int(value) for value in page_numbers if int(value) >= 0})[:AI_MAX_SCAN_PAGES]
    if not page_numbers:
        return []
    upload.seek(0)
    document = pdfium.PdfDocument(upload.read())
    images = []
    try:
        for page_number in page_numbers:
            if page_number >= len(document):
                continue
            page = document[page_number]
            bitmap = page.render(scale=3)
            image = bitmap.to_pil().convert("RGB")
            image.thumbnail((AI_SCAN_MAX_SIDE, AI_SCAN_MAX_SIDE))
            output = BytesIO()
            image.save(output, format="JPEG", quality=92, optimize=True)
            images.append(f"data:image/jpeg;base64,{base64.b64encode(output.getvalue()).decode('ascii')}")
            bitmap.close()
            page.close()
    finally:
        document.close()
        upload.seek(0)
    return images


def _visual_gateway_responses(prompt, upload, max_tokens):
    """Send a long scan in bounded page batches and preserve page order."""
    page_count = _pdf_page_count(upload)
    if page_count <= AI_MAX_SCAN_PAGES:
        return [_ai_gateway_json(prompt, upload=upload, scan_ocr=True, max_tokens=max_tokens)]
    responses = []
    for start_page in range(0, page_count, AI_MAX_SCAN_PAGES):
        end_page = min(page_count, start_page + AI_MAX_SCAN_PAGES)
        images = _scan_pdf_images(upload, start_page=start_page, page_limit=AI_MAX_SCAN_PAGES)
        page_prompt = (
            f"{prompt}\n\nПАКЕТ СТРАНИЦ: {start_page + 1}–{end_page} из {page_count}. "
            "Верни только позиции и требования, действительно видимые на этих страницах; "
            "не повторяй товары из контекста, если их нет в этом пакете."
        )
        image_urls = [f"data:image/jpeg;base64,{value}" for value in images]
        responses.append(_ai_gateway_json(page_prompt, max_tokens=max_tokens, image_data_urls=image_urls))
    return responses


def _json_from_model(content):
    content = content.strip()
    if content.startswith("```"):
        content = re.sub(r"^```(?:json)?\s*|\s*```$", "", content, flags=re.IGNORECASE)
    decoder = json.JSONDecoder(strict=False)
    starts = [index for index, value in enumerate(content) if value in "{["]
    last_error = None
    for start in starts or [0]:
        try:
            result, _ = decoder.raw_decode(content[start:])
            return result
        except json.JSONDecodeError as exc:
            last_error = exc
    raise TenderAIError("Модель вернула ответ в неожиданном формате. Попробуйте ещё раз.") from last_error


def _ai_gateway_json(prompt, upload=None, scan_ocr=False, max_tokens=6000, image_data_urls=None):
    api_key = os.getenv("TIMEWEB_AI_API_KEY", "").strip()
    base_url = os.getenv("TIMEWEB_AI_BASE_URL", "https://api.timeweb.ai/v1").rstrip("/")
    model = os.getenv("TIMEWEB_AI_MODEL", "openai/gpt-4.1-mini").strip()
    if not api_key:
        raise TenderAIError("AI Gateway ещё не настроен.")
    user_content = prompt
    if scan_ocr or image_data_urls:
        user_content = [{"type": "text", "text": prompt}]
        if scan_ocr:
            user_content.extend({"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{image}", "detail": "high"}} for image in _scan_pdf_images(upload))
        if image_data_urls:
            user_content.extend({"type": "image_url", "image_url": {"url": image, "detail": "high"}} for image in image_data_urls)
    total_usage = {"prompt_tokens": 0, "completion_tokens": 0}
    for attempt in range(2):
        messages = [
            {"role": "system", "content": "Ты точно анализируешь документы и таблицы. Отвечай только валидным JSON без markdown."},
            {"role": "user", "content": user_content},
        ]
        if attempt:
            messages[0]["content"] += " Предыдущая попытка содержала синтаксическую ошибку. Особенно тщательно проверь кавычки, запятые и закрывающие скобки."
        payload = json.dumps({"model": model, "messages": messages, "temperature": 0, "max_tokens": max_tokens}, ensure_ascii=False).encode("utf-8")
        request = Request(f"{base_url}/chat/completions", data=payload, headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}, method="POST")
        response_data = None
        last_network_error = None
        for network_attempt in range(3):
            try:
                with urlopen(request, timeout=90) as response:
                    response_data = json.loads(response.read().decode("utf-8"))
                break
            except HTTPError as exc:
                if exc.code not in {429, 500, 502, 503, 504}:
                    try:
                        detail = json.loads(exc.read().decode("utf-8")).get("error", {}).get("message")
                    except Exception:
                        detail = None
                    raise TenderAIError(detail or "AI Gateway отклонил запрос.") from exc
                last_network_error = exc
            except (URLError, TimeoutError, ConnectionError, OSError, json.JSONDecodeError) as exc:
                last_network_error = exc
            if network_attempt < 2:
                time.sleep(1 + network_attempt * 2)
        if response_data is None:
            raise TenderAIError("AI Gateway не ответил после трёх попыток. Попробуйте позже.") from last_network_error
        usage = response_data.get("usage", {})
        total_usage["prompt_tokens"] += usage.get("prompt_tokens", 0) or 0
        total_usage["completion_tokens"] += usage.get("completion_tokens", 0) or 0
        try:
            content = response_data["choices"][0]["message"]["content"]
        except (KeyError, IndexError, TypeError) as exc:
            raise TenderAIError("AI Gateway не вернул результат анализа.") from exc
        try:
            return _json_from_model(content), total_usage
        except TenderAIError:
            if not attempt:
                continue
            raise


def _validate_public_url(value):
    parsed = urlparse(value)
    if parsed.scheme not in {"http", "https"} or not parsed.hostname:
        raise TenderAIError("Укажите корректную публичную ссылку http или https.")
    try:
        addresses = {item[4][0] for item in socket.getaddrinfo(parsed.hostname, parsed.port or (443 if parsed.scheme == "https" else 80))}
    except OSError as exc:
        raise TenderAIError("Не удалось найти сайт по указанной ссылке.") from exc
    for address in addresses:
        ip = ipaddress.ip_address(address)
        if not ip.is_global:
            raise TenderAIError("Локальные и служебные адреса нельзя использовать как источник.")
    return value


def _format_html_tables(tables, limit=18_000):
    parts = []
    for table_index, rows in enumerate(tables[:20], start=1):
        rendered = []
        for row_index, row in enumerate(rows[:300], start=1):
            cells = [_cell_text(value).replace("|", "/") for value in row]
            if any(cells):
                rendered.append(f"СТРОКА {row_index}: " + " | ".join(cells))
        if rendered:
            parts.append(f"HTML-ТАБЛИЦА {table_index}\n" + "\n".join(rendered))
        if sum(len(value) for value in parts) >= limit:
            break
    return "\n\n".join(parts)[:limit]


def _tier_bounds(value):
    text = _normalized_text(value)
    numbers = [Decimal(item.replace(",", ".")) for item in re.findall(r"\d+(?:[.,]\d+)?", text)]
    if not numbers:
        return None
    if re.search(r"\b(?:от|свыше|более)\b", text) or "+" in str(value):
        return numbers[0], None
    if re.search(r"\bдо\b", text):
        return Decimal("0"), numbers[0]
    if len(numbers) >= 2 and re.search(r"[-–—]", str(value)):
        return min(numbers[0], numbers[1]), max(numbers[0], numbers[1])
    return None


def _tier_contains(bounds, quantity):
    if not bounds:
        return False
    lower, upper = bounds
    return quantity >= lower and (upper is None or quantity <= upper)


def _price_cell(value):
    text = _cell_text(value).replace("\xa0", " ")
    matches = re.findall(r"(?<!\d)(\d[\d ]*(?:[.,]\d{1,2})?)(?!\d)", text)
    if len(matches) != 1:
        return None
    try:
        return Decimal(matches[0].replace(" ", "").replace(",", "."))
    except InvalidOperation:
        return None


def _dimension_signatures(value):
    normalized = _cell_text(value).lower().replace("×", "х").replace("x", "х")
    signatures = set()
    for match in re.finditer(r"\d+(?:[.,]\d+)?(?:\s*х\s*\d+(?:[.,]\d+)?){1,3}", normalized):
        signatures.add(re.sub(r"\s+", "", match.group(0)).replace(",", "."))
    return signatures


def _select_html_price_quote(tables, context):
    if not tables or not isinstance(context, dict):
        return None
    line = context.get("line") if isinstance(context.get("line"), dict) else {}
    query = " ".join([
        _cell_text(line.get("name")),
        json.dumps(line.get("requirements", {}), ensure_ascii=False),
        _cell_text(context.get("feedback")),
        _cell_text(context.get("target_name")),
    ])
    try:
        quantity = Decimal(str(line.get("quantity") or 0).replace(",", "."))
    except (InvalidOperation, TypeError, ValueError):
        return None
    if quantity <= 0:
        return None
    query_normalized = _normalized_text(query)
    query_dimensions = _dimension_signatures(query)
    markers = {
        "армани": ("армани",), "мокрый шелк": ("мокрый шелк",), "полиэфирный шелк": ("полиэфир",),
        "джерси": ("джерси",), "габардин": ("габардин",), "атлас": ("атлас",),
        "оверлок": ("оверлок",), "горячий рез": ("горяч", "рез"),
        "треугольный": ("треуголь",), "квадратный": ("квадрат",),
    }
    requested_markers = {
        label for label, roots in markers.items() if all(root in query_normalized for root in roots)
    }
    candidates = []
    for table_index, rows in enumerate(tables, start=1):
        headers = None
        for row_index, row in enumerate(rows, start=1):
            tier_columns = {index: _tier_bounds(cell) for index, cell in enumerate(row)}
            tier_columns = {index: bounds for index, bounds in tier_columns.items() if bounds}
            if len(tier_columns) >= 2:
                headers = (row, tier_columns)
                continue
            if not headers or not row:
                continue
            header_row, tier_columns = headers
            matching_columns = [index for index, bounds in tier_columns.items() if _tier_contains(bounds, quantity)]
            if len(matching_columns) != 1:
                continue
            column = matching_columns[0]
            if column >= len(row):
                continue
            price = _price_cell(row[column])
            if price is None or price <= 0:
                continue
            label = _cell_text(row[0])
            row_normalized = _normalized_text(label)
            row_dimensions = _dimension_signatures(label)
            score = 0
            evidence = []
            if query_dimensions:
                common_dimensions = query_dimensions & row_dimensions
                if common_dimensions:
                    score += 100
                    evidence.append(f"размер {sorted(common_dimensions)[0]}")
                elif row_dimensions:
                    score -= 100
            for marker in requested_markers:
                roots = markers[marker]
                if all(root in row_normalized for root in roots):
                    score += 18
                    evidence.append(marker)
                elif any(root in row_normalized for root in roots):
                    score += 5
                else:
                    score -= 12
            query_tokens = {value for value in re.findall(r"[a-zа-я]{4,}", query_normalized) if value not in {"товар", "цена", "тираж", "штук", "источник", "требования"}}
            row_tokens = set(re.findall(r"[a-zа-я]{4,}", row_normalized))
            score += min(20, len(query_tokens & row_tokens) * 3)
            candidates.append({
                "score": score,
                "table_index": table_index,
                "row_index": row_index,
                "row_label": label,
                "tier": _cell_text(header_row[column]),
                "quantity": str(quantity),
                "unit_price": str(_money(price)),
                "amount_total": str(_money(price * quantity)),
                "evidence": evidence,
                "method": "html_table_tier",
            })
    if not candidates:
        return None
    candidates.sort(key=lambda value: value["score"], reverse=True)
    best = candidates[0]
    conflicting = [value for value in candidates[1:] if value["unit_price"] != best["unit_price"] and value["score"] >= best["score"] - 8]
    minimum_score = 90 if query_dimensions else 30
    if best["score"] < minimum_score or conflicting:
        return None
    best["confidence"] = "exact" if best["score"] >= 110 else "probable"
    return best


def _fetch_public_page(value):
    value = _validate_public_url(value)
    request = Request(value, headers={"User-Agent": "PSAdmin tender calculator/1.0", "Accept": "text/html,text/plain;q=0.9,*/*;q=0.5"})
    try:
        with urlopen(request, timeout=15) as response:
            final_url = _validate_public_url(response.geturl())
            content_type = response.headers.get_content_type()
            raw = response.read(1_500_001)
            charset = response.headers.get_content_charset() or "utf-8"
    except (HTTPError, URLError, TimeoutError) as exc:
        raise TenderAIError("Не удалось прочитать страницу поставщика. Можно приложить PDF или скриншот.") from exc
    if len(raw) > 1_500_000:
        raise TenderAIError("Страница поставщика слишком большая. Приложите нужный фрагмент или прайс.")
    if content_type not in {"text/html", "text/plain", "application/xhtml+xml"}:
        raise TenderAIError("По ссылке нет читаемой страницы. Приложите PDF или скриншот.")
    text = raw.decode(charset, errors="replace")
    tables = []
    if content_type != "text/plain":
        parser = _VisibleTextParser()
        parser.feed(text)
        tables = parser.tables
        visible_text = "\n".join(parser.parts)
        table_text = _format_html_tables(tables)
        text = f"{table_text}\n\nТЕКСТ СТРАНИЦЫ:\n{visible_text}" if table_text else visible_text
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if not text:
        raise TenderAIError("На странице не найден читаемый прайс. Приложите скриншот.")
    return text[:40_000], final_url, tables


def _image_data_url(upload):
    upload.seek(0)
    try:
        image = Image.open(upload).convert("RGB")
    except Exception as exc:
        raise TenderAIError("Не удалось прочитать изображение источника.") from exc
    image.thumbnail((1800, 1800))
    output = BytesIO()
    image.save(output, format="JPEG", quality=88, optimize=True)
    upload.seek(0)
    return f"data:image/jpeg;base64,{base64.b64encode(output.getvalue()).decode('ascii')}"


def extract_calculation_source(source_text="", source_url="", upload=None, selection_context=None):
    parts, source_type, resolved_url = [], "text", source_url
    structured_data = {}
    if source_text.strip():
        parts.append(source_text.strip()[:12_000])
    if source_url.strip():
        page_text, resolved_url, tables = _fetch_public_page(source_url.strip())
        quote = _select_html_price_quote(tables, selection_context)
        if quote:
            structured_data["price_quote"] = quote
            page_text = (
                "ЦЕНА, ПРОВЕРЕННАЯ БЭКЕНДОМ ПО HTML-ТАБЛИЦЕ:\n"
                f"{quote['row_label']} | диапазон {quote['tier']} | {quote['unit_price']} ₽/шт. | "
                f"{quote['quantity']} шт. = {quote['amount_total']} ₽\n\n{page_text}"
            )
        structured_data["html_table_count"] = len(tables)
        parts.append(f"СТРАНИЦА ПОСТАВЩИКА:\n{page_text}")
        source_type = "link"
    if upload is not None:
        suffix = Path(upload.name).suffix.lower()
        if suffix in {".png", ".jpg", ".jpeg", ".webp"}:
            prompt = "Прочитай скриншот прайса или страницы поставщика. Извлеки название поставщика, товар/услугу, цену, единицу цены, тиражные условия, дату и существенные примечания. Не выдумывай. Верни JSON: {\"text\":\"подробно извлечённые данные\"}"
            result, _ = _ai_gateway_json(prompt, max_tokens=1600, image_data_urls=[_image_data_url(upload)])
            parts.append(_cell_text(result.get("text"))[:12_000])
            source_type = "image"
        elif suffix in {".pdf", ".doc", ".docx", ".xlsx", ".xls"}:
            extracted, _ = extract_tender_source(upload)
            parts.append(extracted[:12_000])
            source_type = "document"
        else:
            raise TenderAIError("Источник можно приложить как PDF, Word, Excel, PNG или JPG.")
    content = "\n\n".join(value for value in parts if value).strip()
    if not content:
        raise TenderAIError("Добавьте ссылку, текст или файл источника.")
    return {"content": content[:20_000], "source_type": source_type, "url": resolved_url[:1000], "structured_data": structured_data}


def recognize_tender_items(upload):
    suffix = Path(upload.name).suffix.lower()
    package = _document_package(upload)
    if package and package["embedded_spreadsheets"]:
        structured = []
        for spreadsheet in package["embedded_spreadsheets"]:
            structured.extend(_recognize_structured_nmck_xlsx(spreadsheet))
        if structured:
            structured, usage, warning = _shorten_structured_item_names(structured)
            warnings = [*package["warnings"], *([warning] if warning else [])]
            warnings.extend(_nmck_validation_warnings(structured, package["text"]))
            return {
                "items": structured,
                "warnings": warnings,
                "scan_ocr": False,
                "processing_mode": "embedded",
                "usage": usage,
                "local_parse": True,
                "components": package["components"],
            }
    if suffix in {".xlsx", ".xlsm", ".xls"}:
        structured = _recognize_structured_nmck_xlsx(upload)
        if structured:
            structured, usage, warning = _shorten_structured_item_names(structured)
            source, _ = extract_tender_source(upload)
            warnings = [warning] if warning else []
            warnings.extend(_nmck_validation_warnings(structured, source))
            return {"items": structured, "warnings": warnings, "scan_ocr": False, "processing_mode": "structured", "usage": usage, "local_parse": True}
    source, truncated = extract_tender_source(upload)
    scan_ocr = _requires_visual_recognition(upload, source)
    docx_images = package["images"] if package and suffix in {".docx", ".docm"} and not package["components"].get("tables") and not package["components"].get("embedded_spreadsheets") else []
    mixed_pdf_pages = [] if scan_ocr else _pdf_unreadable_pages(upload)
    visual_mode = scan_ocr or bool(docx_images) or bool(mixed_pdf_pages)
    schema = '{"items":[{"name":"товар","quantity":"исходное значение","nmck_unit":"исходное значение или null","nmck_total":"исходное значение или null","confidence":0.95}],"warnings":[]}'
    prompt = f"""Извлеки из документа позиции НМЦК для расчёта тендера.
Для каждой товарной позиции нужны: короткое рабочее наименование, количество, НМЦК за единицу и итоговая НМЦК всей позиции.
Убирай из названия канцелярские вводные вроде «услуги по изготовлению и поставке продукции», но сохраняй сам вид товара, номер варианта и отличающие его характеристики.
Если указаны цены коммерческих предложений поставщиков (КП 1, КП 2, КП 3 или источники цены), не используй ни одну из них как НМЦК.
В таких таблицах выбирай конечную рассчитанную колонку «Средняя цена» или «Средняя рыночная цена» — это нужная НМЦК за единицу.
Колонка «НМЦК» рядом со средней ценой обычно содержит общую стоимость позиции: перенеси её в nmck_total, а не в nmck_unit.
Не выполняй никакие вычисления. Переноси quantity, nmck_unit и nmck_total только как исходные числа из документа. Если одного из двух ценовых значений нет, верни для него null: деление и умножение выполнит сервер.
Не включай заголовки, итоги, НДС, доставку и пустые строки как товары.
Не выдумывай значения. Сомнения кратко перечисли в warnings.
Верни только JSON строго такого вида: {schema}

ДОКУМЕНТ:
{source if not scan_ocr else 'Перед тобой страницы документа в исходном порядке. Текстовый слой отсутствует или повреждён; внимательно прочитай таблицу на изображениях.'}"""
    if scan_ocr:
        responses = _visual_gateway_responses(prompt, upload, max_tokens=6000)
    elif docx_images:
        responses = [_ai_gateway_json(prompt, max_tokens=6000, image_data_urls=docx_images)]
    elif mixed_pdf_pages:
        responses = []
        for start in range(0, len(mixed_pdf_pages), AI_MAX_SCAN_PAGES):
            page_batch = mixed_pdf_pages[start:start + AI_MAX_SCAN_PAGES]
            batch_prompt = f"{prompt}\n\nДОПОЛНИТЕЛЬНО РАСПОЗНАЮТСЯ СТРАНИЦЫ: {', '.join(str(value + 1) for value in page_batch)}."
            responses.append(_ai_gateway_json(batch_prompt, max_tokens=6000, image_data_urls=_scan_pdf_selected_images(upload, page_batch)))
    else:
        responses = [_ai_gateway_json(prompt)]
    result = {"items": [], "warnings": []}
    usage = {"prompt_tokens": 0, "completion_tokens": 0}
    for partial, partial_usage in responses:
        result["items"].extend(partial.get("items", []) if isinstance(partial.get("items"), list) else [])
        result["warnings"].extend(partial.get("warnings", []) if isinstance(partial.get("warnings"), list) else [])
        usage["prompt_tokens"] += partial_usage.get("prompt_tokens", 0) or 0
        usage["completion_tokens"] += partial_usage.get("completion_tokens", 0) or 0
    items = []
    for raw in result.get("items", []):
        try:
            name = _compact_item_name(raw.get("name"))
            quantity = _parse_document_decimal(raw.get("quantity"))
            nmck_unit = _parse_document_decimal(raw.get("nmck_unit"))
            raw_total = raw.get("nmck_total")
            nmck_total = _parse_document_decimal(raw_total)
            confidence = max(0, min(1, float(raw.get("confidence", 0))))
        except (TypeError, ValueError):
            continue
        if not quantity or quantity <= 0:
            continue
        if (not nmck_unit or nmck_unit <= 0) and nmck_total and nmck_total > 0:
            nmck_unit = nmck_total / quantity
        if (not nmck_total or nmck_total <= 0) and nmck_unit and nmck_unit > 0:
            nmck_total = quantity * nmck_unit
        if name and nmck_unit and nmck_unit > 0 and nmck_total and nmck_total > 0:
            source_total = _money(nmck_total)
            calculated_total = _money(quantity * nmck_unit)
            # The source unit price is often displayed rounded to kopecks while the
            # source line total is calculated from a more precise hidden value.
            rounding_tolerance = max(Decimal("0.05"), quantity * Decimal("0.005") + Decimal("0.02"))
            items.append({
                "name": name[:500],
                "quantity": _decimal_text(quantity),
                "nmck_unit": str(_money(nmck_unit)),
                "nmck_total": str(source_total),
                "total_from_source": raw_total not in (None, "") and _parse_document_decimal(raw_total) is not None,
                "total_matches": abs(source_total - calculated_total) <= rounding_tolerance,
                "confidence": confidence,
            })
    if not items:
        details = "; ".join(package["warnings"][:3]) if package and package["warnings"] else ""
        suffix = f" Обнаруженные ограничения: {details}." if details else ""
        raise TenderAIError(f"Не удалось уверенно найти позиции с количеством и НМЦК.{suffix}")
    warnings = [str(value)[:300] for value in result.get("warnings", []) if str(value).strip()]
    if package:
        warnings = [*package["warnings"], *warnings]
    warnings.extend(_nmck_validation_warnings(items, source))
    if truncated:
        warnings.append("Документ был слишком большим: обработана основная часть содержимого.")
    if visual_mode:
        warnings.insert(0, "Документ распознан по изображению. Внимательно проверьте названия, количество, цены и итоговые суммы.")
    return {"items": items, "warnings": warnings, "scan_ocr": visual_mode, "processing_mode": "visual" if visual_mode else "text", "components": package["components"] if package else {}, "usage": {"prompt_tokens": usage.get("prompt_tokens", 0), "completion_tokens": usage.get("completion_tokens", 0)}}


def _recognize_structured_nmck_xlsx(upload):
    """Read common multi-row XLSX/XLS NMCK tables before spending AI tokens."""
    upload.seek(0)
    suffix = Path(upload.name).suffix.lower()
    workbook = None
    try:
        if suffix in {".xlsx", ".xlsm"}:
            workbook = load_workbook(upload, read_only=True, data_only=True)
            sheets = []
            for sheet in workbook.worksheets:
                if not sheet.max_row or not sheet.max_column:
                    sheet.calculate_dimension(force=True)
                sheets.append(list(sheet.iter_rows(
                    min_row=1,
                    max_row=min(sheet.max_row or 1, 500),
                    max_col=min(sheet.max_column or 1, 60),
                    values_only=True,
                )))
        elif suffix == ".xls":
            book = xlrd.open_workbook(file_contents=upload.read())
            sheets = [
                [tuple(sheet.row_values(index)[:60]) for index in range(min(sheet.nrows, 500))]
                for sheet in book.sheets()
            ]
        else:
            return []
    except Exception:
        upload.seek(0)
        return []
    best = []
    for rows in sheets:
        header_row = None
        # NMCK workbooks often split one logical header across 2–3 rows and
        # use abbreviated labels such as «Кол-во». Inspect a short window
        # instead of requiring every marker to live in one physical row.
        for index in range(min(60, len(rows))):
            window = rows[index:min(len(rows), index + 3)]
            text = " ".join(
                _cell_text(value).lower().replace("ё", "е")
                for row in window for value in row if _cell_text(value)
            )
            has_quantity = any(marker in text for marker in ("колич", "кол-во", "кол во", "объем"))
            has_name = "наимен" in text or ("характеристик" in text and "объект" in text)
            if has_name and has_quantity:
                header_row = index
                break
        if header_row is None:
            continue
        header_end = min(len(rows), header_row + 4)
        for candidate in range(header_row + 1, min(len(rows), header_row + 8)):
            row = rows[candidate]
            has_text = any(_cell_text(value) and not _cell_text(value).replace(".", "", 1).isdigit() for value in row)
            numeric_count = sum(isinstance(value, (int, float, Decimal)) for value in row)
            if has_text and numeric_count >= 2:
                header_end = candidate
                break
        headers = []
        width = max((len(row) for row in rows), default=0)
        for column in range(width):
            headers.append(" ".join(_cell_text(rows[row][column]).lower().replace("ё", "е") for row in range(header_row, header_end) if column < len(rows[row]) and _cell_text(rows[row][column])))
        name_col = next((index for index, value in enumerate(headers) if "наимен" in value), None)
        if name_col is None:
            name_col = next((index for index, value in enumerate(headers) if "характеристик" in value and "объект" in value), None)
        quantity_col = next((
            index for index, value in enumerate(headers)
            if any(marker in value for marker in ("колич", "кол-во", "кол во")) and "цен" not in value
        ), None)
        if quantity_col is None:
            # In some forms «Объём» is a group heading and the actual numeric
            # column below it is left unnamed after merged-cell expansion.
            quantity_col = next((index for index, value in enumerate(headers) if "объем" in value and "цен" not in value), None)
            if quantity_col is not None:
                adjacent = quantity_col + 1
                if adjacent < width and any(marker in headers[adjacent] for marker in ("колич", "кол-во", "кол во")):
                    quantity_col = adjacent
        # The customer's selected unit price is authoritative. Average market
        # price is only a fallback when the selected column is absent.
        unit_candidates = [index for index, value in enumerate(headers) if "минималь" in value and "за единиц" in value]
        unit_candidates += [index for index, value in enumerate(headers) if "выбран" in value and "за единиц" in value]
        unit_candidates += [index for index, value in enumerate(headers) if "началь" in value and "за единиц" in value]
        unit_candidates += [
            index for index, value in enumerate(headers)
            if "средн" in value and any(marker in value for marker in ("цен", "стоим"))
            and not any(marker in value for marker in ("сумм", "за все кол", "общая стоимость", "стоимость позиции"))
        ]
        unit_candidates += [
            index for index, value in enumerate(headers)
            if "за единиц" in value and any(marker in value for marker in ("цен", "стоим"))
        ]
        total_candidates = [
            index for index, value in enumerate(headers)
            if any(marker in value for marker in ("сумм", "за все кол", "общая стоимость", "стоимость позиции", "нмцк"))
            or ("началь" in value and "за единиц" not in value)
        ]
        if name_col is None or quantity_col is None:
            continue
        parsed = []
        for row in rows[header_end:]:
            name = _compact_item_name(row[name_col] if name_col < len(row) else "")
            try:
                quantity = Decimal(str(row[quantity_col]).replace(",", "."))
            except (InvalidOperation, TypeError, ValueError):
                continue
            if not name or quantity <= 0:
                continue
            numeric = {}
            for column, value in enumerate(row):
                try:
                    number = Decimal(str(value).replace(",", "."))
                except (InvalidOperation, TypeError, ValueError):
                    continue
                if number > 0:
                    numeric[column] = number
            unit = next((numeric[column] for column in unit_candidates if column in numeric), None)
            total = next((numeric[column] for column in reversed(total_candidates) if column in numeric and column != quantity_col), None)
            if unit and total and abs(total - unit * quantity) > max(Decimal("0.10"), quantity * Decimal("0.02")):
                matching_total = next((value for column, value in sorted(numeric.items(), reverse=True) if column != quantity_col and abs(value - unit * quantity) <= max(Decimal("0.10"), quantity * Decimal("0.02"))), None)
                total = matching_total or total
            if not unit and total:
                unit = total / quantity
            if not total and unit:
                total = unit * quantity
            if not unit or not total:
                continue
            parsed.append({
                "name": name[:500],
                "quantity": _decimal_text(quantity),
                "nmck_unit": str(_money(unit)),
                "nmck_total": str(_money(total)),
                "total_from_source": True,
                "total_matches": abs(_money(total) - _money(unit * quantity)) <= max(Decimal("0.10"), quantity * Decimal("0.02")),
                "confidence": .92,
            })
        if len(parsed) > len(best):
            best = parsed
    if workbook is not None:
        workbook.close()
    upload.seek(0)
    return best


def _classify_tender_source(file_name, source):
    """Classify a readable document while keeping contracts/instructions out."""
    text = source.lower().replace("ё", "е")
    opening = text[:6000]
    file_hint = re.sub(r"[^a-zа-я0-9]+", " ", Path(file_name).stem.lower().replace("ё", "е")).strip()
    excluded_file = any(marker in file_hint for marker in (
        "проект контракт", "проект договора", "государственн контракт",
        "требовани к содержанию", "составу заявки", "инструкц", "извещен",
        "лист согласован", "информационная карта",
    ))
    nmck_file = any(marker in file_hint for marker in ("нмцк", "обоснован", "расчет", "расчёт"))
    technical_file = (
        any(marker in file_hint for marker in ("техническ", "техзадан", "описание объекта", "ооз", "требования к товар"))
        or "тз" in file_hint.split()
    ) and not excluded_file
    nmck_score = sum(phrase in text for phrase in (
        "обоснование начальной", "расчет начальной", "нмцк",
        "средняя арифметич", "цена исполнителя", "ценовое предложение",
        "минимальная цена выбранная", "начальная (максимальная) цена контракта",
    )) + (2 if nmck_file else 0)
    technical_score = sum(phrase in opening for phrase in (
        "описание объекта закупки", "техническое задание", "технические характеристики",
        "требования к товар", "требования к услуг", "характеристики объекта закупки",
    )) + (2 if technical_file else 0)
    if excluded_file and not nmck_file:
        return "unknown"
    if nmck_score >= 2 and technical_score >= 3:
        return "mixed"
    if nmck_score >= 2:
        return "nmck"
    if technical_score >= 2:
        return "technical"
    return "unknown"


def inspect_tender_document(upload):
    """Run the cheap preflight used by the UI before a potentially slow pass."""
    suffix = Path(upload.name).suffix.lower()
    package = _document_package(upload)
    if package:
        source = package["text"][:AI_MAX_SOURCE_CHARS]
        truncated = len(package["text"]) > AI_MAX_SOURCE_CHARS
    else:
        source, truncated = extract_tender_source(upload)
    quality = _source_text_quality(source)
    visual = _requires_visual_recognition(upload, source)
    unreadable_pdf_pages = _pdf_unreadable_pages(upload) if suffix == ".pdf" else []
    mode = "visual" if visual else "text"
    if package and package["components"]["embedded_spreadsheets"]:
        mode = "embedded"
    elif package and package["images"] and not package["components"]["tables"]:
        mode = "visual"
    elif unreadable_pdf_pages:
        mode = "visual"
    role = "unknown" if visual else _classify_tender_source(upload.name, source)
    if role == "unknown" and Path(upload.name).suffix.lower() in {".xlsx", ".xls"}:
        upload.seek(0)
        if _recognize_structured_nmck_xlsx(upload):
            role = "nmck"
    upload.seek(0)
    return {
        "document_type": role,
        "processing_mode": mode,
        "truncated": truncated,
        "quality": quality,
        "components": package["components"] if package else ({"unreadable_pdf_pages": [value + 1 for value in unreadable_pdf_pages]} if unreadable_pdf_pages else {}),
        "warnings": package["warnings"] if package else [],
        "status_message": (
            "Обнаружена встроенная таблица Excel — извлекаю её локально."
            if mode == "embedded" else
            "Часть страниц не имеет надёжного текстового слоя — распознаю их изображения."
            if mode == "visual" else
            "Структура документа прочитана."
        ),
    }


def detect_tender_document_type(upload):
    """Determine the document's role independently from the upload control/order."""
    inspection = inspect_tender_document(upload)
    if inspection["document_type"] != "unknown":
        return inspection["document_type"]
    # Text extraction can differ between Excel readers and operating systems.
    # The structured parser is a stronger signal: it only succeeds when it
    # finds item names, quantities, unit NMCK and final line totals.
    return "unknown"


def _short_text_list(values, limit=12):
    if not isinstance(values, list):
        return []
    return [_cell_text(value)[:300] for value in values[:limit] if _cell_text(value)]


def _requirement_list(values):
    if not isinstance(values, list):
        return []
    requirements = []
    for value in values[:30]:
        if not isinstance(value, dict):
            continue
        label = _cell_text(value.get("label"))[:120]
        item_value = _cell_text(value.get("value"))[:500]
        source = _cell_text(value.get("source"))[:200]
        if label and item_value:
            requirements.append({"label": label, "value": item_value, "source": source})
    return requirements


def _normalized_item_name(value):
    value = _cell_text(value).lower().replace("ё", "е")
    value = re.sub(r"\b\d{2}(?:\.\d{2}){2,4}\b", " ", value)
    return re.sub(r"[^a-zа-я0-9]+", " ", value).strip()


def _match_score(source_name, source_quantity, line):
    source = _normalized_item_name(source_name)
    target = _normalized_item_name(line.get("name"))
    if not source or not target:
        return 0
    name_score = SequenceMatcher(None, source, target).ratio()
    source_tokens, target_tokens = set(source.split()), set(target.split())
    token_score = len(source_tokens & target_tokens) / max(1, len(source_tokens | target_tokens))
    score = max(name_score, token_score)
    try:
        source_qty = Decimal(str(source_quantity).replace(",", "."))
        target_qty = Decimal(str(line.get("quantity", "")).replace(",", "."))
        if source_qty > 0 and target_qty > 0:
            score += 0.08 if source_qty == target_qty else -0.12
    except (InvalidOperation, TypeError, ValueError):
        pass
    return max(0, min(1, float(score)))


def _resolve_line_match(raw_index, source_name, quantity, current_lines, used_indexes):
    try:
        model_index = int(raw_index) if raw_index is not None else None
    except (TypeError, ValueError):
        model_index = None
    if model_index is not None and 0 <= model_index < len(current_lines) and model_index not in used_indexes:
        score = _match_score(source_name, quantity, current_lines[model_index])
        # In batched documents the model may return an index local to its
        # fragment. Trust it only when the actual name/quantity also agree.
        if score >= .58:
            return model_index, max(score, 0.72), "Совпадение по названию и позиции"
    candidates = [
        (_match_score(source_name, quantity, line), index)
        for index, line in enumerate(current_lines)
        if index not in used_indexes
    ]
    if not candidates:
        return None, 0, "Подходящей строки НМЦК нет"
    score, index = max(candidates)
    if score >= 0.66:
        return index, score, "Совпадение по названию и количеству"
    return None, score, "Нужно проверить соответствие вручную"


def _technical_source_chunks(source, max_chars=5200):
    """Split a large table without separating characteristics of one product."""
    if len(source) <= max_chars:
        return [source]
    lines = [line.strip() for line in source.splitlines() if line.strip()]
    product_lines = [line for line in lines if re.match(r"^\d{1,3}\s*[.|)]?\s*\|", line)]
    if len(product_lines) < 4:
        return [source]
    shared_lines = [line for line in lines if line not in product_lines]
    shared = "\n".join(shared_lines)[:1800].strip()

    # DOCX specifications commonly repeat the product number and name on every
    # characteristic row. Splitting those physical rows independently gives
    # the model only half a product and produces duplicate/contradictory items.
    groups = []
    for line in product_lines:
        columns = [value.strip() for value in line.split("|")]
        row_number = re.sub(r"\D+", "", columns[0]) if columns else ""
        product_name = _normalized_item_name(columns[1]) if len(columns) > 1 else ""
        key = (row_number, product_name) if row_number or product_name else (line, "")
        if groups and groups[-1][0] == key:
            groups[-1][1].append(line)
        else:
            groups.append((key, [line]))

    chunks, current_groups = [], []
    for _, group_lines in groups:
        candidate_lines = [line for group in current_groups for line in group] + group_lines
        candidate = "\n".join(([shared] if shared else []) + candidate_lines)
        if current_groups and len(candidate) > max_chars:
            current_lines = [line for group in current_groups for line in group]
            chunks.append("\n".join(([shared] if shared else []) + current_lines))
            current_groups = [group_lines]
        else:
            current_groups.append(group_lines)
    if current_groups:
        current_lines = [line for group in current_groups for line in group]
        chunks.append("\n".join(([shared] if shared else []) + current_lines))
    return chunks or [source]


def _gap_is_covered(gap, requirements):
    """Return whether an extracted requirement already answers a model gap."""
    normalized_gap = _normalized_item_name(gap)
    if not normalized_gap:
        return False
    requirement_texts = [
        _normalized_item_name(f"{value.get('label', '')} {value.get('value', '')}")
        for value in requirements
    ]
    category_roots = (
        ("материал", "сырь", "бумаг", "картон", "пластик", "ткан"),
        ("размер", "ширин", "высот", "длин", "диаметр", "толщин", "формат"),
        ("нанесен", "печат", "лак", "тиснен", "вышив", "гравир"),
        ("упаков", "фасов"),
        ("цвет", "красоч", "cmyk", "pantone"),
        ("плотност", "грамм"),
    )
    for roots in category_roots:
        if any(root in normalized_gap for root in roots):
            return any(any(root in text for root in roots) for text in requirement_texts)
    gap_tokens = {token for token in normalized_gap.split() if len(token) >= 5}
    return any(gap_tokens & set(text.split()) for text in requirement_texts)


def _merge_technical_items(raw_items):
    """Coalesce partial model answers for the same source product."""
    merged = []
    positions = {}
    for raw in raw_items if isinstance(raw_items, list) else []:
        if not isinstance(raw, dict):
            continue
        source_name = _cell_text(raw.get("source_name"))[:500]
        quantity = _cell_text(raw.get("quantity"))[:50]
        key = (_normalized_item_name(source_name), _normalized_item_name(quantity))
        if not key[0]:
            key = (f"line:{raw.get('line_index')}", key[1])
        if key not in positions:
            positions[key] = len(merged)
            merged.append({
                **raw,
                "source_name": source_name,
                "quantity": quantity,
                "requirements": [],
                "missing": [],
                "questions": [],
            })
        target = merged[positions[key]]
        if len(source_name) > len(_cell_text(target.get("source_name"))):
            target["source_name"] = source_name
        if not _cell_text(target.get("quantity")) and quantity:
            target["quantity"] = quantity
        try:
            if float(raw.get("confidence") or 0) > float(target.get("confidence") or 0):
                target["confidence"] = raw.get("confidence")
        except (TypeError, ValueError):
            pass

        known_requirements = {
            (_normalized_item_name(value.get("label")), _normalized_item_name(value.get("value")))
            for value in target["requirements"] if isinstance(value, dict)
        }
        for requirement in _requirement_list(raw.get("requirements")):
            requirement_key = (_normalized_item_name(requirement["label"]), _normalized_item_name(requirement["value"]))
            if requirement_key not in known_requirements:
                known_requirements.add(requirement_key)
                target["requirements"].append(requirement)
        for field, limit in (("missing", 12), ("questions", 4)):
            known_values = {_normalized_item_name(value) for value in target[field]}
            for value in _short_text_list(raw.get(field), limit=limit):
                normalized = _normalized_item_name(value)
                if normalized and normalized not in known_values:
                    known_values.add(normalized)
                    target[field].append(value)

    for target in merged:
        requirements = target.get("requirements", [])
        target["missing"] = [value for value in target.get("missing", []) if not _gap_is_covered(value, requirements)]
        target["questions"] = [value for value in target.get("questions", []) if not _gap_is_covered(value, requirements)][:4]
    return merged


def analyze_tender_requirements(upload, current_lines):
    suffix = Path(upload.name).suffix.lower()
    package = _document_package(upload)
    if package:
        source = package["text"][:AI_MAX_SOURCE_CHARS]
        truncated = len(package["text"]) > AI_MAX_SOURCE_CHARS
    else:
        source, truncated = extract_tender_source(upload)
    scan_ocr = _requires_visual_recognition(upload, source)
    docx_images = package["images"] if package and suffix in {".docx", ".docm"} and not package["components"].get("tables") and not package["components"].get("embedded_spreadsheets") else []
    mixed_pdf_pages = [] if scan_ocr else _pdf_unreadable_pages(upload)
    visual_mode = scan_ocr or bool(docx_images) or bool(mixed_pdf_pages)
    compact_lines = [
        {"line_index": index, "name": _cell_text(line.get("name"))[:500], "quantity": _cell_text(line.get("quantity"))[:50]}
        for index, line in enumerate(current_lines[:100]) if isinstance(line, dict) and _cell_text(line.get("name"))
    ]
    schema = '{"document_summary":"кратко","global_requirements":[{"label":"Общий цвет","value":"чёрный","source":"стр. 2"}],"items":[{"line_index":0,"source_name":"товар из ТЗ","quantity":10,"requirements":[{"label":"Материал","value":"пластик","source":"таблица 1"}],"missing":["реально отсутствующий параметр"],"questions":["один важный вопрос"],"confidence":0.9}],"warnings":[]}'
    def build_prompt(source_part):
        return f"""Проанализируй ООЗ или техническое задание для будущего расчёта заказа.
Извлеки только значимые технические требования к каждой товарной позиции: параметры, которые влияют на выбор товара, материала, технологии, себестоимость или проверяемое соответствие. Все числовые технические характеристики значимы: размеры, количество, плотность, диапазоны, допуски, сроки службы и иные значения с единицами измерения сохраняй всегда.
Для текстовых требований сохраняй вид продукции, материал, цвет, печать или нанесение, постобработку, комплектность и упаковку, только если они помогают выбрать товар или маршрут производства. Не включай общие декларации вроде «продукция новая» и качественные формулировки вроде «изображение стойкое», если рядом нет измеримого критерия, стандарта, метода испытания или явного влияния на технологию.
Не извлекай условия поставки, оплаты, приёмки, исполнения контракта, адреса и общие календарные сроки: они будут анализироваться отдельно. Срок указывай только когда он является техническим свойством производства конкретного изделия и влияет на выбор технологии.
Сопоставь требования с уже имеющимися строками по названию, смыслу и количеству. line_index должен быть индексом подходящей строки, а confidence — уверенностью именно в этом сопоставлении от 0 до 1. Всегда возвращай source_name, quantity и confidence. Если подходящей строки нет, верни null и сохрани исходное название и количество, чтобы позицию можно было создать.
quantity переноси только из документа; количество из ТЕКУЩИХ СТРОК используй лишь для проверки сопоставления.
Возвращай только позиции, которые действительно присутствуют в переданном фрагменте ДОКУМЕНТА. Не копируй остальные ТЕКУЩИЕ СТРОКИ и не создавай для них элементы без требований.
Не рассчитывай цены и себестоимость. Не выдумывай отсутствующие сведения.
Сначала внимательно прочитай все вложенные таблицы характеристик. Не превращай найденные требования в вопросы.
Числа, единицы измерения и ограничения переписывай без вычислений и сокращений. Критично сохраняй частицы «не», «не менее», «не более», «от» и «до»: их потеря меняет смысл требования. Перед ответом повторно сверь каждое числовое ограничение с документом или изображением.
В missing перечисли только параметры, которых действительно нет во всём документе и без которых нельзя выбрать технологию или посчитать себестоимость. В questions — не более 4 коротких вопросов менеджеру только по этим критическим пробелам. Если данных достаточно для предварительного расчёта, верни пустые массивы.
В source укажи страницу, раздел или таблицу, если это можно определить.
В global_requirements помещай только общие для всех позиций технические характеристики. Не помещай туда закупочные и договорные условия.
Верни только JSON строго такого вида: {schema}

ТЕКУЩИЕ СТРОКИ:
{json.dumps(compact_lines, ensure_ascii=False)}

ДОКУМЕНТ:
{source_part or 'Перед тобой страницы сканированного документа в исходном порядке.'}"""

    source_chunks = [""] if scan_ocr else _technical_source_chunks(source)
    prompts = [build_prompt(value) for value in source_chunks]
    if scan_ocr:
        responses = _visual_gateway_responses(prompts[0], upload, max_tokens=7000)
    elif docx_images:
        responses = [_ai_gateway_json(prompts[0], max_tokens=7000, image_data_urls=docx_images)]
    elif mixed_pdf_pages:
        responses = []
        for start in range(0, len(mixed_pdf_pages), AI_MAX_SCAN_PAGES):
            page_batch = mixed_pdf_pages[start:start + AI_MAX_SCAN_PAGES]
            batch_prompt = f"{prompts[0]}\n\nДОПОЛНИТЕЛЬНО РАСПОЗНАЮТСЯ СТРАНИЦЫ: {', '.join(str(value + 1) for value in page_batch)}."
            responses.append(_ai_gateway_json(batch_prompt, max_tokens=7000, image_data_urls=_scan_pdf_selected_images(upload, page_batch)))
    elif len(prompts) == 1:
        responses = [_ai_gateway_json(prompts[0], max_tokens=7000)]
    else:
        # The work is network-bound. Small parallel requests stay within the
        # hosting request timeout and use little additional application RAM.
        with ThreadPoolExecutor(max_workers=min(3, len(prompts))) as executor:
            responses = list(executor.map(lambda value: _ai_gateway_json(value, max_tokens=4200), prompts))
    result = {"items": [], "global_requirements": [], "warnings": []}
    usage = {"prompt_tokens": 0, "completion_tokens": 0}
    summaries = []
    for partial, partial_usage in responses:
        if _cell_text(partial.get("document_summary")):
            summaries.append(_cell_text(partial.get("document_summary")))
        result["items"].extend(partial.get("items", []) if isinstance(partial.get("items"), list) else [])
        result["global_requirements"].extend(partial.get("global_requirements", []) if isinstance(partial.get("global_requirements"), list) else [])
        result["warnings"].extend(partial.get("warnings", []) if isinstance(partial.get("warnings"), list) else [])
        usage["prompt_tokens"] += partial_usage.get("prompt_tokens", 0) or 0
        usage["completion_tokens"] += partial_usage.get("completion_tokens", 0) or 0
    result["document_summary"] = summaries[0] if summaries else ""
    items = []
    used_indexes = set()
    for raw in _merge_technical_items(result.get("items", [])):
        if not isinstance(raw, dict):
            continue
        source_name = _cell_text(raw.get("source_name"))[:500]
        quantity = _cell_text(raw.get("quantity"))[:50]
        requirements = _requirement_list(raw.get("requirements"))
        missing = _short_text_list(raw.get("missing"))
        questions = _short_text_list(raw.get("questions"), limit=4)
        # Batched models sometimes echo context lines that were not present in
        # the current document fragment. An empty echo must not reserve a match
        # or appear as a new technical position.
        if not requirements and not missing and not questions:
            continue
        line_index, fallback_confidence, match_reason = _resolve_line_match(raw.get("line_index"), source_name, quantity, current_lines, used_indexes)
        try:
            model_confidence = max(0, min(1, float(raw.get("confidence"))))
        except (TypeError, ValueError):
            model_confidence = 0
        confidence = max(model_confidence, fallback_confidence) if line_index is not None else fallback_confidence
        if line_index is not None:
            used_indexes.add(line_index)
            # The NMCK line is the calculation authority for quantity. The
            # technical model may echo context or misread a dense scanned cell.
            quantity = _cell_text(current_lines[line_index].get("quantity"))[:50] or quantity
        if source_name or requirements or line_index is not None:
            items.append({"line_index": line_index, "source_name": source_name, "quantity": quantity, "requirements": requirements, "missing": missing, "questions": questions, "confidence": confidence, "match_status": "matched" if line_index is not None else "unmatched", "match_reason": match_reason})
    if not items and not result.get("global_requirements"):
        raise TenderAIError("Не удалось найти технические требования к позициям.")
    warnings = _short_text_list(result.get("warnings"))
    if package:
        warnings = [*package["warnings"], *warnings]
    unmatched_nmck_lines = [
        _cell_text(line.get("name"))[:160]
        for index, line in enumerate(current_lines[:100])
        if isinstance(line, dict) and _cell_text(line.get("name")) and index not in used_indexes
    ]
    if unmatched_nmck_lines:
        unmatched_count = len(unmatched_nmck_lines)
        count_mod_100 = unmatched_count % 100
        count_mod_10 = unmatched_count % 10
        line_word = (
            "строку" if count_mod_10 == 1 and count_mod_100 != 11
            else "строки" if 2 <= count_mod_10 <= 4 and not 12 <= count_mod_100 <= 14
            else "строк"
        )
        preview = "; ".join(unmatched_nmck_lines[:5])
        remainder = unmatched_count - 5
        suffix = f"; и ещё {remainder}" if remainder > 0 else ""
        warnings.append(
            f"ООЗ/ТЗ не покрывает {unmatched_count} "
            f"{line_word} НМЦК: {preview}{suffix}."
        )
    if truncated:
        warnings.append("Документ был слишком большим: обработана основная часть содержимого.")
    if visual_mode:
        warnings.insert(0, "ТЗ распознано по изображению. Проверьте извлечённые требования.")
    global_requirements = []
    seen_global = set()
    supply_condition_markers = (
        "срок постав", "достав", "приемк", "приёмк", "оплат", "контракт",
        "место постав", "адрес постав", "обеспечен", "штраф", "гарантийн",
    )
    for requirement in _requirement_list(result.get("global_requirements")):
        searchable = f"{requirement['label']} {requirement['value']}".lower().replace("ё", "е")
        if any(marker.replace("ё", "е") in searchable for marker in supply_condition_markers):
            continue
        key = (requirement["label"].lower(), requirement["value"].lower())
        if key not in seen_global:
            seen_global.add(key)
            global_requirements.append(requirement)
    return {
        "document_summary": _cell_text(result.get("document_summary"))[:1000],
        "global_requirements": global_requirements,
        "items": items,
        "warnings": warnings,
        "scan_ocr": visual_mode,
        "processing_mode": "visual" if visual_mode else ("embedded" if package and package["components"]["embedded_spreadsheets"] else "text"),
        "components": package["components"] if package else ({"unreadable_pdf_pages": [value + 1 for value in mixed_pdf_pages]} if mixed_pdf_pages else {}),
        "usage": {"prompt_tokens": usage.get("prompt_tokens", 0), "completion_tokens": usage.get("completion_tokens", 0)},
    }


SHEET_FORMATS_MM = {
    "A5": (148, 210),
    "A4": (210, 297),
    "A3": (297, 420),
    "SRA3": (320, 450),
}


def _sheet_format_from_name(value):
    value = _cell_text(value).upper().replace("×", "X")
    for name in ("SRA3", "A5", "A4", "A3"):
        if re.search(rf"(?<![A-ZА-Я0-9]){name}(?![A-ZА-Я0-9])", value):
            return name
    return ""


def _paper_grammage(value):
    values = [int(number) for number in re.findall(r"(?<!\d)(\d{2,3})(?:\s*(?:Г|ГР|G)(?:/М[²2])?)?", _cell_text(value).upper())]
    plausible = [number for number in values if 60 <= number <= 500]
    return plausible[-1] if plausible else None


def calculate_sheet_imposition(piece_width, piece_height, stock_width, stock_height, bleed=0):
    piece_width = float(piece_width) + float(bleed) * 2
    piece_height = float(piece_height) + float(bleed) * 2
    if min(piece_width, piece_height, stock_width, stock_height) <= 0:
        return {"ups": 0, "rotation": False}
    straight = math.floor(float(stock_width) / piece_width) * math.floor(float(stock_height) / piece_height)
    rotated = math.floor(float(stock_width) / piece_height) * math.floor(float(stock_height) / piece_width)
    return {"ups": max(straight, rotated), "rotation": rotated > straight}


def _paper_matches_material(item_name, material_query):
    name = _normalized_item_name(item_name)
    query = _normalized_item_name(material_query)
    if not query:
        return True
    if "офсет" in query:
        return any(value in name for value in ("обыч", "maestro", "немел", "офсет"))
    if "мелован" in query or "мелов" in query:
        return "мелов" in name
    if "самокле" in query:
        return any(value in name for value in ("ритрама", "ritrama", "oracal", "с к"))
    if "дизайнер" in query or "металлиз" in query:
        return any(value in name for value in ("majestic", "curious", "touche", "лен", "крафт", "калька"))
    return True


def _paper_candidates(component, product_quantity, paper_items, waste_percent=3):
    width = float(component.get("finished_width_mm") or 0)
    height = float(component.get("finished_height_mm") or 0)
    units_per_product = float(component.get("units_per_product") or 0)
    grammage = component.get("grammage_gsm")
    try:
        grammage = int(float(grammage)) if grammage not in (None, "") else None
    except (TypeError, ValueError):
        grammage = None
    if not width or not height or not units_per_product:
        return []
    total_pieces = math.ceil(float(product_quantity) * units_per_product)
    candidates, covered_formats = [], set()
    for item in paper_items:
        format_name = _sheet_format_from_name(item.name)
        if format_name not in SHEET_FORMATS_MM:
            continue
        item_grammage = _paper_grammage(item.name)
        if not _paper_matches_material(item.name, component.get("material_query")):
            continue
        if grammage and item_grammage and not grammage <= item_grammage <= grammage + 50:
            continue
        if grammage and item_grammage is None:
            continue
        stock_width, stock_height = SHEET_FORMATS_MM[format_name]
        layout = calculate_sheet_imposition(width, height, stock_width, stock_height, component.get("bleed_mm") or 0)
        if not layout["ups"]:
            continue
        base_sheets = math.ceil(total_pieces / layout["ups"])
        sheets = math.ceil(base_sheets * (100 + waste_percent) / 100)
        candidates.append({
            "catalog_item_id": item.pk,
            "name": item.name,
            "format": format_name,
            "stock_width_mm": stock_width,
            "stock_height_mm": stock_height,
            "grammage_gsm": item_grammage,
            "ups": layout["ups"],
            "rotation": layout["rotation"],
            "sheets": sheets,
            "waste_percent": waste_percent,
            "unit_price": str(item.effective_unit_price),
            "material_cost": str(_money(Decimal(sheets) * item.effective_unit_price)),
            "price_missing": False,
        })
        if grammage is None or item_grammage is None or abs(item_grammage - grammage) <= 5:
            covered_formats.add(format_name)
    for format_name in ("A4", "A3", "SRA3"):
        if format_name in covered_formats:
            continue
        stock_width, stock_height = SHEET_FORMATS_MM[format_name]
        layout = calculate_sheet_imposition(width, height, stock_width, stock_height, component.get("bleed_mm") or 0)
        if not layout["ups"]:
            continue
        base_sheets = math.ceil(total_pieces / layout["ups"])
        candidates.append({
            "catalog_item_id": None,
            "name": f"{component.get('material_query') or 'Бумага'} {format_name}",
            "format": format_name,
            "stock_width_mm": stock_width,
            "stock_height_mm": stock_height,
            "grammage_gsm": grammage,
            "ups": layout["ups"],
            "rotation": layout["rotation"],
            "sheets": math.ceil(base_sheets * (100 + waste_percent) / 100),
            "waste_percent": waste_percent,
            "unit_price": None,
            "material_cost": None,
            "price_missing": True,
        })
    return sorted(candidates, key=lambda value: (value["material_cost"] is None, Decimal(value["material_cost"] or "999999999")))


def _production_evidence_text(line):
    """Collect facts already extracted from the position and its TZ."""
    values = [line.get("name"), line.get("quantity")]
    info = line.get("requirements") if isinstance(line.get("requirements"), dict) else {}
    for item in info.get("requirements", []) if isinstance(info.get("requirements"), list) else []:
        if isinstance(item, dict):
            values.extend((item.get("label"), item.get("value"), item.get("source")))
        else:
            values.append(item)
    return _cell_text(" ".join(str(value) for value in values if value is not None)).lower()


def _question_is_answered_by_evidence(question, evidence):
    """Detect application questions already answered by explicit TZ facts."""
    question = _cell_text(question).lower()
    if not question or not evidence:
        return False
    subjects = (
        ("герб",), ("надпис", "текст"), ("логотип", "лого"),
        ("изображен", "рисунок"),
    )
    operations = (
        "тиснен", "конгрев", "печать", "напечат", "уф-печат", "уф печат",
        "шелкограф", "вышив", "гравиров", "металлизац", "фольг",
    )
    if not any(value in question for value in ("нанес", "метод", "способ", "печат", "офсет", "цифров")):
        return False
    mentioned_subjects = [aliases for aliases in subjects if any(alias in question for alias in aliases)]
    if not mentioned_subjects:
        return False
    for aliases in mentioned_subjects:
        subject_answered = False
        for alias in aliases:
            for match in re.finditer(alias, evidence):
                context = evidence[max(0, match.start() - 280):match.end() + 280]
                if any(operation in context for operation in operations):
                    subject_answered = True
                    break
            if subject_answered:
                break
        if not subject_answered:
            return False
    return True


def _unanswered_production_questions(raw_questions, line, features=None, limit=2):
    evidence = _cell_text(f"{_production_evidence_text(line)} {' '.join(features or [])}").lower()
    result = []
    for raw in raw_questions if isinstance(raw_questions, list) else []:
        question = raw.get("question") if isinstance(raw, dict) else raw
        question = _cell_text(question)
        if question and not _question_is_answered_by_evidence(question, evidence) and question not in result:
            result.append(question)
        if len(result) >= limit:
            break
    return result


def classify_production_type(line):
    from .models import ProcessDefinition, ProductionTrainingExample, ProductionType

    production_types = list(ProductionType.objects.filter(is_active=True))
    # Ограничиваем контекст, чтобы обучение не раздувало время и стоимость
    # каждого запроса на малом тарифе приложения.
    examples = list(ProductionTrainingExample.objects.filter(is_active=True).select_related("production_type")[:25])
    type_payload = [{"code": value.code, "name": value.name, "description": value.description} for value in production_types]
    example_payload = [
        {"id": value.pk, "name": value.position_name, "type": value.production_type.code, "features": value.features, "routes": value.routes}
        for value in examples
    ]
    process_definitions = list(ProcessDefinition.objects.filter(is_active=True))
    process_payload = [{"name": value.name, "role": value.role, "description": value.description} for value in process_definitions]
    schema = '{"suggested_type":"digital_sheet","confidence":0.45,"reason":"почему","features":["существенный признак"],"alternatives":[{"type":"offset_print","reason":"почему возможно"}],"routes":[{"name":"Под ключ","reason":"почему","processes":[{"role":"supply|production|completion","name":"Процесс","reason":"зачем"}]}],"matched_example_ids":[1],"questions":[{"question":"один вопрос","missing_fact":"какого факта нет","why_it_changes_route":"что изменит ответ"}]}'
    prompt = f"""Определи технологический тип и предложи 1–3 возможных маршрута из процессов. Пока НЕ ищи конкретного поставщика, НЕ выбирай калькулятор и НЕ считай цену.
Используй только типы из справочника. Выдели фактические признаки из названия и ТЗ: тираж, формат, конструкцию, способ печати, материал и обязательные операции. Не считай само упоминание бумаги или печати доказательством цифровой листовой печати.
Маршрут — последовательность процессов, приводящая к готовому изделию. Процесс может относиться к снабжению, производству или завершению/логистике. Один процесс «изготовление под ключ» допустим, если он реалистичен. Альтернативный маршрут может разделять снабжение и производство, например готовый бланк + нанесение. Не добавляй процессы ради количества.
Используй известные процессы из справочника, но если необходимого процесса нет — предложи ясное новое название. DTF и УФ-DTF относятся к возможностям цифровой типографии; прямая УФ-печать на станке — отдельный процесс.
Если данных недостаточно или подходят несколько технологий, честно снизь confidence и задай не более двух вопросов, ответ на которые действительно изменит классификацию.
Перед каждым вопросом проверь все извлечённые requirements. Нельзя спрашивать способ нанесения, материал, формат или операцию, если они уже прямо указаны в ТЗ. Несколько операций могут сосуществовать: конгрев герба и тиснение надписи не являются офсетной или цифровой печатью. Вопрос допустим только о конкретном отсутствующем факте; укажи его в missing_fact и объясни в why_it_changes_route, какой выбор маршрута зависит от ответа. Если явного пробела нет, questions должен быть пустым.
matched_example_ids указывай только для действительно похожих подтверждённых примеров. Не завышай уверенность: без близкого подтверждённого примера значение не должно превышать 0.55.
Верни только JSON: {schema}

ПОЗИЦИЯ:
{json.dumps(line, ensure_ascii=False)}

СПРАВОЧНИК ТИПОВ:
{json.dumps(type_payload, ensure_ascii=False)}

ИЗВЕСТНЫЕ ПРОЦЕССЫ:
{json.dumps(process_payload, ensure_ascii=False)}

ПОДТВЕРЖДЁННЫЕ АДМИНИСТРАТОРОМ ПРИМЕРЫ:
{json.dumps(example_payload, ensure_ascii=False)}"""
    result, usage = _ai_gateway_json(prompt, max_tokens=1800)
    valid_codes = {value.code for value in production_types}
    suggested = result.get("suggested_type") if result.get("suggested_type") in valid_codes else "other"
    features = _short_text_list(result.get("features"), limit=10)
    questions = _unanswered_production_questions(result.get("questions"), line, features=features, limit=2)
    valid_example_ids = {value.pk for value in examples}
    matched_ids = []
    for value in result.get("matched_example_ids", []) if isinstance(result.get("matched_example_ids"), list) else []:
        try:
            value = int(value)
        except (TypeError, ValueError):
            continue
        if value in valid_example_ids:
            matched_ids.append(value)
    try:
        model_confidence = max(0, min(1, float(result.get("confidence", 0))))
    except (TypeError, ValueError):
        model_confidence = 0
    # Процент является программной оценкой доказательств, а не уверенностью,
    # которую модель может объявить сама.
    confidence_cap = .45 if not matched_ids else (.72 if len(matched_ids) == 1 else .82)
    if questions:
        confidence_cap = min(confidence_cap, .50)
    confidence = min(model_confidence, confidence_cap)
    alternatives = []
    for raw in result.get("alternatives", [])[:3]:
        if isinstance(raw, dict) and raw.get("type") in valid_codes and raw.get("type") != suggested:
            alternatives.append({"type": raw["type"], "reason": _cell_text(raw.get("reason"))[:300]})
    routes = []
    for route_index, raw_route in enumerate(result.get("routes", [])[:3]):
        if not isinstance(raw_route, dict):
            continue
        processes = []
        for raw_process in raw_route.get("processes", [])[:10]:
            if not isinstance(raw_process, dict):
                continue
            role = raw_process.get("role") if raw_process.get("role") in {"supply", "production", "completion"} else "production"
            name = _cell_text(raw_process.get("name"))[:200]
            if name:
                processes.append({"role": role, "name": name, "reason": _cell_text(raw_process.get("reason"))[:300]})
        if processes:
            routes.append({"name": _cell_text(raw_route.get("name"))[:120] or f"Маршрут {route_index + 1}", "reason": _cell_text(raw_route.get("reason"))[:300], "processes": processes})
    return {
        "stage": "production_classification",
        "suggested_type": suggested,
        "confidence": confidence,
        "reason": _cell_text(result.get("reason"))[:700],
        "features": features,
        "alternatives": alternatives,
        "routes": routes,
        "matched_example_ids": matched_ids,
        "questions": questions,
        "production_types": type_payload,
        "process_definitions": process_payload,
        "training_examples_count": len(examples),
        "manager_answers": line.get("manager_answers") or line.get("requirements", {}).get("manager_answers") or {},
        "usage": {"prompt_tokens": usage.get("prompt_tokens", 0), "completion_tokens": usage.get("completion_tokens", 0)},
    }


def _embeddings_enabled():
    return os.getenv("TIMEWEB_EMBEDDINGS_ENABLED", "0") == "1"


def _embedding_model():
    return os.getenv("TIMEWEB_EMBEDDING_MODEL", "openai/text-embedding-3-small").strip()


def _embedding_vector(text, model=None):
    api_key = os.getenv("TIMEWEB_AI_API_KEY", "").strip()
    base_url = os.getenv("TIMEWEB_AI_BASE_URL", "https://api.timeweb.ai/v1").rstrip("/")
    model = model or _embedding_model()
    if not api_key:
        raise TenderAIError("AI Gateway ещё не настроен для смыслового поиска.")
    payload = json.dumps({"model": model, "input": _cell_text(text)[:12_000]}, ensure_ascii=False).encode("utf-8")
    request = Request(
        f"{base_url}/embeddings", data=payload,
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}, method="POST",
    )
    try:
        with urlopen(request, timeout=45) as response:
            result = json.loads(response.read().decode("utf-8"))
        vector = result["data"][0]["embedding"]
        if not isinstance(vector, list) or not vector:
            raise ValueError
        return [float(value) for value in vector]
    except (HTTPError, URLError, TimeoutError, ConnectionError, OSError, KeyError, IndexError, TypeError, ValueError, json.JSONDecodeError) as exc:
        raise TenderAIError("Не удалось построить смысловой индекс.") from exc


def _training_example_embedding_text(example):
    return json.dumps({
        "position": example.position_name,
        "requirements": example.requirements,
        "features": example.features,
        "routes": example.routes,
    }, ensure_ascii=False)


def refresh_training_example_embedding(example):
    if not _embeddings_enabled():
        return False
    try:
        example.embedding = _embedding_vector(_training_example_embedding_text(example))
    except TenderAIError:
        logger.warning("Could not refresh training embedding for example %s", example.pk)
        return False
    example.embedding_model = _embedding_model()
    example.embedding_updated_at = timezone.now()
    example.save(update_fields=["embedding", "embedding_model", "embedding_updated_at"])
    return True


def _cosine_similarity(left, right):
    if not left or not right or len(left) != len(right):
        return 0
    denominator = math.sqrt(sum(value * value for value in left)) * math.sqrt(sum(value * value for value in right))
    return sum(a * b for a, b in zip(left, right)) / denominator if denominator else 0


def _training_examples_for_line(line, limit=12):
    from .models import ProductionTrainingExample

    target = _normalized_item_name(line.get("name"))
    examples = list(ProductionTrainingExample.objects.filter(is_active=True).select_related("production_type")[:200])
    semantic_scores = {}
    model = _embedding_model()
    embedded = [value for value in examples if value.embedding_model == model and isinstance(value.embedding, list) and value.embedding]
    if _embeddings_enabled() and embedded:
        source = json.dumps(line, ensure_ascii=False, sort_keys=True)
        cache_key = f"training-query-embedding:{model}:{hashlib.sha256(source.encode('utf-8')).hexdigest()}"
        query_embedding = cache.get(cache_key)
        if query_embedding is None:
            try:
                query_embedding = _embedding_vector(source, model=model)
                cache.set(cache_key, query_embedding, 24 * 60 * 60)
            except TenderAIError:
                query_embedding = []
        semantic_scores = {value.pk: max(0, _cosine_similarity(query_embedding, value.embedding)) for value in embedded}
    ranked = []
    for example in examples:
        source = _normalized_item_name(example.position_name)
        lexical_score = SequenceMatcher(None, target, source).ratio() if target and source else 0
        target_tokens, source_tokens = set(target.split()), set(source.split())
        if target_tokens and source_tokens:
            lexical_score = max(lexical_score, len(target_tokens & source_tokens) / len(target_tokens | source_tokens))
        semantic_score = semantic_scores.get(example.pk)
        score = semantic_score * .8 + lexical_score * .2 if semantic_score is not None else lexical_score
        ranked.append((score, lexical_score, semantic_score, example))
    ranked.sort(key=lambda value: (value[0], value[3].created_at), reverse=True)
    return [value for score, lexical, semantic, value in ranked[:limit] if lexical >= .12 or (semantic or 0) >= .2]


def _knowledge_sources_for_line(line, limit=4):
    from .models import TenderKnowledgeSource

    source_text = json.dumps(line, ensure_ascii=False).lower().replace("ё", "е")
    ignored = {"который", "изделие", "позиция", "количество", "требования", "материал", "печать", "цвет", "размер"}
    tokens = {value for value in re.findall(r"[a-zа-я0-9]{4,}", source_text) if value not in ignored}
    ranked = []
    for source in TenderKnowledgeSource.objects.filter(is_active=True)[:100]:
        title = f"{source.supplier_name} {source.title}".lower().replace("ё", "е")
        content = source.content_summary.lower().replace("ё", "е")
        title_hits = sum(1 for value in tokens if value in title)
        content_hits = sum(1 for value in tokens if value in content)
        score = title_hits * 3 + content_hits
        if score:
            ranked.append((score, source.updated_at, source))
    ranked.sort(key=lambda value: (value[0], value[1]), reverse=True)
    return [{
        "id": source.pk,
        "supplier": source.supplier_name,
        "title": source.title,
        "type": source.source_type,
        "url": source.url,
        "data": source.content_summary[:1600],
    } for _, _, source in ranked[:limit]]


def _canonical_process_name(value):
    text = _cell_text(value)
    lowered = text.lower().replace("ё", "е")
    turnkey = "под ключ" in lowered
    rules = [
        (("закуп", "материал"), "Закупка материала"),
        (("постав", "материал"), "Закупка материала"),
        (("постав", "бумаг"), "Закупка материала"),
        (("закуп", "бумаг"), "Закупка материала"),
        (("готов", "издел"), "Закупка готового изделия"),
        (("готов", "бланк"), "Закупка готового изделия"),
        (("универсаль", "типограф"), "Универсальная типография"),
        (("цифров", "типограф"), "Цифровая типография"),
        (("офсет", "типограф"), "Офсетная типография"),
        (("нанесен",), "Нанесение"),
    ]
    for markers, name in rules:
        if all(marker in lowered for marker in markers):
            if turnkey and name in {"Универсальная типография", "Цифровая типография", "Офсетная типография", "Нанесение"}:
                return f"{name} под ключ"
            return name
    return text[:80] or "Процесс не определён"


def _normalize_route_processes(route):
    raw_processes = route.get("processes") if isinstance(route.get("processes"), list) else []
    if not raw_processes:
        raw_processes = [{"name": value, "details": []} for value in _short_text_list(route.get("steps"), limit=6)]
    processes = []
    for value in raw_processes[:6]:
        if isinstance(value, dict):
            name = _canonical_process_name(value.get("name"))
            details = _short_text_list(value.get("details"), limit=8)
        else:
            name, details = _canonical_process_name(value), []
        if not processes or processes[-1]["name"].lower() != name.lower():
            processes.append({"name": name, "details": details})
        else:
            processes[-1]["details"].extend(item for item in details if item not in processes[-1]["details"])
    processes = processes or [{"name": "Маршрут пока не определён", "details": []}]
    # A single turnkey contractor is a production process, not procurement.
    # Correct this contradiction deterministically even if the model labelled
    # the route as material purchase while describing full manufacturing.
    if len(processes) == 1 and processes[0]["name"] == "Закупка материала":
        evidence = " ".join([
            _cell_text(route.get("reason")),
            *processes[0].get("details", []),
        ]).lower().replace("ё", "е")
        if "под ключ" in evidence or ("изготов" in evidence and any(word in evidence for word in ("типограф", "производств", "подрядчик"))):
            if "цифров" in evidence:
                processes[0]["name"] = "Цифровая типография под ключ"
            elif "офсет" in evidence:
                processes[0]["name"] = "Офсетная типография под ключ"
            elif "универсаль" in evidence or "типограф" in evidence:
                processes[0]["name"] = "Универсальная типография под ключ"
            else:
                processes[0]["name"] = "Производство под ключ"
    return processes


def _decimal_input(inputs, key, default=None):
    try:
        value = Decimal(str(inputs.get(key, default)).replace(",", "."))
    except (InvalidOperation, TypeError, ValueError, AttributeError):
        return None
    return value


def _evaluate_cost_recipe(recipe, quantity):
    if not isinstance(recipe, dict):
        return None, []
    method = recipe.get("method")
    inputs = recipe.get("inputs") if isinstance(recipe.get("inputs"), dict) else {}
    total = None
    steps = []
    if method == "sheet_yield":
        unit_price = _decimal_input(inputs, "unit_price")
        units_per_sheet = _decimal_input(inputs, "units_per_sheet")
        waste_percent = _decimal_input(inputs, "waste_percent", 0)
        if unit_price is None or units_per_sheet is None or units_per_sheet <= 0 or waste_percent is None:
            return None, []
        sheets_exact = quantity / units_per_sheet * (Decimal("1") + waste_percent / Decimal("100"))
        sheets = sheets_exact.to_integral_value(rounding=ROUND_CEILING)
        total = _money(sheets * unit_price)
        steps = [
            f"Тираж: {_decimal_text(quantity)} шт.; выход: {_decimal_text(units_per_sheet)} шт. с исходного листа",
            f"С учётом отходов {_decimal_text(waste_percent)}%: {_decimal_text(sheets_exact)} → {sheets} листов",
            f"{sheets} листов × {_money(unit_price)} ₽ = {total} ₽",
        ]
    elif method == "unit_rate":
        unit_rate = _decimal_input(inputs, "unit_rate")
        if unit_rate is None:
            return None, []
        total = _money(quantity * unit_rate)
        steps = [f"{_decimal_text(quantity)} шт. × {_money(unit_rate)} ₽/шт. = {total} ₽"]
    elif method == "fixed":
        fixed_amount = _decimal_input(inputs, "fixed_amount")
        if fixed_amount is None:
            return None, []
        total = _money(fixed_amount)
        steps = [f"Фиксированная стоимость на тираж: {total} ₽"]
    elif method == "history_scaled":
        base_total = _decimal_input(inputs, "base_total")
        base_quantity = _decimal_input(inputs, "base_quantity")
        if base_total is None or base_quantity is None or base_quantity <= 0:
            return None, []
        unit_rate = base_total / base_quantity
        total = _money(unit_rate * quantity)
        steps = [
            f"Исходный кейс: {_money(base_total)} ₽ за {_decimal_text(base_quantity)} шт. = {_money(unit_rate)} ₽/шт.",
            f"Текущий тираж: {_decimal_text(quantity)} шт. × {_money(unit_rate)} ₽/шт. = {total} ₽",
        ]
    if total is None:
        return None, []

    modifiers = recipe.get("modifiers") if isinstance(recipe.get("modifiers"), list) else []
    for modifier in modifiers[:10]:
        if not isinstance(modifier, dict):
            return None, []
        modifier_type = _cell_text(modifier.get("type")).lower()
        value = _decimal_input(modifier, "value")
        if value is None:
            return None, []
        before = total
        if modifier_type == "discount_percent" and Decimal("0") <= value <= Decimal("100"):
            multiplier = Decimal("1") - value / Decimal("100")
            total = _money(before * multiplier)
            steps.append(f"Скидка {_decimal_text(value)}%: {before} ₽ × {multiplier} = {total} ₽")
        elif modifier_type == "markup_percent" and value >= 0:
            multiplier = Decimal("1") + value / Decimal("100")
            total = _money(before * multiplier)
            steps.append(f"Наценка {_decimal_text(value)}%: {before} ₽ × {multiplier} = {total} ₽")
        elif modifier_type == "add_fixed" and value >= 0:
            total = _money(before + value)
            steps.append(f"Дополнительный фиксированный расход: {before} ₽ + {_money(value)} ₽ = {total} ₽")
        elif modifier_type == "subtract_fixed" and Decimal("0") <= value <= before:
            total = _money(before - value)
            steps.append(f"Фиксированная скидка: {before} ₽ − {_money(value)} ₽ = {total} ₽")
        else:
            return None, []
    return total, steps


def _extract_productivity_per_hour(*values):
    text = " ".join(json.dumps(value, ensure_ascii=False) if not isinstance(value, str) else value for value in values if value)
    text = text.lower().replace("ё", "е")
    patterns = (
        r"(\d+(?:[.,]\d+)?)\s*(?:шт\.?|штук\w*|издел\w*|футбол\w*)\s*(?:/|в)\s*(?:1\s*)?час",
        r"за\s*(?:1\s*)?час\D{0,20}(\d+(?:[.,]\d+)?)\s*(?:шт\.?|штук\w*|издел\w*|футбол\w*)",
    )
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            try:
                value = Decimal(match.group(1).replace(",", "."))
            except InvalidOperation:
                continue
            if value > 0:
                return value
    return None


def _apply_psodin_calculation(hypothesis, raw, line, current=None, feedback="", confirmed=None):
    """Replace model arithmetic with the existing PSODIN backend calculator."""
    from calculator.models import CalculatorSettings
    from calculator.services import calculate_sheet_estimate

    raw_calculation = raw.get("psodin_calculation") if isinstance(raw.get("psodin_calculation"), dict) else {}
    current_calculation = current.get("psodin_calculation") if isinstance(current, dict) and isinstance(current.get("psodin_calculation"), dict) else {}
    confirmed_calculation = confirmed if isinstance(confirmed, dict) else {}
    feedback_text = feedback.lower().replace("ё", "е")
    authorized = any(marker in feedback_text for marker in ("psodin", "псодин", "печатный салон №1")) or bool(current_calculation.get("authorized")) or bool(confirmed_calculation.get("authorized"))
    if not authorized:
        return hypothesis

    productivity = _decimal_input(raw_calculation, "productivity_per_hour")
    if productivity is None or productivity <= 0:
        productivity = _decimal_input(current_calculation, "productivity_per_hour")
    if productivity is None or productivity <= 0:
        productivity = _decimal_input(confirmed_calculation, "productivity_per_hour")
    if productivity is None or productivity <= 0:
        productivity = _extract_productivity_per_hour(feedback, current, raw)
    questions = list(hypothesis.get("questions") or [])
    questions = [value for value in questions if "psodin" not in str(value).lower() and "час" not in str(value).lower()]
    if productivity is None or productivity <= 0:
        questions.append("Сколько изделий в час PSODIN выполняет эту работу?")
        hypothesis["questions"] = list(dict.fromkeys(questions))[:3]
        hypothesis["psodin_calculation"] = {"authorized": True, "status": "missing_productivity", "calculator": "sheet"}
        return hypothesis

    try:
        quantity = max(Decimal("1"), Decimal(str(line.get("quantity", 1)).replace(",", ".")))
    except (InvalidOperation, TypeError, ValueError):
        quantity = Decimal("1")
    exact_hours = quantity / productivity
    billed_hours = (exact_hours * Decimal("2")).to_integral_value(rounding=ROUND_CEILING) / Decimal("2")
    settings = CalculatorSettings.objects.get_or_create(pk=1)[0]
    calculated = calculate_sheet_estimate([], quantity, billed_hours, settings, "sheet")
    tariff = raw_calculation.get("tariff") or current_calculation.get("tariff") or confirmed_calculation.get("tariff") or "partner"
    if tariff not in {"standard", "regular", "partner", "urgent"}:
        tariff = "partner"
    tariff_labels = {"standard": "стандартный", "regular": "постоянник", "partner": "контрагент", "urgent": "без очереди"}
    amount = _money(calculated[tariff])
    process_name = _cell_text(raw_calculation.get("process_name"))[:80] or "Работа PSODIN"
    existing_costs = hypothesis.get("costs") if isinstance(hypothesis.get("costs"), list) else []
    costs = [item for item in existing_costs if "psodin" not in f"{item.get('name', '')} {item.get('process_name', '')} {item.get('source', '')}".lower()]
    base_formula = f"{_decimal_text(billed_hours)} ч × {_money(settings.hourly_rate)} ₽/ч × {_decimal_text(settings.time_coefficient)}"
    steps = [
        f"Тираж {_decimal_text(quantity)} шт. ÷ {_decimal_text(productivity)} шт./ч = {_decimal_text(exact_hours)} ч",
        f"Оплачиваемое время с шагом 0,5 ч: {_decimal_text(billed_hours)} ч",
        f"Стандартная цена: {base_formula} = {_money(calculated['standard'])} ₽",
    ]
    if tariff == "regular":
        multiplier = Decimal("1") - settings.regular_discount / Decimal("100")
        steps.append(f"Скидка {settings.regular_discount}% к стоимости: {_money(calculated['standard'])} ₽ × {multiplier} = {amount} ₽")
    elif tariff == "partner":
        multiplier = Decimal("1") - settings.partner_discount / Decimal("100")
        steps.append(f"Скидка {settings.partner_discount}% к стоимости: {_money(calculated['standard'])} ₽ × {multiplier} = {amount} ₽")
    elif tariff == "urgent":
        steps.append(f"Коэффициент срочности {settings.urgency_multiplier}: {_money(calculated['standard'])} ₽ × {settings.urgency_multiplier} = {amount} ₽")
    costs.append({
        "category": "application", "name": "Работа PSODIN", "amount_total": str(amount), "process_name": process_name,
        "source": "Калькулятор PSODIN · Листовая печать", "source_type": "calculator", "source_url": "", "source_date": "",
        "basis": f"{base_formula}; тариф «{tariff_labels[tariff]}»",
        "adaptation": "Трудоёмкость получена из тиража и подтверждённой производительности; цена полностью рассчитана бэкендом.",
        "calculation_steps": steps,
        "recipe": {"method": "psodin_backend", "inputs": {"quantity": str(quantity), "productivity_per_hour": str(productivity), "billed_hours": str(billed_hours), "tariff": tariff}},
        "confirmed": False,
    })
    totals = {"material": Decimal("0"), "application": Decimal("0"), "logistics": Decimal("0")}
    for item in costs:
        category = item.get("category") if item.get("category") in totals else "application"
        try:
            totals[category] += max(Decimal("0"), Decimal(str(item.get("amount_total", 0)).replace(",", ".")))
        except (InvalidOperation, TypeError, ValueError):
            continue
    total = sum(totals.values(), Decimal("0"))
    hypothesis["costs"] = costs
    hypothesis["totals"] = {
        "material_unit": str(_money(totals["material"] / quantity)), "application_unit": str(_money(totals["application"] / quantity)),
        "logistics_unit": str(_money(totals["logistics"] / quantity)), "cost_unit": str(_money(total / quantity)), "cost_total": str(_money(total)),
    }
    hypothesis["questions"] = list(dict.fromkeys(questions))[:3]
    hypothesis["psodin_calculation"] = {
        "authorized": True, "status": "calculated", "calculator": "sheet", "scope": "labour_only", "process_name": process_name,
        "productivity_per_hour": str(productivity), "exact_hours": str(exact_hours), "billed_hours": str(billed_hours), "tariff": tariff,
    }
    return hypothesis


def _normalize_training_hypothesis(raw, line, production_types, matched_ids):
    valid_types = {value.code for value in production_types}
    product_type = raw.get("product_type") if raw.get("product_type") in valid_types else "other"
    route = raw.get("route") if isinstance(raw.get("route"), dict) else {}
    processes = _normalize_route_processes(route)
    process_names = [value["name"] for value in processes]
    try:
        quantity = max(Decimal("1"), Decimal(str(line.get("quantity", 1)).replace(",", ".")))
    except (InvalidOperation, TypeError, ValueError):
        quantity = Decimal("1")
    costs = []
    learning_warnings = []
    requirements_text = json.dumps(line.get("requirements", {}), ensure_ascii=False).lower().replace("ё", "е")
    requirements_contain_price = bool(re.search(r"(?:₽|\bруб\.?\b|цен[аы]|стоимост)", requirements_text))
    totals = {"material": Decimal("0"), "application": Decimal("0"), "logistics": Decimal("0")}
    for item in raw.get("costs", [])[:12] if isinstance(raw.get("costs"), list) else []:
        if not isinstance(item, dict):
            continue
        category = item.get("category") if item.get("category") in totals else "application"
        name = _cell_text(item.get("name"))[:200]
        recipe = item.get("recipe") if isinstance(item.get("recipe"), dict) else {}
        calculated_amount, calculated_steps = _evaluate_cost_recipe(recipe, quantity)
        try:
            amount = max(Decimal("0"), Decimal(str(item.get("amount_total", 0)).replace(",", ".")))
        except (InvalidOperation, TypeError, ValueError):
            amount = Decimal("0")
        if not name or (amount <= 0 and calculated_amount is None):
            continue
        source_type = item.get("source_type") if item.get("source_type") in {"calculator", "catalog", "supplier", "history", "manager"} else "manager"
        source = _cell_text(item.get("source"))[:300]
        manual_unit = _decimal_input(line, {"material": "material_unit", "application": "application_unit", "logistics": "logistics_unit"}[category], 0) or Decimal("0")
        if "тз" in source.lower().replace("ё", "е") and manual_unit > 0 and abs(amount - _money(manual_unit * quantity)) <= Decimal("0.02"):
            source = "Введено администратором в расчёте"
            source_type = "manager"
        elif "тз" in source.lower().replace("ё", "е") and not requirements_contain_price:
            source = "Источник цены не подтверждён"
            source_type = "manager"
            learning_warnings.append(f"Для статьи «{name}» цена ошибочно приписана ТЗ.")
        if amount > 0 and calculated_amount is None and source_type == "manager":
            # A fixed amount explicitly supplied by the administrator is already
            # a complete backend input. Do not depend on the LLM to serialize the
            # equivalent recipe correctly and do not show a false warning.
            recipe = {"method": "fixed", "inputs": {"fixed_amount": str(_money(amount))}}
            calculated_amount, calculated_steps = _evaluate_cost_recipe(recipe, quantity)
        if amount > 0 and calculated_amount is None:
            learning_warnings.append(f"Для статьи «{name}» нет проверяемой серверной формулы.")
        amount = calculated_amount if calculated_amount is not None else _money(amount)
        totals[category] += amount
        raw_steps = _short_text_list(item.get("calculation_steps"), limit=12)
        costs.append({
            "category": category,
            "name": name,
            "amount_total": str(amount),
            "process_name": _canonical_process_name(item.get("process_name") or process_names[min(len(costs), len(process_names) - 1)]),
            "source": source,
            "source_type": source_type,
            "source_url": _cell_text(item.get("source_url"))[:1000],
            "source_date": _cell_text(item.get("source_date"))[:50],
            "basis": _cell_text(item.get("basis"))[:700],
            "adaptation": _cell_text(item.get("adaptation"))[:700],
            "calculation_steps": calculated_steps or raw_steps,
            "recipe": recipe,
            "confirmed": bool(item.get("confirmed")),
        })
    total = sum(totals.values(), Decimal("0"))
    confidence = raw.get("confidence", 0)
    try:
        confidence = max(0, min(1, float(confidence)))
    except (TypeError, ValueError):
        confidence = 0
    if not matched_ids:
        confidence = min(confidence, .55)
    return {
        "stage": "training_dialogue",
        "product_type": product_type,
        "summary": _cell_text(raw.get("summary"))[:700],
        "confidence": confidence,
        "facts": _short_text_list(raw.get("facts"), limit=10),
        "route": {
            "name": " → ".join(process_names),
            "reason": _cell_text(route.get("reason"))[:700],
            "steps": process_names,
            "processes": processes,
            "is_turnkey": len(processes) == 1 and "под ключ" in processes[0]["name"].lower().replace("ё", "е"),
        },
        "costs": costs,
        "totals": {
            "material_unit": str(_money(totals["material"] / quantity)),
            "application_unit": str(_money(totals["application"] / quantity)),
            "logistics_unit": str(_money(totals["logistics"] / quantity)),
            "cost_unit": str(_money(total / quantity)),
            "cost_total": str(_money(total)),
        },
        "questions": _short_text_list(raw.get("questions"), limit=3),
        "assumptions": _short_text_list(raw.get("assumptions"), limit=6),
        "understood_changes": _short_text_list(raw.get("understood_changes"), limit=8),
        "learning_warnings": learning_warnings,
        "matched_example_ids": matched_ids,
    }


def apply_verified_source_quote(hypothesis, line, quote, source_name, source_url="", target=None):
    """Replace an LLM-read web price with the exact row/tier chosen by the backend."""
    from .models import ProductionType

    if not isinstance(hypothesis, dict) or not isinstance(quote, dict) or quote.get("confidence") != "exact":
        return hypothesis
    try:
        quantity = max(Decimal("1"), Decimal(str(line.get("quantity", 1)).replace(",", ".")))
        unit_price = Decimal(str(quote.get("unit_price", "")).replace(",", "."))
    except (InvalidOperation, TypeError, ValueError):
        return hypothesis
    if unit_price <= 0:
        return hypothesis
    raw = json.loads(json.dumps(hypothesis, ensure_ascii=False))
    costs = [value for value in raw.get("costs", []) if isinstance(value, dict)]
    target = target if isinstance(target, dict) else None
    matching = None
    if target:
        matching = next((value for value in costs if _normalized_text(value.get("name")) == _normalized_text(target.get("name"))), None)
    if matching is None and source_url:
        matching = next((value for value in costs if _cell_text(value.get("source_url")) == source_url), None)
    category = (matching or target or {}).get("category")
    if category not in {"material", "application", "logistics"}:
        category = "material"
    route = raw.get("route") if isinstance(raw.get("route"), dict) else {}
    processes = _normalize_route_processes(route)
    process_name = _cell_text((matching or target or {}).get("process_name"))
    if not process_name:
        process_name = processes[0]["name"] if processes else "Изготовление под ключ"
    cost_name = _cell_text((target or matching or {}).get("name")) or _cell_text(quote.get("row_label")) or "Изготовление по прайсу поставщика"
    retained = []
    for cost in costs:
        same_target = target and _normalized_text(cost.get("name")) == _normalized_text(target.get("name"))
        same_source = source_url and _cell_text(cost.get("source_url")) == source_url
        if same_target or same_source:
            continue
        retained.append(cost)
    amount_total = _money(unit_price * quantity)
    retained.insert(0, {
        "category": category,
        "process_name": process_name,
        "name": cost_name,
        "amount_total": str(amount_total),
        "source": _cell_text(source_name)[:300],
        "source_type": "supplier",
        "source_url": _cell_text(source_url)[:1000],
        "source_date": timezone.localdate().isoformat(),
        "basis": f"{_decimal_text(quantity)} шт. × {_money(unit_price)} ₽/шт.; строка «{_cell_text(quote.get('row_label'))}»; диапазон «{_cell_text(quote.get('tier'))}»",
        "recipe": {"method": "unit_rate", "inputs": {"unit_rate": str(_money(unit_price))}},
        "calculation_steps": [],
        "adaptation": "Строка товара выбрана по характеристикам, а ценовой столбец — бэкендом по текущему тиражу из HTML-таблицы поставщика.",
        "confirmed": False,
    })
    raw["costs"] = retained
    raw["questions"] = [
        value for value in raw.get("questions", []) if "цен" not in _normalized_text(value)
    ] if isinstance(raw.get("questions"), list) else []
    change = f"Цена поставщика проверена по строке «{quote.get('row_label')}» и диапазону «{quote.get('tier')}»: {_money(unit_price)} ₽/шт."
    raw["understood_changes"] = [*_short_text_list(raw.get("understood_changes"), limit=7), change]
    production_types = list(ProductionType.objects.filter(is_active=True))
    normalized = _normalize_training_hypothesis(raw, line, production_types, raw.get("matched_example_ids", []))
    for key in ("catalog_candidates", "catalog_selection", "catalog_intent", "catalog_warning", "psodin_calculation", "sources", "usage"):
        if key in hypothesis:
            normalized[key] = hypothesis[key]
    normalized["production_types"] = [{"code": value.code, "name": value.name} for value in production_types]
    normalized["verified_source_quote"] = quote
    return _attach_memory_preview(normalized)


def _attach_memory_preview(hypothesis):
    preview = []
    intent = hypothesis.get("catalog_intent") if isinstance(hypothesis.get("catalog_intent"), dict) else {}
    if intent.get("product_class"):
        preview.append(f"Тип товара: {intent['product_class']}")
    route = hypothesis.get("route") if isinstance(hypothesis.get("route"), dict) else {}
    if route.get("name"):
        preview.append(f"Маршрут: {route['name']}")
    for cost in hypothesis.get("costs", []) if isinstance(hypothesis.get("costs"), list) else []:
        if not isinstance(cost, dict):
            continue
        recipe = cost.get("recipe") if isinstance(cost.get("recipe"), dict) else {}
        method = recipe.get("method") or "без формулы"
        preview.append(f"Цена: {cost.get('name', 'статья')} — {cost.get('amount_total', '0')} ₽; расчёт: {method}; источник: {cost.get('source') or 'не указан'}")
    for change in hypothesis.get("understood_changes", []) if isinstance(hypothesis.get("understood_changes"), list) else []:
        preview.append(f"Корректировка: {change}")
    hypothesis["memory_preview"] = preview[:12]
    return hypothesis


def build_training_hypothesis(line, current=None, feedback=""):
    started_at = time.perf_counter()
    from .models import ProductionType
    from .catalog import CatalogSyncError, catalog_candidates_for_line

    production_types = list(ProductionType.objects.filter(is_active=True))
    examples = _training_examples_for_line(line)
    knowledge_sources = _knowledge_sources_for_line(line)
    example_payload = [{
        "id": value.pk,
        "position": value.position_name,
        "type": value.production_type.code,
        "features": value.features,
        "approved_route": value.routes[0] if value.routes else {},
    } for value in examples]
    schema = '{"product_type":"digital_sheet","summary":"как понята позиция","confidence":0.5,"facts":["факт"],"route":{"reason":"почему выбран маршрут","processes":[{"name":"Закупка материала","details":["операции и характеристики внутри процесса"]}]},"costs":[{"process_name":"Закупка материала","category":"material|application|logistics","name":"статья расхода","amount_total":0,"source":"точное название справочника, расчёта, поставщика или записи истории","source_type":"calculator|catalog|supplier|history|manager","source_url":"https://... или пусто","source_date":"дата цены или пусто","basis":"краткая итоговая формула","recipe":{"method":"sheet_yield|unit_rate|fixed|history_scaled|none","inputs":{"unit_price":380,"units_per_sheet":4,"waste_percent":5},"modifiers":[{"type":"discount_percent|markup_percent|add_fixed|subtract_fixed","value":15}]},"calculation_steps":["исходный формат и цена","выход изделий с листа","число листов с браком","арифметика стоимости"],"adaptation":"как исходная цена адаптирована к текущему формату, тиражу и условиям","confirmed":false}],"questions":["только критичный вопрос"],"assumptions":["допущение"],"matched_example_ids":[1],"understood_changes":["как понята обратная связь"]}'
    schema = schema[:-1] + ',"psodin_calculation":{"requested":false,"calculator":"sheet","scope":"labour_only","process_name":"Работа PSODIN","productivity_per_hour":10,"tariff":"standard|regular|partner|urgent"}}'
    schema = schema[:-1] + ',"catalog_intent":{"product_class":"канонический тип товара в единственном числе","synonyms":["синоним из позиции"],"hard_constraints":["обязательное требование без вычислений"],"preferences":["желательное свойство"]}}'
    prompt = f"""Ты — ассистент администратора по расчёту тендеров. Предложи ровно ОДИН наиболее вероятный маршрут и его калькуляцию. Не строй дерево и не дроби производство на мелкие физические операции: шаг маршрута — крупный самостоятельно заказываемый блок (например, готовое изделие, нанесение, изготовление под ключ).
Маршрут описывай универсальными процессами по 2–5 слов: «Закупка материала», «Универсальная типография», «Закупка готового изделия», «Нанесение». Не включай в название процесса конкретный продукт, тираж, материал или перечень операций. Конкретные резку, биговку, печать, тиснение и характеристики перечисляй в details процесса. Логистика и другие дополнительные расходы не являются процессом маршрута, если администратор явно не сказал обратное.
«Закупка материала» используй только когда материал покупается отдельно и затем передаётся следующему исполнителю. Если один исполнитель сам предоставляет материал и выполняет весь заказ, это один производственный процесс «Цифровая типография под ключ», «Универсальная типография под ключ», «Швейное производство под ключ» и т. п. Не называй изготовление под ключ закупкой материала. Свой или сторонний исполнитель — атрибут конкретного предложения и источника цены, а не название процесса.
Не выдумывай цены. В costs добавляй только цену, явно указанную в подтверждённых примерах, текущей гипотезе или обратной связи администратора. amount_total — сумма статьи на весь тираж. Если цены нет, оставь её вопросом, а не нулевой выдуманной статьёй.
Для каждой статьи costs дай проверяемый след расчёта. В source укажи конкретный источник, в basis — итоговую формулу, а в calculation_steps — максимально подробную арифметику по шагам: исходную единицу и цену, раскладку/выход, требуемое количество с отходами, операции, скидки и итог. В adaptation объясни, как цена источника приведена к текущему тиражу, формату и характеристикам. Для калькулятора перечисли материалы и операции отдельно. Для истории или поставщика укажи исходный кейс/товар и все коэффициенты пересчёта. Не придумывай отсутствующие детали: если подробного основания нет, прямо напиши это в adaptation и задай вопрос администратору.
Если переносишь опыт подтверждённого примера, переноси его ПРАВИЛО и заново подставляй текущие параметры, а не копируй готовую сумму. Для воспроизводимых правил заполняй recipe: sheet_yield использует unit_price, units_per_sheet и waste_percent; unit_rate — unit_rate; fixed — fixed_amount; history_scaled — base_total и base_quantity. Скидки, наценки и фиксированные поправки передавай только в recipe.modifiers в порядке применения. Никогда не меняй amount_total самостоятельно из-за скидки: сервер пересчитает сумму и сам сформирует объяснение. amount_total должен соответствовать recipe.
Значения material_unit, application_unit и logistics_unit в ПОЗИЦИИ — ручные поля текущего расчёта, а не факты из ТЗ. Если используешь их, source_type=manager и source="Введено администратором". Нельзя писать «дано в ТЗ», если цена не находится внутри requirements с явным source.
Подтверждённые примеры важнее общих предположений. matched_example_ids указывай только для действительно похожих примеров. Без подтверждённого близкого примера confidence не выше 0.55.
ПРОВЕРЕННЫЕ ИСТОЧНИКИ ИЗ БАЗЫ — это кандидаты цен и предложений, а не готовый ответ. Используй только источник, характеристики которого подходят текущей позиции. В source пиши поставщика и название источника, в source_url — его ссылку. Если условия нельзя надёжно адаптировать, задай вопрос вместо выдумывания цены.
Если передана ОБРАТНАЯ СВЯЗЬ, обнови всю гипотезу и запиши в understood_changes краткий структурированный список того, что изменил. Не повторяй закрытые вопросы. Найденные в ТЗ факты не спрашивай повторно.
Калькулятор PSODIN реально доступен на бэкенде. Если администратор явно сказал, что работу делает PSODIN, заполни psodin_calculation. Не считай часы, скидку и сумму: это сделает бэкенд. Передай только явно названную администратором производительность в штуках в час и тариф. Не добавляй работу PSODIN в costs: сервер добавит её сам.
Для поиска готового товара заполни catalog_intent. Здесь работает только понимание слов: приведи название к каноническому типу в единственном числе (например, «майка брендированная» → «футболка», «жилетка» → «жилет»), выдели синонимы и раздели обязательные требования от пожеланий. Не подбирай артикулы, не сравнивай числа и ничего не рассчитывай — это выполнит бэкенд.
Верни только JSON: {schema}

ПОЗИЦИЯ:
{json.dumps(line, ensure_ascii=False)}

ТЕКУЩАЯ ГИПОТЕЗА:
{json.dumps(current or {}, ensure_ascii=False)}

ОБРАТНАЯ СВЯЗЬ АДМИНИСТРАТОРА:
{feedback or 'нет — это первая гипотеза'}

ТИПЫ ПРОДУКЦИИ:
{json.dumps([{"code": value.code, "name": value.name, "description": value.description} for value in production_types], ensure_ascii=False)}

ПОДТВЕРЖДЁННЫЕ ПРИМЕРЫ:
{json.dumps(example_payload, ensure_ascii=False)}

ПРОВЕРЕННЫЕ ИСТОЧНИКИ ИЗ БАЗЫ:
{json.dumps(knowledge_sources, ensure_ascii=False)}"""
    ai_started_at = time.perf_counter()
    result, usage = _ai_gateway_json(prompt, max_tokens=3600)
    ai_seconds = round(time.perf_counter() - ai_started_at, 3)
    valid_ids = {value.pk for value in examples}
    matched_ids = []
    for value in result.get("matched_example_ids", []) if isinstance(result.get("matched_example_ids"), list) else []:
        try:
            value = int(value)
        except (TypeError, ValueError):
            continue
        if value in valid_ids:
            matched_ids.append(value)
    hypothesis = _normalize_training_hypothesis(result, line, production_types, matched_ids)
    confirmed_psodin = next((
        route.get("psodin_calculation")
        for example in examples if example.pk in matched_ids
        for route in example.routes[:1] if isinstance(route, dict) and isinstance(route.get("psodin_calculation"), dict)
    ), None)
    hypothesis = _apply_psodin_calculation(hypothesis, result, line, current=current, feedback=feedback, confirmed=confirmed_psodin)
    raw_intent = result.get("catalog_intent") if isinstance(result.get("catalog_intent"), dict) else {}
    catalog_intent = {
        "product_class": _cell_text(raw_intent.get("product_class"))[:100],
        "synonyms": _short_text_list(raw_intent.get("synonyms"), limit=8),
        "hard_constraints": _short_text_list(raw_intent.get("hard_constraints"), limit=12),
        "preferences": _short_text_list(raw_intent.get("preferences"), limit=8),
    }
    catalog_started_at = time.perf_counter()
    try:
        catalog_candidates = catalog_candidates_for_line(line, limit=3, intent=catalog_intent)
    except CatalogSyncError as exc:
        catalog_candidates = []
        hypothesis["catalog_warning"] = str(exc)[:300]
    except Exception:
        # Oasis is an optional price source. A malformed external row or a
        # temporary API problem must not discard the already valid LLM route.
        logger.exception("Unexpected Oasis catalog failure while building a training hypothesis")
        catalog_candidates = []
        hypothesis["catalog_warning"] = "Не удалось проверить каталог Oasis. Маршрут сохранён без цены поставщика."
    catalog_seconds = round(time.perf_counter() - catalog_started_at, 3)
    hypothesis["catalog_intent"] = catalog_intent
    hypothesis["catalog_candidates"] = catalog_candidates
    if isinstance(current, dict) and isinstance(current.get("catalog_selection"), dict):
        selected_id = current["catalog_selection"].get("id")
        if any(value.get("id") == selected_id and value.get("fit") == "exact" for value in catalog_candidates):
            hypothesis["catalog_selection"] = current["catalog_selection"]
    if isinstance(current, dict) and isinstance(current.get("sources"), list):
        hypothesis["sources"] = current["sources"][:20]
    hypothesis["usage"] = {"prompt_tokens": usage.get("prompt_tokens", 0), "completion_tokens": usage.get("completion_tokens", 0)}
    hypothesis["timings"] = {
        "ai_seconds": ai_seconds,
        "catalog_seconds": catalog_seconds,
        "total_seconds": round(time.perf_counter() - started_at, 3),
    }
    hypothesis["production_types"] = [{"code": value.code, "name": value.name} for value in production_types]
    # A fully matching live offer is an executable backend price source, not
    # merely a visual suggestion. Apply it immediately so the displayed total
    # and the tender material field cannot remain zero while showing a product.
    exact_candidate = next((
        value for value in catalog_candidates
        if value.get("fit") == "exact" and value.get("price") not in (None, "")
    ), None)
    if exact_candidate and not hypothesis.get("catalog_selection"):
        preserved_usage = hypothesis.get("usage", {})
        preserved_warning = hypothesis.get("catalog_warning")
        hypothesis = apply_catalog_candidate(hypothesis, line, exact_candidate.get("id"))
        hypothesis["catalog_selection"]["selection_mode"] = "automatic"
        hypothesis["usage"] = preserved_usage
        if preserved_warning:
            hypothesis["catalog_warning"] = preserved_warning
    return _attach_memory_preview(hypothesis)


def apply_catalog_candidate(hypothesis, line, product_id):
    from .models import ProductionType

    candidates = hypothesis.get("catalog_candidates", []) if isinstance(hypothesis, dict) else []
    try:
        candidate = next(value for value in candidates if str(value.get("id")) == str(product_id))
    except (StopIteration, TypeError, ValueError):
        raise TenderAIError("Товар больше не входит в актуальную подборку. Обновите гипотезу.")
    supplier_code = _cell_text(candidate.get("supplier_code")).lower()
    supplier_name = _cell_text(candidate.get("supplier_name"))[:200] or {"oasis": "Oasis"}.get(supplier_code, supplier_code or "Поставщик")
    supplier_site = _cell_text(candidate.get("supplier_site"))[:200]
    supplier_label = " · ".join(value for value in (supplier_name, supplier_site) if value)
    try:
        quantity = max(Decimal("1"), Decimal(str(line.get("quantity", 1)).replace(",", ".")))
    except (InvalidOperation, TypeError, ValueError):
        raise TenderAIError("Не удалось определить количество для товара поставщика.")
    try:
        price = Decimal(str(candidate.get("price")))
    except (InvalidOperation, TypeError, ValueError):
        raise TenderAIError("У товара поставщика больше нет актуальной цены.")
    raw = json.loads(json.dumps(hypothesis, ensure_ascii=False)) if isinstance(hypothesis, dict) else {}
    costs = [
        value for value in raw.get("costs", [])
        if isinstance(value, dict) and not (value.get("category") == "material" and value.get("source_type") in {"catalog", "supplier", "history", "manager"})
    ]
    costs.insert(0, {
        "category": "material",
        "process_name": "Закупка готового изделия",
        "name": candidate.get("name") or "Товар поставщика",
        "amount_total": str(_money(price * quantity)),
        "source": f"{supplier_label} · арт. {candidate.get('article', '')}",
        "source_type": "catalog",
        "source_url": candidate.get("url", ""),
        "source_date": timezone.localdate().isoformat(),
        "basis": f"{_decimal_text(quantity)} шт. × {_money(price)} ₽/шт.",
        "recipe": {"method": "unit_rate", "inputs": {"unit_rate": str(price)}},
        "calculation_steps": [],
        "adaptation": f"Цена и свободный остаток получены из каталога поставщика {supplier_name}; перед закупкой требуется повторная проверка актуальности.",
        "confirmed": False,
    })
    raw["costs"] = costs
    route = raw.get("route") if isinstance(raw.get("route"), dict) else {}
    processes = route.get("processes") if isinstance(route.get("processes"), list) else []
    processes = [value for value in processes if isinstance(value, dict) and _canonical_process_name(value.get("name")) not in {"Закупка материала", "Закупка готового изделия"}]
    route["processes"] = [{"name": "Закупка готового изделия", "details": [f"{supplier_name}, арт. {candidate.get('article', '')}"]}, *processes]
    if candidate.get("fit") == "exact":
        route["reason"] = f"Готовое изделие найдено у поставщика {supplier_name} и соответствует проверенным требованиям ТЗ. Его актуальная цена автоматически включена в закупочную себестоимость; нанесение считается отдельным процессом."
    else:
        mismatch_text = "; ".join(_short_text_list(candidate.get("mismatches"), limit=3))
        route["reason"] = f"Товар поставщика {supplier_name} выбран администратором как рабочая альтернатива. Расхождения, которые нужно учитывать: {mismatch_text or 'часть характеристик требует проверки'}."
    raw["route"] = route
    raw["questions"] = [
        value for value in raw.get("questions", []) if not (
            "цен" in _normalized_text(value)
            and any(marker in _normalized_text(value) for marker in ("закуп", "готов", "товар", "майк", "футбол", "oasis", "поставщик"))
        )
    ] if isinstance(raw.get("questions"), list) else []
    production_types = list(ProductionType.objects.filter(is_active=True))
    normalized = _normalize_training_hypothesis(raw, line, production_types, raw.get("matched_example_ids", []))
    normalized["catalog_candidates"] = candidates[:3]
    normalized["catalog_selection"] = {
        **candidate,
        "price": str(price),
        "cost_total": str(_money(price * quantity)),
        "selection_mode": "manual",
        "accepted_mismatches": candidate.get("mismatches", []) if candidate.get("fit") != "exact" else [],
        "selected_at": timezone.now().isoformat(),
    }
    normalized["production_types"] = [{"code": value.code, "name": value.name} for value in production_types]
    if isinstance(hypothesis, dict) and isinstance(hypothesis.get("catalog_intent"), dict):
        normalized["catalog_intent"] = hypothesis["catalog_intent"]
    existing_sources = hypothesis.get("sources", []) if isinstance(hypothesis, dict) and isinstance(hypothesis.get("sources"), list) else []
    normalized["sources"] = [
        value for value in existing_sources
        if isinstance(value, dict) and value.get("source_type") != "catalog"
    ][:19]
    normalized["sources"].append({
        "source_type": "catalog",
        "supplier_name": supplier_name,
        "title": candidate.get("name") or "Товар поставщика",
        "url": candidate.get("url", ""),
        "article": candidate.get("article", ""),
        "price": str(price),
        "cost_name": candidate.get("name") or "Закупка готового изделия",
        "is_pending": False,
    })
    if isinstance(hypothesis, dict) and isinstance(hypothesis.get("usage"), dict):
        normalized["usage"] = hypothesis["usage"]
    if isinstance(hypothesis, dict) and hypothesis.get("catalog_warning"):
        normalized["catalog_warning"] = _cell_text(hypothesis.get("catalog_warning"))[:300]
    return _attach_memory_preview(normalized)


def analyze_production_route(line):
    from calculator.models import PriceItem

    items = list(PriceItem.objects.filter(is_active=True).select_related("base_item", "production_rule"))
    catalog = [
        {"id": item.pk, "category": item.category, "name": item.name, "aliases": item.aliases, "unit": item.unit_name, "price": str(item.effective_unit_price)}
        for item in items
    ]
    schema = '{"product_class":"тип продукции","source_candidates":[{"source":"psodin_sheet|psodin_canon|supplier_price|supplier_api|history|white_site|open_web|manager","title":"название источника","priority":1,"fit":"high|medium|low","reason":"почему"}],"selected_source":"psodin_sheet|psodin_canon|supplier_price|supplier_api|history|white_site|open_web|manager|unknown","calculator":"sheet|canon|none","route":"internal|hybrid|outsourcing|unknown","confidence":0.8,"reason":"почему выбран источник","components":[{"name":"Компонент","source":"internal|outsourcing|unknown","source_reason":"почему","kind":"sheet|operation|material","finished_width_mm":148,"finished_height_mm":210,"units_per_product":60,"material_query":"офсетная бумага","grammage_gsm":80,"print_required":false,"bleed_mm":0,"operation_item_ids":[]}],"questions":[],"warnings":[]}'
    prompt = f"""Подбери источники коммерческого предложения для позиции. Цель — найти наиболее выгодный надёжный вариант, а не обязательно изготовить изделие у нас.

Сначала классифицируй изделие, затем составь source_candidates в порядке полезности:
1) точный загруженный прайс или специализированный калькулятор;
2) подтверждённая история расчётов;
3) API поставщика;
4) проверенные сайты поставщиков;
5) открытый интернет;
6) уточнение у менеджера.
Перечисляй только реально подходящие источники. Не придумывай наличие цены или интеграции.

«Печатный салон №1 · psodin.ru» — один обычный подрядчик среди остальных:
- psodin_sheet подходит только для цифровой листовой печати и постпечатки, когда изделие технологически укладывается в лист SRA3; типичные признаки — небольшой тираж, листовая полиграфия без сложной фигурной вырубки и сборки;
- psodin_canon подходит только для широкоформатной рулонной печати;
- пакет, твёрдый переплёт, сложная фигурная вырубка, офсетный/промышленный тираж или готовый сувенир нельзя уверенно отправлять в psodin_sheet только потому, что в описании есть бумага или печать.
Если признаков недостаточно, selected_source=manager или unknown, calculator=none и задай один короткий вопрос, который поможет выбрать источник.

Компоненты и детальный расчёт по каталогу создавай только если selected_source=psodin_sheet или psodin_canon. Для листового компонента укажи чистовой размер и количество листов/заготовок на конечное изделие. Не подменяй формат готового изделия форматом исходного листа.
Если материал или операция отсутствуют в каталоге, не выбирай похожее молча: сохрани описание и добавь короткий вопрос. operation_item_ids может содержать только существующие ID из каталога.
В позиции могут быть manager_answers — ответы опытного менеджера. Они приоритетны: примени их и не повторяй уже закрытый вопрос. Ответ может быть названием позиции каталога, ссылкой или свободным описанием цены и параметров.
Не рассчитывай раскладку и стоимость — это сделает программа детерминированно.
Верни только JSON строго такого вида: {schema}

ПОЗИЦИЯ:
{json.dumps(line, ensure_ascii=False)}

НАШ КАТАЛОГ МАТЕРИАЛОВ И ОПЕРАЦИЙ:
{json.dumps(catalog, ensure_ascii=False)}"""
    result, usage = _ai_gateway_json(prompt, max_tokens=3200)
    route = result.get("route") if result.get("route") in {"internal", "hybrid", "outsourcing", "unknown"} else "unknown"
    calculator = result.get("calculator") if result.get("calculator") in {"sheet", "canon", "none"} else "none"
    valid_sources = {"psodin_sheet", "psodin_canon", "supplier_price", "supplier_api", "history", "white_site", "open_web", "manager", "unknown"}
    selected_source = result.get("selected_source") if result.get("selected_source") in valid_sources else ({"sheet": "psodin_sheet", "canon": "psodin_canon"}.get(calculator, "unknown"))
    if selected_source not in {"psodin_sheet", "psodin_canon"}:
        calculator = "none"
    try:
        confidence = max(0, min(1, float(result.get("confidence", 0))))
    except (TypeError, ValueError):
        confidence = 0
    paper_items = [item for item in items if item.category == PriceItem.CATEGORY_PAPER]
    components = []
    for raw in result.get("components", [])[:20]:
        if not isinstance(raw, dict):
            continue
        component = {
            "name": _cell_text(raw.get("name"))[:200] or "Компонент",
            "source": raw.get("source") if raw.get("source") in {"internal", "outsourcing", "unknown"} else ("internal" if route == "internal" else "unknown"),
            "source_reason": _cell_text(raw.get("source_reason"))[:300],
            "kind": raw.get("kind") if raw.get("kind") in {"sheet", "operation", "material"} else "material",
            "finished_width_mm": raw.get("finished_width_mm"),
            "finished_height_mm": raw.get("finished_height_mm"),
            "units_per_product": raw.get("units_per_product"),
            "material_query": _cell_text(raw.get("material_query"))[:300],
            "grammage_gsm": raw.get("grammage_gsm"),
            "print_required": bool(raw.get("print_required")),
            "bleed_mm": raw.get("bleed_mm") or 0,
        }
        valid_ids = []
        for value in raw.get("operation_item_ids", []) if isinstance(raw.get("operation_item_ids"), list) else []:
            try:
                item_id = int(value)
            except (TypeError, ValueError):
                continue
            if any(item.pk == item_id for item in items):
                valid_ids.append(item_id)
        component["operations"] = [entry for entry in catalog if entry["id"] in valid_ids]
        component["paper_candidates"] = _paper_candidates(component, line.get("quantity") or 0, paper_items) if component["kind"] == "sheet" and component["source"] == "internal" else []
        components.append(component)
    questions = _short_text_list(result.get("questions"), limit=6)
    manager_answers = line.get("manager_answers") or line.get("requirements", {}).get("manager_answers") or {}
    source_text = json.dumps(line, ensure_ascii=False).lower()
    cost_options = []

    # ИИ предлагает маршрут, но известные операции проверяем по устойчивым
    # технологическим синонимам. Так «горячее тиснение» не потеряет категорию
    # калькулятора, которая коротко называется «Тиснение».
    internal_allowed = selected_source in {"psodin_sheet", "psodin_canon"} and any(component["source"] == "internal" for component in components)
    if internal_allowed and any(value in source_text for value in ("тиснен", "фольгир", "горячая фольга")):
        for item in items:
            if item.category != PriceItem.CATEGORY_EMBOSSING:
                continue
            quantity = Decimal("1") if "прилад" in item.name.lower() else Decimal(str(line.get("quantity") or 0))
            total = quantity * item.effective_unit_price
            cost_options.append({
                "group": "embossing",
                "mode": "required",
                "catalog_item_id": item.pk,
                "name": item.name,
                "calculation": "однократно" if quantity == 1 else f"{quantity} × {item.effective_unit_price} ₽",
                "total_cost": str(_money(total)),
            })

    if internal_allowed and "пружин" in source_text:
        quantity = Decimal(str(line.get("quantity") or 0))
        finished_edges = []
        for component in components:
            try:
                width = Decimal(str(component.get("finished_width_mm") or 0))
                height = Decimal(str(component.get("finished_height_mm") or 0))
            except (InvalidOperation, TypeError, ValueError):
                continue
            if width > 0 and height > 0:
                finished_edges.append(min(width, height))
        binding_mm = min(finished_edges) if finished_edges else Decimal("0")
        spring_items = [item for item in items if item.category == PriceItem.CATEGORY_POSTPRESS and "пружин" in item.name.lower()]
        for item in spring_items:
            rule = getattr(item, "production_rule", None)
            if rule and rule.calculation_kind == "linear" and rule.package_quantity > 0 and binding_mm > 0:
                required_m = binding_mm / Decimal("1000") * quantity * (Decimal("1") + rule.waste_percent / Decimal("100"))
                packages = math.ceil(required_m / rule.package_quantity)
                total = Decimal(packages) * item.effective_unit_price
                calculation = f"{required_m.quantize(Decimal('0.01'))} м → {packages} уп."
            else:
                # Текущие позиции каталога заданы отрезками около 30 см.
                # Показываем геометрический выход, но оставляем подтверждение технологу.
                pieces_per_segment = max(1, math.floor(Decimal("300") / binding_mm)) if binding_mm else 1
                segments = math.ceil(quantity / pieces_per_segment)
                total = Decimal(segments) * item.effective_unit_price
                calculation = f"{pieces_per_segment} шт. с 30 см · {segments} отрезков"
            cost_options.append({
                "group": "spring",
                "mode": "alternative",
                "catalog_item_id": item.pk,
                "name": item.name,
                "calculation": calculation,
                "total_cost": str(_money(total)),
            })
        spring_answered = any("пружин" in str(question).lower() and str(answer).strip() for question, answer in manager_answers.items())
        if spring_items and not spring_answered:
            questions.append("Подтвердите тип пружины и допустимо ли получать две пружины А5 из отрезка 30 см.")

    questions = [question for question in questions if not str(manager_answers.get(question, "")).strip()]
    source_candidates = []
    for raw in result.get("source_candidates", [])[:8]:
        if not isinstance(raw, dict) or raw.get("source") not in valid_sources:
            continue
        source_candidates.append({
            "source": raw["source"],
            "title": _cell_text(raw.get("title"))[:120] or raw["source"],
            "fit": raw.get("fit") if raw.get("fit") in {"high", "medium", "low"} else "medium",
            "reason": _cell_text(raw.get("reason"))[:300],
        })
    return {
        "product_class": _cell_text(result.get("product_class"))[:200],
        "selected_source": selected_source,
        "source_candidates": source_candidates,
        "route": route,
        "calculator": calculator,
        "confidence": confidence,
        "reason": _cell_text(result.get("reason"))[:700],
        "components": components,
        "cost_options": cost_options,
        "questions": list(dict.fromkeys(questions))[:8],
        "manager_answers": manager_answers,
        "catalog_choices": [{"id": item.pk, "name": item.name, "category": item.get_category_display()} for item in items],
        "warnings": _short_text_list(result.get("warnings"), limit=6),
        "assumptions": ["Раскладка геометрическая, с поворотом заготовки.", "Технологические отходы: 3%.", "Направление волокна, поля оборудования и порядок реза должен подтвердить технолог."],
        "usage": {"prompt_tokens": usage.get("prompt_tokens", 0), "completion_tokens": usage.get("completion_tokens", 0)},
    }


def _money(value):
    return Decimal(value).quantize(MONEY, rounding=ROUND_HALF_UP)


def calculate_tender(lines, reduction_percent, russia_delivery, vat_rate):
    coefficient = (Decimal("100") - reduction_percent) / Decimal("100")
    totals = {"nmck_total": Decimal("0"), "rrp_total": Decimal("0"), "purchase_total": Decimal("0"), "gross_profit": Decimal("0")}
    calculated_lines = []
    for line in lines:
        nmck_total = line["quantity"] * line["nmck_unit"]
        rrp_unit = line["nmck_unit"] * coefficient
        rrp_total = line["quantity"] * rrp_unit
        purchase_unit = line["material_unit"] + line["application_unit"] + line["logistics_unit"]
        purchase_total = line["quantity"] * purchase_unit
        profit = rrp_total - purchase_total
        calculated_lines.append({"nmck_total": _money(nmck_total), "rrp_unit": _money(rrp_unit), "rrp_total": _money(rrp_total), "purchase_unit": _money(purchase_unit), "purchase_total": _money(purchase_total), "profit": _money(profit)})
        totals["nmck_total"] += nmck_total
        totals["rrp_total"] += rrp_total
        totals["purchase_total"] += purchase_total
        totals["gross_profit"] += profit
    vat = totals["rrp_total"] * vat_rate / Decimal("100")
    all_expenses = totals["purchase_total"] + russia_delivery + vat
    net_profit = totals["rrp_total"] - all_expenses
    roi = net_profit / all_expenses * Decimal("100") if all_expenses else Decimal("0")
    summary = {**{key: _money(value) for key, value in totals.items()}, "vat": _money(vat), "all_expenses": _money(all_expenses), "net_profit": _money(net_profit), "roi": roi.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP), "coefficient": coefficient.quantize(Decimal("0.0001"))}
    return calculated_lines, summary
